
# BAB IV DOCX Generator - Part 2 (4.5 - 4.8)
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document(os.path.join(os.path.dirname(__file__), '_bab4_temp.docx'))

def h2(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(text); r.bold = True; r.font.name = 'Times New Roman'; r.font.size = Pt(12); r.font.color.rgb = RGBColor(0,0,0)
    p.paragraph_format.line_spacing = 1.5; p.paragraph_format.space_before = Pt(18)

def h3(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(text); r.bold = True; r.font.name = 'Times New Roman'; r.font.size = Pt(12); r.font.color.rgb = RGBColor(0,0,0)
    p.paragraph_format.line_spacing = 1.5; p.paragraph_format.space_before = Pt(12)

def ap(text, indent=True):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY; p.paragraph_format.line_spacing = 1.5
    if indent: p.paragraph_format.first_line_indent = Cm(1.27)
    r = p.add_run(text); r.font.name = 'Times New Roman'; r.font.size = Pt(12); r.font.color.rgb = RGBColor(0,0,0)

def ap_rich(parts, indent=True):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY; p.paragraph_format.line_spacing = 1.5
    if indent: p.paragraph_format.first_line_indent = Cm(1.27)
    for text, style in parts:
        r = p.add_run(text); r.font.name = 'Times New Roman'; r.font.size = Pt(12); r.font.color.rgb = RGBColor(0,0,0)
        if style == 'bold': r.bold = True
        elif style == 'italic': r.italic = True

def li(text, num=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY; p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.left_indent = Cm(1.27); p.paragraph_format.first_line_indent = Cm(0)
    prefix = f"{num}. " if num else "\u2022 "
    parts = (prefix + text).split('**')
    for i, part in enumerate(parts):
        r = p.add_run(part); r.font.name = 'Times New Roman'; r.font.size = Pt(12); r.font.color.rgb = RGBColor(0,0,0)
        if i % 2 == 1: r.bold = True

def tbl(headers, rows):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = 'Table Grid'; table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = table.rows[0].cells[i]; c.text = ''; p = c.paragraphs[0]
        r = p.add_run(h); r.bold = True; r.font.name = 'Times New Roman'; r.font.size = Pt(11); r.font.color.rgb = RGBColor(0,0,0)
        p.paragraph_format.line_spacing = 1.5
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            c = table.rows[ri+1].cells[ci]; c.text = ''; p = c.paragraphs[0]
            r = p.add_run(val); r.font.name = 'Times New Roman'; r.font.size = Pt(11); r.font.color.rgb = RGBColor(0,0,0)
            p.paragraph_format.line_spacing = 1.5
    doc.add_paragraph()

# ===== 4.5 =====
h2('4.5 Implementasi Basis Data')
ap('Struktur basis data Sistem ERP Prototype dibangun menggunakan fitur migration pada Laravel yang memungkinkan pembuatan dan pengelolaan skema basis data secara terprogram dan terversioning. Basis data menggunakan MySQL dengan lima tabel utama yang saling berelasi. Berikut adalah penjelasan detail setiap tabel beserta relasinya.')

h3('4.5.1 Tabel Roles')
ap('Tabel roles berfungsi menyimpan data hak akses pengguna dalam sistem. Tabel ini menjadi referensi utama untuk menentukan level akses setiap pengguna. Terdapat tiga role yang dikonfigurasi melalui RoleSeeder, yaitu: super_admin, admin, dan karyawan.')
tbl(['Kolom','Tipe Data','Keterangan'],[
    ['id','bigint unsigned (PK, AI)','Primary key, auto increment'],
    ['nama_role','varchar (unique)','Nama role: super_admin, admin, karyawan'],
    ['created_at','timestamp','Waktu pembuatan record'],
    ['updated_at','timestamp','Waktu pembaruan record'],
])

h3('4.5.2 Tabel Divisions')
ap('Tabel divisions menyimpan data divisi atau departemen yang ada di perusahaan. Tabel ini digunakan untuk mengelompokkan pengguna berdasarkan unit kerja. Terdapat delapan divisi yang dikonfigurasi melalui DivisionSeeder, yaitu: Marketing, Sales, Keuangan, Konten Kreator, Gudang, CRM, YouTube, dan Admin Marketplace.')
tbl(['Kolom','Tipe Data','Keterangan'],[
    ['id','bigint unsigned (PK, AI)','Primary key, auto increment'],
    ['nama_divisi','varchar (unique)','Nama divisi perusahaan'],
    ['created_at','timestamp','Waktu pembuatan record'],
    ['updated_at','timestamp','Waktu pembaruan record'],
])

h3('4.5.3 Tabel Users')
ap('Tabel users merupakan tabel inti yang menyimpan data seluruh pengguna sistem. Tabel ini memiliki relasi foreign key ke tabel roles dan divisions untuk menentukan hak akses dan unit kerja setiap pengguna.')
tbl(['Kolom','Tipe Data','Keterangan'],[
    ['id','bigint unsigned (PK, AI)','Primary key, auto increment'],
    ['name','varchar(255)','Nama lengkap pengguna'],
    ['email','varchar(255) (unique)','Alamat email sebagai identitas login'],
    ['email_verified_at','timestamp (nullable)','Waktu verifikasi email'],
    ['password','varchar(255) (hashed)','Kata sandi terenkripsi bcrypt'],
    ['role_id','bigint unsigned (FK)','Foreign key ke tabel roles'],
    ['division_id','bigint unsigned (FK, nullable)','Foreign key ke tabel divisions'],
    ['remember_token','varchar(100)','Token untuk fitur remember me'],
    ['created_at','timestamp','Waktu pembuatan akun'],
    ['updated_at','timestamp','Waktu pembaruan data akun'],
])

h3('4.5.4 Tabel Activities')
ap('Tabel activities berfungsi mencatat seluruh log aktivitas yang dilakukan pengguna dalam sistem. Tabel ini memiliki foreign key ke tabel users dengan constraint cascade on delete, yang berarti apabila data user dihapus maka seluruh aktivitas terkait juga akan terhapus secara otomatis.')
tbl(['Kolom','Tipe Data','Keterangan'],[
    ['id','bigint unsigned (PK, AI)','Primary key, auto increment'],
    ['user_id','bigint unsigned (FK, cascade delete)','Foreign key ke tabel users'],
    ['tanggal','datetime','Tanggal dan waktu aktivitas dilakukan'],
    ['deskripsi','text','Deskripsi detail aktivitas'],
    ['status','varchar (default: submitted)','Status aktivitas (submitted, created, dll)'],
    ['created_at','timestamp','Waktu pembuatan record'],
    ['updated_at','timestamp','Waktu pembaruan record'],
])

h3('4.5.5 Tabel Reports')
ap('Tabel reports menyimpan data laporan yang dibuat oleh Super Admin melalui modul Reports Center. Tabel ini mendukung empat jenis laporan yaitu financial, department, growth, dan custom.')
tbl(['Kolom','Tipe Data','Keterangan'],[
    ['id','bigint unsigned (PK, AI)','Primary key, auto increment'],
    ['title','varchar(255)','Judul laporan'],
    ['type','varchar','Jenis laporan (financial, department, growth, custom)'],
    ['description','text (nullable)','Deskripsi atau isi laporan'],
    ['start_date','date (nullable)','Tanggal mulai periode laporan'],
    ['end_date','date (nullable)','Tanggal akhir periode laporan'],
    ['status','enum (generated, ready, failed)','Status pemrosesan laporan'],
    ['created_by','bigint unsigned (FK, nullable)','Foreign key ke tabel users (pembuat)'],
    ['created_at','timestamp','Waktu pembuatan laporan'],
    ['updated_at','timestamp','Waktu pembaruan laporan'],
])

h3('4.5.6 Relasi Antar Tabel')
ap('Seluruh relasi antar tabel dalam sistem ERP ini menggunakan pola Eloquent Relationship pada Laravel. Hubungan antar entitas didefinisikan dalam model masing-masing menggunakan method belongsTo() dan hasMany(). Berikut adalah ringkasan relasi antar tabel:')
tbl(['Relasi','Tipe','Penjelasan'],[
    ['Role \u2192 Users','One-to-Many (hasMany)','Satu role dapat dimiliki oleh banyak user'],
    ['Division \u2192 Users','One-to-Many (hasMany)','Satu divisi dapat memiliki banyak user'],
    ['User \u2192 Role','Many-to-One (belongsTo)','Setiap user memiliki satu role'],
    ['User \u2192 Division','Many-to-One (belongsTo)','Setiap user memiliki satu divisi (nullable)'],
    ['User \u2192 Activities','One-to-Many (hasMany)','Satu user dapat memiliki banyak aktivitas'],
    ['Activity \u2192 User','Many-to-One (belongsTo)','Setiap aktivitas dimiliki oleh satu user'],
    ['User \u2192 Reports','One-to-Many','Satu user dapat membuat banyak laporan'],
    ['Report \u2192 User','Many-to-One (belongsTo)','Setiap laporan dibuat oleh satu user'],
])
ap('Seluruh foreign key diimplementasikan dengan constraint yang sesuai. Foreign key user_id pada tabel activities menggunakan cascadeOnDelete() sehingga data aktivitas otomatis terhapus ketika user dihapus. Sementara foreign key created_by pada tabel reports menggunakan nullOnDelete() sehingga data laporan tetap tersimpan meskipun user pembuat dihapus.')

# ===== 4.6 =====
h2('4.6 Implementasi Backend (MVC Laravel)')

h3('4.6.1 Arsitektur Model-View-Controller')
ap('Sistem ERP Prototype ini mengimplementasikan pola arsitektur Model-View-Controller (MVC) yang merupakan standar bawaan framework Laravel. Pola MVC memisahkan logika aplikasi menjadi tiga komponen utama yang saling terhubung namun memiliki tanggung jawab yang berbeda:')
li('**Model** \u2014 Terdapat lima model utama yaitu User, Role, Division, Activity, dan Report. Setiap model mendefinisikan fillable attributes untuk mass assignment protection, casting untuk konversi tipe data otomatis, dan relasi Eloquent untuk menghubungkan antar entitas.', 1)
li('**View** \u2014 Menggunakan Blade Template Engine dengan sistem layout inheritance. Layout utama layouts/admin.blade.php menjadi template induk untuk seluruh halaman panel admin. TailwindCSS 4.0 dan Alpine.js digunakan untuk styling dan interaktivitas.', 2)
li('**Controller** \u2014 Terdapat delapan controller yang menangani logika bisnis, meliputi AdminDashboardController, UserController, ActivityController, ReportController, LoginController, ProfileController, DashboardController, dan AuthenticatedSessionController.', 3)

h3('4.6.2 Sistem Autentikasi')
ap('Proses autentikasi diimplementasikan melalui LoginController dengan alur sebagai berikut: (1) Pengguna mengakses halaman login melalui route GET /login; (2) Pengguna mengisi email dan password, kemudian mengirim form melalui route POST /login; (3) Sistem memvalidasi input dengan aturan required dan email; (4) Auth::attempt() memverifikasi kredensial terhadap data di tabel users; (5) Jika valid, session diregenerasi untuk keamanan dan pengguna diarahkan ke route /dashboard; (6) Route /dashboard memeriksa role pengguna dan melakukan redirect ke dashboard yang sesuai.')

h3('4.6.3 Sistem Routing')
ap('Routing sistem dikonfigurasi dalam file routes/web.php dengan pengelompokan berdasarkan middleware dan prefix. Terdapat empat kelompok route utama:')
tbl(['Kelompok Route','Prefix','Middleware','Keterangan'],[
    ['Auth Routes','/login','guest','Route untuk halaman login (GET dan POST)'],
    ['Admin Routes','/admin','auth','Dashboard, Users, Activities, Reports'],
    ['Karyawan Routes','/karyawan','auth','Dashboard karyawan'],
    ['Utility Routes','/','-','Profile, Logout, Root redirect'],
])

h3('4.6.4 Fitur Dark Mode')
ap('Sistem mendukung fitur dark mode yang diimplementasikan menggunakan Alpine.js dengan penyimpanan preferensi di localStorage browser. TailwindCSS dikonfigurasi dengan darkMode: \'class\' untuk mendukung styling kondisional berdasarkan tema aktif. Fitur ini meningkatkan kenyamanan pengguna dalam menggunakan sistem.')

# ===== 4.7 =====
h2('4.7 Pengujian Sistem')
ap_rich([
    ('Pengujian sistem dilakukan menggunakan metode ', 'normal'),
    ('Black Box Testing', 'italic'),
    (', yaitu metode pengujian yang berfokus pada fungsionalitas sistem tanpa memperhatikan struktur internal kode program. Pengujian dilakukan untuk memastikan bahwa setiap fitur sistem berjalan sesuai dengan kebutuhan yang telah didefinisikan pada tahap perancangan.', 'normal'),
])

h3('4.7.1 Hasil Pengujian Black Box')
ap('Berikut adalah tabel hasil pengujian Black Box Testing pada Sistem ERP Prototype:', indent=False)
tbl(['No','Skenario Pengujian','Langkah Pengujian','Hasil yang Diharapkan','Status'],[
    ['1','Login Valid','Memasukkan email dan password yang benar, klik Sign In','Pengguna berhasil login dan diarahkan ke dashboard sesuai role','Berhasil'],
    ['2','Login Tidak Valid','Memasukkan email atau password yang salah, klik Sign In','Muncul pesan error "Email atau password salah"','Berhasil'],
    ['3','Redirect Super Admin','Login menggunakan akun Super Admin','Diarahkan ke halaman /admin/dashboard','Berhasil'],
    ['4','Redirect Karyawan','Login menggunakan akun Karyawan','Diarahkan ke halaman /karyawan/dashboard','Berhasil'],
    ['5','Tampilan Dashboard Admin','Mengakses halaman /admin/dashboard','Menampilkan KPI, grafik analytics, dan quick actions','Berhasil'],
    ['6','Menambahkan User Baru','Klik Add User, mengisi form lengkap, klik submit','User baru tersimpan dan muncul di daftar pengguna','Berhasil'],
    ['7','Validasi Form User','Submit form tambah user dengan field kosong','Muncul pesan validasi error pada field yang kosong','Berhasil'],
    ['8','Pencarian User','Mengetik nama/email pada kolom search','Tabel menampilkan hasil pencarian yang sesuai','Berhasil'],
    ['9','Reset Password User','Klik Reset Password pada user, konfirmasi reset','Password user berhasil direset ke default','Berhasil'],
    ['10','Monitoring Aktivitas','Mengakses halaman /admin/activities','Menampilkan log aktivitas dengan KPI dan paginasi','Berhasil'],
    ['11','Generate Report','Memilih jenis report, mengisi form, submit','Report berhasil disimpan dan tampil di daftar','Berhasil'],
    ['12','Melihat Detail Report','Klik View pada report di tabel recent reports','Menampilkan halaman detail report','Berhasil'],
    ['13','Akses Halaman Profil','Klik avatar pengguna, pilih Profile','Menampilkan informasi akun dan ringkasan aktivitas','Berhasil'],
    ['14','Toggle Dark Mode','Klik tombol toggle Dark/Light di header','Tampilan berubah sesuai tema yang dipilih','Berhasil'],
    ['15','Logout','Klik tombol Logout pada dropdown profil','Session dihapus dan diarahkan kembali ke halaman login','Berhasil'],
    ['16','Akses Tanpa Login','Mengakses URL /admin/dashboard tanpa login','Diarahkan kembali ke halaman login','Berhasil'],
    ['17','Duplikasi Email','Menambahkan user dengan email yang sudah terdaftar','Muncul pesan error validasi email unique','Berhasil'],
    ['18','Paginasi Data','Mengakses halaman activities dengan data >15','Data ditampilkan dengan pagination 15 per halaman','Berhasil'],
])

h3('4.7.2 Pengujian Validasi Input')
ap_rich([
    ('Sistem menerapkan validasi input pada sisi server (', 'normal'),
    ('server-side validation', 'italic'),
    (') menggunakan fitur Request Validation bawaan Laravel. Validasi ini memastikan bahwa data yang dikirimkan oleh pengguna memenuhi aturan yang telah ditetapkan sebelum diproses lebih lanjut. Berikut adalah aturan validasi yang diterapkan pada setiap form dalam sistem:', 'normal'),
])
tbl(['Form','Field','Aturan Validasi'],[
    ['Login','email','required, email'],
    ['Login','password','required'],
    ['Tambah User','name','required, string, max:255'],
    ['Tambah User','email','required, email, unique:users'],
    ['Tambah User','role_id','required, exists:roles,id'],
    ['Tambah User','division_id','required, exists:divisions,id'],
    ['Generate Report','title','required, string, max:255'],
    ['Generate Report','type','required, string'],
    ['Generate Report','report_date','required, date'],
    ['Generate Report','attachment','nullable, file, max:2048'],
])

# ===== 4.8 =====
h2('4.8 Pembahasan')
ap('Pada subbab ini diuraikan pembahasan mengenai hasil implementasi Sistem ERP Prototype yang telah dibangun selama pelaksanaan Kerja Praktik. Pembahasan mencakup analisis kelebihan dan kekurangan sistem berdasarkan hasil implementasi dan pengujian yang telah dilakukan.')

h3('4.8.1 Kelebihan Sistem')
ap('Berdasarkan hasil implementasi dan pengujian, Sistem ERP Prototype PT. Golden Intan Berlian memiliki beberapa kelebihan sebagai berikut:')
li('**Arsitektur Terstruktur** \u2014 Penggunaan pola MVC pada framework Laravel memastikan pemisahan yang jelas antara logika bisnis (Controller), tampilan antarmuka (View), dan pengelolaan data (Model), sehingga kode program lebih terorganisir, mudah dipelihara, dan dapat dikembangkan secara modular.', 1)
li('**Keamanan Data** \u2014 Sistem menerapkan beberapa mekanisme keamanan yang meliputi: hashing password menggunakan algoritma bcrypt, proteksi CSRF (Cross-Site Request Forgery) pada setiap form, regenerasi session setelah proses login, dan validasi input pada sisi server untuk mencegah injeksi data berbahaya.', 2)
li('**Antarmuka Responsif dan Modern** \u2014 Antarmuka pengguna dibangun menggunakan TailwindCSS 4.0 yang memastikan tampilan responsif di berbagai ukuran layar. Dukungan fitur dark mode meningkatkan kenyamanan dan fleksibilitas penggunaan sistem.', 3)
li('**Pencatatan Aktivitas Otomatis** \u2014 Setiap aksi penting yang dilakukan dalam sistem dicatat secara otomatis ke dalam tabel activities, sehingga tersedia audit trail yang komprehensif untuk keperluan monitoring dan evaluasi.', 4)
li('**Visualisasi Data Informatif** \u2014 Penggunaan Chart.js untuk menampilkan grafik tren aktivitas pada dashboard membantu pengambilan keputusan berbasis data secara visual dan intuitif.', 5)
li('**Pengembangan Cepat dengan RAD** \u2014 Penerapan metode RAD memungkinkan pengembangan sistem secara iteratif dan cepat, sesuai dengan keterbatasan waktu pelaksanaan Kerja Praktik.', 6)

h3('4.8.2 Kekurangan Sistem')
ap('Meskipun demikian, sebagai sebuah prototype, sistem ini juga memiliki beberapa kekurangan dan keterbatasan yang perlu diperhatikan secara realistis:')
li('**Status Prototype** \u2014 Sistem masih berupa prototype sehingga beberapa fitur, khususnya fitur untuk aktor Karyawan seperti pencatatan aktivitas harian dan dashboard karyawan, belum diimplementasikan secara lengkap dan memerlukan pengembangan lebih lanjut.', 1)
li('**Belum Terdapat Middleware Role** \u2014 Sistem belum memiliki middleware khusus untuk memvalidasi role pengguna pada setiap route admin. Saat ini pembatasan akses hanya berdasarkan redirect di route /dashboard, sehingga secara teoritis pengguna dengan role karyawan yang mengetahui URL admin dapat mengaksesnya.', 2)
li('**Penyimpanan File Lokal** \u2014 Fitur upload lampiran pada modul report masih menggunakan penyimpanan lokal (local storage) dan belum terintegrasi dengan layanan cloud storage seperti Amazon S3 atau Google Cloud Storage untuk skalabilitas yang lebih baik.', 3)
li('**Belum Terdapat Sistem Notifikasi** \u2014 Sistem belum memiliki fitur notifikasi real-time untuk memberitahu admin tentang aktivitas penting atau perubahan data dalam sistem.', 4)
li('**Ekspor Data Terbatas** \u2014 Fitur download report belum sepenuhnya diimplementasikan untuk menghasilkan file dalam format PDF atau Excel, yang merupakan kebutuhan umum dalam konteks pelaporan bisnis.', 5)
li('**Belum Dilakukan Performance Testing** \u2014 Pengujian yang dilakukan masih sebatas black box testing fungsional. Belum dilakukan pengujian performa (performance testing) untuk mengukur kemampuan sistem dalam menangani beban pengguna dalam jumlah besar.', 6)

h3('4.8.3 Kesesuaian dengan Metode RAD')
ap_rich([
    ('Penerapan metode Rapid Application Development (RAD) dalam pengembangan Sistem ERP Prototype ini telah berjalan sesuai dengan tahapan yang didefinisikan. Fase ', 'normal'),
    ('Requirements Planning', 'italic'),
    (' dilaksanakan melalui analisis kebutuhan dan observasi proses bisnis perusahaan. Fase ', 'normal'),
    ('User Design', 'italic'),
    (' direalisasikan melalui perancangan diagram UML dan ERD. Fase ', 'normal'),
    ('Construction', 'italic'),
    (' dilaksanakan melalui implementasi kode program menggunakan Laravel. Fase ', 'normal'),
    ('Cutover', 'italic'),
    (' dilaksanakan melalui pengujian black box testing untuk memastikan fungsionalitas sistem.', 'normal'),
])
ap('Metode RAD terbukti efektif untuk pengembangan prototype ini karena memungkinkan iterasi yang cepat dalam proses pembangunan sistem. Namun demikian, keterbatasan waktu pelaksanaan Kerja Praktik menyebabkan beberapa fitur belum dapat diselesaikan secara sempurna dan memerlukan iterasi pengembangan tambahan di masa mendatang.')

# Save final
base = os.path.dirname(__file__)
output = os.path.join(base, 'laporan_bab4_v2.docx')
doc.save(output)
# Cleanup temp
try: os.remove(os.path.join(base, '_bab4_temp.docx'))
except: pass
print(f'BAB IV DOCX saved: {output}')
