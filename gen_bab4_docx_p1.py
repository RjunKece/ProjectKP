
# BAB IV DOCX Generator - Part 1 (4.1 - 4.4)
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import pickle, os

doc = Document()
for section in doc.sections:
    section.top_margin = Cm(3)
    section.bottom_margin = Cm(3)
    section.left_margin = Cm(4)
    section.right_margin = Cm(3)

style_normal = doc.styles['Normal']
style_normal.font.name = 'Times New Roman'
style_normal.font.size = Pt(12)
style_normal.font.color.rgb = RGBColor(0, 0, 0)
style_normal.paragraph_format.line_spacing = 1.5
style_normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
style_normal.element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

def hc(text, size=14):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text); r.bold = True; r.font.name = 'Times New Roman'; r.font.size = Pt(size); r.font.color.rgb = RGBColor(0,0,0)
    p.paragraph_format.line_spacing = 1.5
    return p

def h2(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(text); r.bold = True; r.font.name = 'Times New Roman'; r.font.size = Pt(12); r.font.color.rgb = RGBColor(0,0,0)
    p.paragraph_format.line_spacing = 1.5; p.paragraph_format.space_before = Pt(18)
    return p

def h3(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(text); r.bold = True; r.font.name = 'Times New Roman'; r.font.size = Pt(12); r.font.color.rgb = RGBColor(0,0,0)
    p.paragraph_format.line_spacing = 1.5; p.paragraph_format.space_before = Pt(12)
    return p

def ap(text, indent=True):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY; p.paragraph_format.line_spacing = 1.5
    if indent: p.paragraph_format.first_line_indent = Cm(1.27)
    r = p.add_run(text); r.font.name = 'Times New Roman'; r.font.size = Pt(12); r.font.color.rgb = RGBColor(0,0,0)
    return p

def ap_rich(parts, indent=True):
    """Add paragraph with mixed bold/normal/italic parts.
    parts = [('text', 'normal'), ('bold text', 'bold'), ('italic', 'italic')]"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY; p.paragraph_format.line_spacing = 1.5
    if indent: p.paragraph_format.first_line_indent = Cm(1.27)
    for text, style in parts:
        r = p.add_run(text); r.font.name = 'Times New Roman'; r.font.size = Pt(12); r.font.color.rgb = RGBColor(0,0,0)
        if style == 'bold': r.bold = True
        elif style == 'italic': r.italic = True
    return p

def li(text, num=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY; p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.left_indent = Cm(1.27); p.paragraph_format.first_line_indent = Cm(0)
    prefix = f"{num}. " if num else "\u2022 "
    parts = (prefix + text).split('**')
    for i, part in enumerate(parts):
        r = p.add_run(part); r.font.name = 'Times New Roman'; r.font.size = Pt(12); r.font.color.rgb = RGBColor(0,0,0)
        if i % 2 == 1: r.bold = True
    return p

def tbl(headers, rows):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = 'Table Grid'; table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = table.rows[0].cells[i]; c.text = ''; p = c.paragraphs[0]
        r = p.add_run(h); r.bold = True; r.font.name = 'Times New Roman'; r.font.size = Pt(11); r.font.color.rgb = RGBColor(0,0,0)
        p.paragraph_format.line_spacing = 1.5
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            c = table.rows[ri+1].cells[ci]; c.text = ''; p = c.paragraphs[0]
            r = p.add_run(val); r.font.name = 'Times New Roman'; r.font.size = Pt(11); r.font.color.rgb = RGBColor(0,0,0)
            p.paragraph_format.line_spacing = 1.5
    doc.add_paragraph()
    return table

def img_placeholder(num, caption):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(f'[Gambar 4.{num} {caption}]')
    r.font.name = 'Times New Roman'; r.font.size = Pt(12); r.font.color.rgb = RGBColor(128,128,128)
    r.italic = True
    # Caption
    pc = doc.add_paragraph()
    pc.alignment = WD_ALIGN_PARAGRAPH.CENTER; pc.paragraph_format.line_spacing = 1.5
    rc = pc.add_run(f'Gambar 4.{num} '); rc.bold = True; rc.font.name = 'Times New Roman'; rc.font.size = Pt(12); rc.font.color.rgb = RGBColor(0,0,0)
    rc2 = pc.add_run(caption); rc2.font.name = 'Times New Roman'; rc2.font.size = Pt(12); rc2.font.color.rgb = RGBColor(0,0,0)

# ===== TITLE =====
hc('BAB IV')
hc('HASIL DAN PEMBAHASAN')

# ===== 4.1 =====
h2('4.1 Gambaran Umum Sistem')
ap('Pada Bab IV ini diuraikan hasil perancangan dan implementasi Sistem Enterprise Resource Planning (ERP) sederhana berbasis web yang dikembangkan selama pelaksanaan Kerja Praktik di PT. Golden Intan Berlian. Sistem ini dirancang sebagai prototype pendukung proses bisnis perusahaan dengan tujuan mengintegrasikan fungsi-fungsi operasional ke dalam satu platform terpadu yang dapat diakses melalui browser.')
ap('Pengembangan sistem dilakukan dengan menggunakan metode Rapid Application Development (RAD), yang memungkinkan proses pembangunan aplikasi secara cepat dan iteratif. Metode RAD dipilih karena sesuai dengan keterbatasan waktu pelaksanaan Kerja Praktik, serta memungkinkan penyesuaian kebutuhan secara fleksibel berdasarkan umpan balik yang diperoleh selama proses pengembangan.')
ap('Sistem ERP Prototype ini dibangun menggunakan framework Laravel 12 dengan bahasa pemrograman PHP 8.2 dan berjalan di atas arsitektur Model-View-Controller (MVC). Basis data yang digunakan adalah MySQL, sementara antarmuka pengguna dikembangkan menggunakan Blade Template Engine dengan dukungan TailwindCSS 4.0 untuk styling responsif, Alpine.js untuk interaktivitas sisi klien, dan Chart.js untuk visualisasi data dalam bentuk grafik. Vite 7 digunakan sebagai build tool untuk pengelolaan aset frontend. Sistem dijalankan pada server lokal menggunakan Laragon sebagai lingkungan pengembangan.')
ap('Sistem ini dirancang untuk dua aktor utama, yaitu Super Admin dan Karyawan. Super Admin memiliki hak akses penuh terhadap seluruh fitur sistem, termasuk manajemen pengguna, monitoring aktivitas, dan pembuatan laporan. Sementara itu, aktor Karyawan dirancang untuk memiliki akses terbatas pada dashboard pribadi dan pencatatan aktivitas harian. Meskipun implementasi fitur untuk aktor Karyawan pada tahap prototype ini belum sepenuhnya lengkap, perancangan sistem telah mempertimbangkan peran tersebut untuk pengembangan di masa mendatang.')

# ===== 4.2 =====
h2('4.2 Analisis Sistem')

h3('4.2.1 Analisis Sistem Berjalan')
ap('Berdasarkan hasil observasi selama pelaksanaan Kerja Praktik, proses bisnis yang berjalan di PT. Golden Intan Berlian masih menggunakan metode konvensional dalam beberapa aspek operasionalnya. Pencatatan aktivitas karyawan, pengelolaan data pengguna, dan pembuatan laporan dilakukan secara manual menggunakan spreadsheet atau dokumen tercetak. Hal ini menimbulkan beberapa permasalahan, antara lain:')
li('**Inefisiensi waktu** \u2014 Proses pencarian data dan penyusunan laporan memerlukan waktu yang cukup lama karena harus dilakukan secara manual.', 1)
li('**Risiko kesalahan data** \u2014 Input data secara manual rentan terhadap human error seperti duplikasi data, kesalahan pengetikan, atau inkonsistensi format.', 2)
li('**Keterbatasan akses informasi** \u2014 Data yang tersebar di berbagai dokumen menyulitkan proses monitoring dan pengambilan keputusan secara real-time.', 3)
li('**Tidak adanya audit trail** \u2014 Tidak tersedia mekanisme pencatatan otomatis atas setiap perubahan dan aktivitas yang terjadi dalam proses bisnis.', 4)

h3('4.2.2 Analisis Sistem Usulan')
ap('Berdasarkan permasalahan pada sistem berjalan, diusulkan pengembangan Sistem ERP Prototype berbasis web yang mampu mengintegrasikan proses pengelolaan data pengguna, pencatatan aktivitas, dan pembuatan laporan ke dalam satu platform terpusat. Sistem usulan ini memiliki karakteristik sebagai berikut:')
li('**Sentralisasi data** \u2014 Seluruh data tersimpan dalam satu basis data MySQL yang terstruktur dan dapat diakses secara terpusat.', 1)
li('**Otomatisasi pencatatan aktivitas** \u2014 Setiap aksi yang dilakukan pengguna dalam sistem dicatat secara otomatis melalui service ActivityLogger.', 2)
li('**Manajemen pengguna berbasis role** \u2014 Sistem mendukung pembagian hak akses berdasarkan role (Super Admin dan Karyawan) sehingga keamanan data terjaga.', 3)
li('**Pembuatan laporan digital** \u2014 Laporan dapat dibuat, disimpan, dan diakses kembali secara digital melalui modul Reports Center.', 4)
li('**Antarmuka responsif dan modern** \u2014 Tampilan sistem dirancang responsif menggunakan TailwindCSS sehingga dapat diakses dari berbagai perangkat.', 5)

# ===== 4.3 =====
h2('4.3 Perancangan Sistem')
ap_rich([
    ('Tahap perancangan sistem merupakan fase kedua dalam metode RAD setelah fase perencanaan kebutuhan (', 'normal'),
    ('Requirements Planning', 'italic'),
    ('). Pada tahap ini dilakukan pemodelan proses bisnis dan struktur data menggunakan diagram UML (Unified Modeling Language) serta Entity Relationship Diagram (ERD). Perancangan ini menjadi acuan dalam proses implementasi sistem.', 'normal'),
])

h3('4.3.1 Use Case Diagram')
ap('Use Case Diagram digunakan untuk memodelkan interaksi antara aktor dengan sistem. Dalam Sistem ERP Prototype ini terdapat dua aktor utama, yaitu Super Admin dan Karyawan. Setiap aktor memiliki hak akses dan fungsionalitas yang berbeda sesuai dengan perannya masing-masing dalam sistem.')
ap('Super Admin sebagai aktor dengan hak akses tertinggi dapat melakukan seluruh fungsionalitas sistem yang meliputi: login ke sistem, mengakses dashboard admin yang menampilkan Key Performance Indicator (KPI), mengelola data pengguna (tambah user dan reset password), melakukan monitoring aktivitas seluruh pengguna, membuat dan mengelola laporan (Reports), serta mengelola profil akun. Sementara itu, aktor Karyawan memiliki akses terbatas yang meliputi: login ke sistem, mengakses dashboard karyawan, melihat profil pribadi, dan melakukan logout. Meskipun pada tahap prototype ini fitur Karyawan belum diimplementasikan secara lengkap, use case untuk aktor tersebut tetap dimodelkan sebagai bagian dari perancangan sistem secara keseluruhan.')
img_placeholder(1, 'Use Case Diagram Sistem ERP')

h3('4.3.2 Activity Diagram')
ap_rich([
    ('Activity Diagram digunakan untuk memodelkan alur kerja (', 'normal'),
    ('workflow', 'italic'),
    (') dari proses-proses utama dalam sistem. Diagram aktivitas ini menggambarkan urutan langkah-langkah yang dilakukan oleh aktor dalam menjalankan fungsionalitas tertentu, termasuk percabangan kondisi dan alur alternatif.', 'normal'),
])
ap('Alur utama yang dimodelkan dalam Activity Diagram meliputi: (1) Proses Login, yang dimulai dari pengisian kredensial oleh pengguna, validasi oleh sistem, pengecekan role, hingga redirect ke dashboard yang sesuai; (2) Proses Manajemen User, yang mencakup alur penambahan user baru dengan validasi input dan pencatatan aktivitas; (3) Proses Monitoring Aktivitas, yang menggambarkan alur akses dan filter data aktivitas; serta (4) Proses Pembuatan Laporan, yang meliputi pemilihan jenis laporan, pengisian detail, preview, konfirmasi, dan penyimpanan.')
img_placeholder(2, 'Activity Diagram Sistem ERP')

h3('4.3.3 Entity Relationship Diagram (ERD)')
ap('Entity Relationship Diagram (ERD) digunakan untuk memodelkan struktur basis data serta hubungan antar entitas dalam sistem. ERD Sistem ERP Prototype ini menggambarkan lima entitas utama beserta atribut dan relasinya. Entitas-entitas tersebut adalah: users (menyimpan data pengguna), roles (menyimpan data hak akses), divisions (menyimpan data divisi perusahaan), activities (menyimpan log aktivitas), dan reports (menyimpan data laporan).')
ap('Relasi yang terdapat dalam ERD meliputi: relasi one-to-many antara tabel roles dan users (satu role dimiliki oleh banyak user), relasi one-to-many antara tabel divisions dan users (satu divisi memiliki banyak user), relasi one-to-many antara tabel users dan activities (satu user memiliki banyak aktivitas), serta relasi one-to-many antara tabel users dan reports (satu user dapat membuat banyak laporan). Seluruh relasi diimplementasikan menggunakan foreign key constraint pada MySQL.')
img_placeholder(3, 'Entity Relationship Diagram (ERD) Sistem ERP')

# ===== 4.4 =====
h2('4.4 Implementasi Sistem')
ap_rich([
    ('Tahap implementasi merupakan fase ketiga dalam metode RAD yang disebut ', 'normal'),
    ('Construction Phase', 'italic'),
    ('. Pada tahap ini, hasil perancangan diterjemahkan ke dalam kode program menggunakan teknologi yang telah ditentukan. Berikut adalah uraian implementasi setiap halaman dan fitur utama dalam Sistem ERP Prototype.', 'normal'),
])

h3('4.4.1 Halaman Login')
ap_rich([
    ('Halaman login merupakan titik masuk utama (', 'normal'),
    ('entry point', 'italic'),
    (') bagi seluruh pengguna untuk mengakses Sistem ERP. Halaman ini menampilkan formulir autentikasi yang terdiri dari dua field input, yaitu email dan password. Desain halaman login menggunakan layout dua kolom (', 'normal'),
    ('split screen', 'italic'),
    (') dengan bagian kiri menampilkan formulir login dan bagian kanan menampilkan branding perusahaan PT. Golden Intan Berlian dengan skema warna emas (gold) sebagai identitas visual.', 'normal'),
])
ap('Fungsi utama halaman login adalah melakukan proses autentikasi pengguna menggunakan facade Auth::attempt() pada Laravel. Setelah kredensial divalidasi, sistem melakukan regenerasi session untuk mencegah serangan session fixation, kemudian mengarahkan pengguna ke halaman dashboard sesuai dengan role yang dimiliki. Apabila Super Admin yang login, sistem akan melakukan redirect ke /admin/dashboard, sedangkan Karyawan akan diarahkan ke /karyawan/dashboard. Jika kredensial tidak valid, sistem menampilkan pesan error "Email atau password salah."')
img_placeholder(4, 'Halaman Login Sistem ERP')

h3('4.4.2 Dashboard Admin')
ap_rich([
    ('Dashboard Admin merupakan halaman utama yang ditampilkan setelah Super Admin berhasil melakukan login. Halaman ini berfungsi sebagai pusat informasi (', 'normal'),
    ('control center', 'italic'),
    (') yang menyajikan ringkasan data operasional sistem secara real-time. Dashboard dirancang untuk memberikan gambaran menyeluruh tentang kondisi sistem melalui visualisasi data yang informatif dan mudah dipahami.', 'normal'),
])
ap('Dashboard Admin terdiri dari tiga komponen utama: (1) System Snapshot yang menampilkan enam Key Performance Indicator (KPI) meliputi Total Users, Active Employees, Total Activities, Today Activity, Productivity, dan Growth; (2) Activity Analytics yang menampilkan grafik line chart menggunakan library Chart.js untuk memvisualisasikan tren aktivitas karyawan dan super admin selama 12 bulan terakhir; serta (3) Quick Actions yang menyediakan pintasan navigasi cepat ke fitur Manage Users, Review Activities, dan Generate Reports.')
img_placeholder(5, 'Halaman Dashboard Admin')

h3('4.4.3 Dashboard Karyawan')
ap('Dashboard Karyawan merupakan halaman yang dirancang untuk pengguna dengan role karyawan setelah berhasil login. Halaman ini bertujuan untuk menyediakan antarmuka yang memungkinkan karyawan melihat informasi terkait tugas, aktivitas harian, dan profil pribadi mereka. Pada tahap prototype saat ini, dashboard karyawan telah memiliki halaman dasar dengan tampilan informasi pengguna yang sedang login.')
ap('Secara konseptual, dashboard karyawan dirancang untuk menyediakan fitur-fitur seperti: ringkasan aktivitas pribadi, pencatatan aktivitas harian, notifikasi dari admin, dan akses ke profil. Meskipun implementasi fitur-fitur tersebut belum sepenuhnya lengkap pada tahap prototype ini, arsitektur sistem telah disiapkan untuk mengakomodasi pengembangan di masa mendatang melalui routing terpisah dengan prefix /karyawan dan controller yang independen.')
img_placeholder(6, 'Halaman Dashboard Karyawan')

h3('4.4.4 Manajemen User')
ap_rich([
    ('Halaman Manajemen User (', 'normal'),
    ('User Management', 'italic'),
    (') merupakan fitur yang memungkinkan Super Admin untuk mengelola seluruh akun pengguna yang terdaftar dalam sistem. Fitur ini merupakan salah satu komponen inti dari modul Human Resource Management dalam konsep ERP.', 'normal'),
])
ap('Halaman ini terdiri dari beberapa komponen: (1) Summary Cards yang menampilkan statistik Total Users, Active Users, Super Admin, dan Pending Approval; (2) Filter dan Pencarian berupa toolbar untuk mencari pengguna berdasarkan nama atau email; (3) Tabel Data Pengguna yang menampilkan daftar pengguna dengan informasi nama, email, role, divisi, status, dan tombol aksi; (4) Tambah User melalui modal form dengan field nama, email, role, dan divisi, dimana password default ditetapkan sebagai "password123" yang di-hash menggunakan bcrypt; serta (5) Reset Password untuk mereset password pengguna ke default melalui modal konfirmasi.')
img_placeholder(7, 'Halaman Manajemen User')

h3('4.4.5 Monitoring Aktivitas')
ap('Halaman Monitoring Aktivitas berfungsi sebagai pusat pemantauan seluruh log aktivitas yang tercatat dalam sistem. Fitur ini memungkinkan Super Admin untuk melacak dan mengawasi setiap aksi yang dilakukan oleh pengguna, sehingga mendukung fungsi audit trail dan pengawasan operasional.')
ap('Komponen utama halaman ini meliputi: (1) KPI Summary berupa empat kartu indikator yang menampilkan Total Aktivitas, Aktivitas Hari Ini, Aktivitas Admin, dan Aktivitas Karyawan; (2) Filter Bar berupa toolbar pencarian dengan filter berdasarkan role, divisi, dan tanggal; serta (3) Activity Table berupa tabel yang menampilkan informasi user, role, divisi, deskripsi aktivitas, dan waktu dengan paginasi 15 item per halaman menggunakan fitur pagination bawaan Laravel.')
img_placeholder(8, 'Halaman Monitoring Aktivitas')

h3('4.4.6 Halaman Reports')
ap('Halaman Reports Center merupakan pusat pembuatan dan pengelolaan laporan strategis dalam sistem ERP. Fitur ini memungkinkan Super Admin untuk menghasilkan berbagai jenis laporan yang mendukung proses pengambilan keputusan manajemen.')
ap('Halaman ini terdiri dari: (1) KPI Strip yang menampilkan lima indikator berupa Total Reports, Generated This Month, Scheduled, System Coverage, dan Storage Used; (2) Report Builder dengan empat modul pembuatan laporan yaitu Financial Intelligence, Department Performance, Growth & Trends, dan Custom Intelligence, dimana masing-masing modul membuka modal form untuk mengisi detail laporan; (3) Recent Reports berupa tabel laporan terbaru dengan informasi judul, tipe, periode, status, dan aksi (View/Download) dengan paginasi 10 item per halaman; serta (4) System Insight yang menampilkan statistik penggunaan laporan.')
img_placeholder(9, 'Halaman Reports Center')

h3('4.4.7 Halaman Profil Pengguna')
ap('Halaman Profil menampilkan informasi detail akun pengguna yang sedang login. Halaman ini dapat diakses oleh seluruh pengguna baik Super Admin maupun Karyawan. Fitur ini berfungsi untuk memberikan transparansi informasi akun kepada pengguna.')
ap('Komponen halaman profil meliputi: (1) Profile Overview berupa avatar inisial, nama, email, badge role, dan status aktif; (2) Account Information menampilkan detail role, divisi, dan tanggal registrasi; serta (3) Activity Summary yang menampilkan ringkasan total aktivitas dan aktivitas bulan berjalan.')
img_placeholder(10, 'Halaman Profil Pengguna')

# Save intermediate
doc.save(os.path.join(os.path.dirname(__file__), '_bab4_temp.docx'))
print("Part 1 done - saved _bab4_temp.docx")
