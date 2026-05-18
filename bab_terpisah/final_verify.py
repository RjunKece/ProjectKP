"""Final clean verification of all 5 BABs."""
import re
from docx import Document

def check_doc(filepath, label):
    doc = Document(filepath)
    print(f"\n{'='*60}")
    print(f"{label}")
    print(f"{'='*60}")
    
    # Count stats
    total_paras = len([p for p in doc.paragraphs if p.text.strip()])
    italic_paras = 0
    citation_list = []
    heading_transitions = []
    
    prev_heading = None
    prev_was_heading = False
    
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue
        
        style_name = para.style.name.lower() if para.style else ''
        is_h = 'heading' in style_name or 'head' in style_name
        
        if is_h:
            if prev_was_heading:
                heading_transitions.append(f"  {prev_heading} -> {text[:50]}")
            prev_heading = text[:50]
            prev_was_heading = True
        else:
            prev_was_heading = False
        
        # Count italic paragraphs
        has_italic = any(r.italic and r.text.strip() for r in para.runs)
        if has_italic and not is_h:
            italic_paras += 1
        
        # Collect citations
        cites = re.findall(r'\([^)]*\d{4}[^)]*\)', text)
        citation_list.extend(cites)
    
    print(f"  Total paragraphs: {total_paras}")
    print(f"  Paragraphs with italic foreign words: {italic_paras}")
    print(f"  Total citations: {len(citation_list)}")
    
    if citation_list:
        print(f"  Citation samples:")
        for c in citation_list[:8]:
            print(f"    - {c}")
    
    if heading_transitions:
        print(f"  [!] Heading-to-heading transitions still present:")
        for t in heading_transitions:
            print(t)
    else:
        print(f"  [OK] No direct heading-to-heading transitions")

files = [
    (r'd:\laragon\www\erp-prototype\bab_terpisah\BAB 1 - Pendahuluan.docx', 'BAB 1 - Pendahuluan'),
    (r'd:\laragon\www\erp-prototype\bab_terpisah\BAB 2 - Landasan Teori.docx', 'BAB 2 - Landasan Teori'),
    (r'd:\laragon\www\erp-prototype\bab_terpisah\BAB 3 - Metodologi Penelitian.docx', 'BAB 3 - Metodologi'),
    (r'd:\laragon\www\erp-prototype\bab_terpisah\BAB 4 - Hasil dan Pembahasan.docx', 'BAB 4 - Hasil'),
    (r'd:\laragon\www\erp-prototype\bab_terpisah\BAB 5 - Kesimpulan dan Saran.docx', 'BAB 5 - Kesimpulan'),
]

for f, l in files:
    check_doc(f, l)
