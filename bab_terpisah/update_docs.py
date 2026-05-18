"""
Update all BAB documents:
1. Fix citation formats in BAB 1 & BAB 2 (Landasan Teori section)
2. Italicize foreign/English words in all BABs
3. Add introductory paragraphs where heading jumps directly to sub-heading
"""
import re
import copy
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ============================================================
# PART 1: CITATION FIXER
# ============================================================

def fix_citations_in_paragraph(para):
    """Fix citation format in a paragraph's text runs.
    Rules:
    - 1 author: (Katu, 2020)
    - 2 authors: (Katu and Rizqy, 2020) or (Katu dan Rizqy, 2020)
    - 3+ authors: (Katu et al., 2020)
    """
    full_text = para.text
    if not full_text.strip():
        return False
    
    changed = False
    
    # We need to work at the run level to preserve formatting
    # Strategy: rebuild runs with fixed citation text
    
    # Collect all runs info
    runs_info = []
    for run in para.runs:
        runs_info.append({
            'text': run.text,
            'bold': run.bold,
            'italic': run.italic,
            'underline': run.underline,
            'font_name': run.font.name,
            'font_size': run.font.size,
            'font_color': run.font.color.rgb if run.font.color and run.font.color.rgb else None,
        })
    
    # Join all text
    joined = ''.join(r['text'] for r in runs_info)
    
    # Fix parenthetical citations like (Katu ,2020) -> (Katu, 2020)
    # Fix (Lutfi et al, 2022) -> (Lutfi et al., 2022)
    # Fix (A.S. Rosa dan Shalahuddin., 2018) -> (Rosa dan Shalahuddin, 2018)
    # Fix (Ayu Putri dan Rudi Ramadhan., 2021) -> (Putri dan Ramadhan, 2021)
    # Fix (Mohammad Al-Shboul et al., 2020) -> already correct
    # Fix (Grady Booch et al., 2021) -> (Booch et al., 2021)
    
    # Pattern for parenthetical citations
    def fix_paren_citation(m):
        content = m.group(1)
        # Handle multi-citation like (Martin., 1991; Kendall dan Kendall., 2019)
        if ';' in content:
            parts = content.split(';')
            fixed_parts = []
            for part in parts:
                fixed_parts.append(fix_single_citation(part.strip()))
            return '(' + '; '.join(fixed_parts) + ')'
        else:
            return '(' + fix_single_citation(content) + ')'
    
    def fix_single_citation(content):
        # Remove trailing dots before year in author names
        # e.g., "Martin., 1991" -> "Martin, 1991"
        # e.g., "Pressman dan Maxim., 2019" -> "Pressman dan Maxim, 2019"
        
        # Match: authors part, then year
        m = re.match(r'^(.+?),?\s*(\d{4})\s*$', content)
        if not m:
            return content
        
        authors = m.group(1).strip()
        year = m.group(2)
        
        # Clean up author trailing dots (but keep "et al.")
        # Remove dots at end of author names before comma
        authors = re.sub(r'\.\s*$', '', authors)
        
        # Fix "et al" without dot -> "et al."
        authors = re.sub(r'\bet al\b(?!\.)', 'et al.', authors)
        
        # Fix extra spaces around comma
        authors = re.sub(r'\s*,\s*', ', ', authors)
        authors = re.sub(r'\s+', ' ', authors)
        
        # Remove trailing comma or dot from authors (but keep "et al.")
        if authors.endswith('.') and not authors.endswith('et al.'):
            authors = authors[:-1].strip()
        if authors.endswith(','):
            authors = authors[:-1].strip()
            
        return f'{authors}, {year}'
    
    new_joined = re.sub(r'\(([^)]*\d{4}[^)]*)\)', fix_paren_citation, joined)
    
    # Fix narrative citations: "Author1 dan Author2 (Year)" format is fine
    # Fix "Abd Elmonem et al. (2017)" -> keep as is (narrative style)
    # But we need to check "Lutfi et al, 2022" in parens -> "Lutfi et al., 2022"
    
    if new_joined != joined:
        changed = True
        # Now redistribute text back to runs preserving formatting
        _redistribute_text(para, runs_info, new_joined)
    
    return changed


def _redistribute_text(para, original_runs_info, new_text):
    """Redistribute new_text back into runs, preserving original formatting."""
    if not original_runs_info:
        return
        
    # Remove all existing runs
    for run in para.runs:
        run._element.getparent().remove(run._element)
    
    # Simple approach: put all text in one run with first run's formatting
    # But this loses italic on specific words. Better approach:
    # Map character positions from old to new text, then rebuild.
    
    # Since citation fixes are small changes, use a simpler approach:
    # Create runs that match the original structure as closely as possible
    
    old_text = ''.join(r['text'] for r in original_runs_info)
    
    # If lengths are similar, try to map char-by-char
    # Use a diff-like approach: find common prefix/suffix per run
    
    # Simplest reliable approach: recreate with original formatting segments
    # Map each char in old text to its run index
    char_to_run = []
    for idx, ri in enumerate(original_runs_info):
        char_to_run.extend([idx] * len(ri['text']))
    
    # Use sequence matcher to align old and new
    from difflib import SequenceMatcher
    sm = SequenceMatcher(None, old_text, new_text)
    
    # Build mapping: for each char in new_text, what run format to use
    new_char_formats = [0] * len(new_text)  # default to first run format
    
    for op, i1, i2, j1, j2 in sm.get_opcodes():
        if op == 'equal':
            for k in range(j2 - j1):
                old_idx = i1 + k
                new_idx = j1 + k
                if old_idx < len(char_to_run):
                    new_char_formats[new_idx] = char_to_run[old_idx]
        elif op == 'replace':
            # Use the format of the first char being replaced
            fmt = char_to_run[i1] if i1 < len(char_to_run) else 0
            for k in range(j1, j2):
                new_char_formats[k] = fmt
        elif op == 'insert':
            # Use format of surrounding context
            fmt = char_to_run[i1] if i1 < len(char_to_run) else (char_to_run[i1-1] if i1 > 0 else 0)
            for k in range(j1, j2):
                new_char_formats[k] = fmt
    
    # Group consecutive chars with same format into runs
    if not new_text:
        return
        
    segments = []
    current_fmt = new_char_formats[0]
    current_text = new_text[0]
    
    for i in range(1, len(new_text)):
        if new_char_formats[i] == current_fmt:
            current_text += new_text[i]
        else:
            segments.append((current_text, current_fmt))
            current_fmt = new_char_formats[i]
            current_text = new_text[i]
    segments.append((current_text, current_fmt))
    
    # Create new runs
    for text, fmt_idx in segments:
        if fmt_idx >= len(original_runs_info):
            fmt_idx = len(original_runs_info) - 1
        fmt = original_runs_info[fmt_idx]
        run = para.add_run(text)
        run.bold = fmt['bold']
        run.italic = fmt['italic']
        run.underline = fmt['underline']
        if fmt['font_name']:
            run.font.name = fmt['font_name']
        if fmt['font_size']:
            run.font.size = fmt['font_size']
        if fmt['font_color']:
            run.font.color.rgb = fmt['font_color']


# ============================================================
# PART 2: ITALIC FOREIGN WORDS
# ============================================================

# Words/phrases that should be italicized (case-insensitive matching)
FOREIGN_WORDS = [
    # Tech terms
    'Enterprise Resource Planning', 'Material Requirements Planning',
    'Rapid Application Development', 'Model-View-Controller',
    'Role-Based Access Control', 'Unified Modeling Language',
    'Entity Relationship Diagram', 'Object-Relational Mapping',
    'Blade Template Engine', 'Artisan CLI',
    'Eloquent ORM', 'Command-line interface',
    'Chief Executive Officer', 'Human Resource Development',
    'Customer Relationship Management',
    'Equivalence Partitioning', 'Boundary Value Analysis',
    'Decision Table Testing', 'Black-Box Testing', 'black-box testing',
    'Content Creator', 'User Management',
    'Requirements Planning', 'User Design', 'Construction', 'Cutover',
    # Two-word terms
    'cloud computing', 'on-premise', 'open-source', 'utility-first',
    'tree-shaking', 'dark mode', 'light mode', 'design token',
    'digital marketing', 'version control', 'proof of concept',
    'stress testing', 'load testing', 'cloud storage',
    'bar chart', 'line chart', 'pie chart', 'doughnut chart',
    'radar chart', 'scatter plot', 'split-screen',
    'task management', 'local storage',
    # Single words
    'framework', 'software', 'hardware', 'database', 'network',
    'prototype', 'standalone', 'spreadsheet', 'real-time',
    'usable', 'reusable', 'timeline', 'waterfall',
    'scaffolding', 'middleware', 'routing', 'migration', 'seeder',
    'toggle', 'dropdown', 'modal', 'canvas',
    'dashboard', 'overview', 'engagement', 'marketplace', 'online',
    'advertiser', 'affiliate', 'branding', 'rendering',
    'responsive', 'collapse', 'scroll', 'session', 'hashing',
    'cache', 'browser', 'wireframe', 'glassmorphism', 'gradient',
    'micro-animations', 'hover', 'input', 'output', 'debugging',
    'login', 'logout', 'breakpoint', 'production-ready', 'dummy',
    'backup', 'localhost', 'deployment', 'Finance',
    'et al.', 'production-ready',
]

# Sort by length (longest first) to match longer phrases before shorter ones
FOREIGN_WORDS.sort(key=len, reverse=True)

def italicize_foreign_words_in_paragraph(para):
    """Find foreign words in paragraph text and make them italic."""
    full_text = para.text
    if not full_text.strip():
        return False
    
    # Skip headings - don't italicize words in headings
    style_name = para.style.name.lower() if para.style else ''
    if 'heading' in style_name or 'head' in style_name:
        return False
    # Skip captions
    if 'caption' in style_name:
        return False
    
    # Find all foreign word occurrences with positions
    occurrences = []  # (start, end, matched_word)
    text_lower = full_text.lower()
    
    for word in FOREIGN_WORDS:
        word_lower = word.lower()
        start = 0
        while True:
            idx = text_lower.find(word_lower, start)
            if idx == -1:
                break
            # Check word boundaries for single words (not phrases)
            if ' ' not in word and '-' not in word and '.' not in word:
                # Check boundaries
                before_ok = (idx == 0 or not full_text[idx-1].isalpha())
                after_ok = (idx + len(word) >= len(full_text) or not full_text[idx + len(word)].isalpha())
                if not (before_ok and after_ok):
                    start = idx + 1
                    continue
            
            # Check if already overlapping with a longer match
            overlaps = False
            for s, e, _ in occurrences:
                if idx < e and idx + len(word) > s:
                    overlaps = True
                    break
            
            if not overlaps:
                occurrences.append((idx, idx + len(word), word))
            start = idx + 1
    
    if not occurrences:
        return False
    
    occurrences.sort(key=lambda x: x[0])
    
    # Check which occurrences are already italic
    # Build a map of which characters are italic
    italic_map = [False] * len(full_text)
    pos = 0
    for run in para.runs:
        run_len = len(run.text)
        if run.italic:
            for j in range(pos, min(pos + run_len, len(full_text))):
                italic_map[j] = True
        pos += run_len
    
    # Filter to only non-italic occurrences
    to_italicize = []
    for start, end, word in occurrences:
        # Check if ALL characters in range are already italic
        all_italic = all(italic_map[j] for j in range(start, min(end, len(full_text))))
        if not all_italic:
            to_italicize.append((start, end, word))
    
    if not to_italicize:
        return False
    
    # Now rebuild the paragraph runs with italic applied
    # Collect original run info with positions
    runs_data = []
    pos = 0
    for run in para.runs:
        runs_data.append({
            'text': run.text,
            'start': pos,
            'end': pos + len(run.text),
            'bold': run.bold,
            'italic': run.italic,
            'underline': run.underline,
            'font_name': run.font.name,
            'font_size': run.font.size,
            'font_color': run.font.color.rgb if run.font.color and run.font.color.rgb else None,
        })
        pos += len(run.text)
    
    # Build character-level format map
    char_formats = []
    for rd in runs_data:
        for _ in range(len(rd['text'])):
            char_formats.append({
                'bold': rd['bold'],
                'italic': rd['italic'],
                'underline': rd['underline'],
                'font_name': rd['font_name'],
                'font_size': rd['font_size'],
                'font_color': rd['font_color'],
            })
    
    # Apply italic to foreign word positions
    for start, end, word in to_italicize:
        for j in range(start, min(end, len(char_formats))):
            char_formats[j]['italic'] = True
    
    # Group consecutive chars with same format
    if not char_formats:
        return False
    
    segments = []
    current_fmt = char_formats[0]
    current_text = full_text[0] if full_text else ''
    
    for i in range(1, len(full_text)):
        if i < len(char_formats) and _fmt_equal(char_formats[i], current_fmt):
            current_text += full_text[i]
        else:
            segments.append((current_text, current_fmt))
            current_fmt = char_formats[i] if i < len(char_formats) else char_formats[-1]
            current_text = full_text[i]
    if current_text:
        segments.append((current_text, current_fmt))
    
    # Remove old runs and create new ones
    for run in para.runs:
        run._element.getparent().remove(run._element)
    
    for text, fmt in segments:
        run = para.add_run(text)
        run.bold = fmt['bold']
        run.italic = fmt['italic']
        run.underline = fmt['underline']
        if fmt['font_name']:
            run.font.name = fmt['font_name']
        if fmt['font_size']:
            run.font.size = fmt['font_size']
        if fmt['font_color']:
            run.font.color.rgb = fmt['font_color']
    
    return True


def _fmt_equal(f1, f2):
    return (f1['bold'] == f2['bold'] and 
            f1['italic'] == f2['italic'] and
            f1['underline'] == f2['underline'] and
            f1['font_name'] == f2['font_name'] and
            f1['font_size'] == f2['font_size'] and
            f1['font_color'] == f2['font_color'])


# ============================================================
# PART 3: ADD INTRODUCTORY PARAGRAPHS
# ============================================================

INTRO_TEXTS = {
    # BAB 1
    ('PENDAHULUAN', 'Latar Belakang'): 
        'Bab ini menguraikan latar belakang permasalahan, tujuan, manfaat, serta ruang lingkup pelaksanaan kegiatan Kerja Praktik yang dilakukan di PT. Golden Intan Berlian.',
    
    # BAB 2
    ('LANDASAN TEORI', 'Gambaran Umum Perusahaan'):
        'Bab ini menyajikan gambaran umum perusahaan tempat pelaksanaan Kerja Praktik serta landasan teori yang digunakan sebagai dasar dalam perancangan dan pengembangan sistem.',
    
    ('Gambaran Umum Perusahaan', 'Profil Perusahaan'):
        'Berikut ini diuraikan profil perusahaan, struktur organisasi, serta kegiatan operasional yang dilaksanakan oleh PT. Golden Intan Berlian.',
    
    ('Landasan Teori', 'Sistem Informasi'):
        'Berikut ini diuraikan beberapa teori dan konsep yang menjadi acuan dalam perancangan dan pengembangan sistem ERP Prototype pada kegiatan Kerja Praktik ini.',
    
    # BAB 3  
    ('METODOLOGI PENELITIAN', 'Metode Pengembangan Sistem'):
        'Bab ini menguraikan metodologi yang digunakan dalam pengembangan sistem, meliputi metode pengembangan, teknik pengumpulan data, alat dan bahan penelitian, analisis dan perancangan sistem, serta metode pengujian.',
    
    ('Analisis Sistem', 'Sistem Berjalan'):
        'Berikut ini diuraikan analisis terhadap sistem yang sedang berjalan serta sistem yang diusulkan sebagai solusi atas permasalahan yang ditemukan.',
    
    # BAB 4
    ('HASIL DAN PEMBAHASAN', 'Hasil Implementasi Sistem'):
        'Bab ini menyajikan hasil implementasi sistem berdasarkan perancangan yang telah dilakukan pada bab sebelumnya, serta pembahasan mengenai kelebihan, keterbatasan, dan kesesuaian sistem dengan metode yang digunakan.',
    
    ('Pembahasan', 'Kelebihan Sistem'):
        'Berikut ini diuraikan pembahasan mengenai kelebihan, keterbatasan, serta kesesuaian implementasi sistem dengan metode Rapid Application Development (RAD) yang digunakan.',
    
    # BAB 5
    ('KESIMPULAN DAN SARAN', 'Kesimpulan'):
        'Bab ini menyajikan kesimpulan dari hasil perancangan, implementasi, dan pengujian sistem yang telah dilakukan, serta saran untuk pengembangan sistem di masa yang akan datang.',
}


def is_heading(para):
    style_name = para.style.name.lower() if para.style else ''
    return 'heading' in style_name or 'head' in style_name


def add_intro_paragraphs(doc, label):
    """Add introductory text between heading and immediate sub-heading."""
    paragraphs = list(doc.paragraphs)
    insertions = []  # (index, text, style_to_copy_from)
    
    for i in range(len(paragraphs) - 1):
        if not is_heading(paragraphs[i]):
            continue
        
        # Check if next non-empty paragraph is also a heading
        next_heading_idx = None
        for j in range(i + 1, len(paragraphs)):
            if paragraphs[j].text.strip():
                if is_heading(paragraphs[j]):
                    next_heading_idx = j
                break
        
        if next_heading_idx is None:
            continue
        
        current_text = paragraphs[i].text.strip()
        next_text = paragraphs[next_heading_idx].text.strip()
        
        # Check if we have an intro text for this transition
        intro = None
        for (h1, h2), text in INTRO_TEXTS.items():
            if h1.lower() in current_text.lower() and h2.lower() in next_text.lower():
                intro = text
                break
        
        if intro:
            # Insert after the current heading (at next_heading_idx position)
            # We need to find the right insertion point
            insertions.append((next_heading_idx, intro))
    
    # Insert in reverse order to maintain indices
    for idx, text in sorted(insertions, key=lambda x: x[0], reverse=True):
        # Find a Normal style paragraph to copy formatting from
        ref_style = None
        ref_run = None
        for p in paragraphs:
            if p.style and 'normal' in p.style.name.lower() and p.text.strip():
                ref_style = p.style
                ref_run = p.runs[0] if p.runs else None
                break
        
        # Insert new paragraph before the target index
        target_para = paragraphs[idx]
        new_para = OxmlElement('w:p')
        target_para._element.addprevious(new_para)
        
        # Create proper paragraph with style
        from docx.text.paragraph import Paragraph
        new_p = Paragraph(new_para, doc)
        if ref_style:
            new_p.style = ref_style
        
        run = new_p.add_run(text)
        if ref_run:
            if ref_run.font.name:
                run.font.name = ref_run.font.name
            if ref_run.font.size:
                run.font.size = ref_run.font.size
    
    return len(insertions)


# ============================================================
# MAIN PROCESSING
# ============================================================

def process_bab1():
    filepath = r'd:\laragon\www\erp-prototype\bab_terpisah\BAB 1 - Pendahuluan.docx'
    print(f"\n{'='*60}")
    print(f"Processing BAB 1 - Pendahuluan")
    print(f"{'='*60}")
    
    doc = Document(filepath)
    
    # Fix citations
    citation_count = 0
    for para in doc.paragraphs:
        if fix_citations_in_paragraph(para):
            citation_count += 1
    print(f"  Citations fixed in {citation_count} paragraphs")
    
    # Italicize foreign words
    italic_count = 0
    for para in doc.paragraphs:
        if italicize_foreign_words_in_paragraph(para):
            italic_count += 1
    print(f"  Foreign words italicized in {italic_count} paragraphs")
    
    # Add intro paragraphs
    intro_count = add_intro_paragraphs(doc, 'BAB 1')
    print(f"  Introductory paragraphs added: {intro_count}")
    
    doc.save(filepath)
    print(f"  Saved: {filepath}")


def process_bab2():
    filepath = r'd:\laragon\www\erp-prototype\bab_terpisah\BAB 2 - Landasan Teori.docx'
    print(f"\n{'='*60}")
    print(f"Processing BAB 2 - Landasan Teori")
    print(f"{'='*60}")
    
    doc = Document(filepath)
    
    # Fix citations (only in Landasan Teori section - para index >= 56)
    citation_count = 0
    in_landasan_teori = False
    for i, para in enumerate(doc.paragraphs):
        style_name = para.style.name.lower() if para.style else ''
        if 'heading' in style_name or 'head' in style_name:
            if 'landasan teori' in para.text.lower():
                in_landasan_teori = True
        
        if in_landasan_teori:
            if fix_citations_in_paragraph(para):
                citation_count += 1
    print(f"  Citations fixed in {citation_count} paragraphs (Landasan Teori section)")
    
    # Italicize foreign words (all sections)
    italic_count = 0
    for para in doc.paragraphs:
        if italicize_foreign_words_in_paragraph(para):
            italic_count += 1
    print(f"  Foreign words italicized in {italic_count} paragraphs")
    
    # Add intro paragraphs
    intro_count = add_intro_paragraphs(doc, 'BAB 2')
    print(f"  Introductory paragraphs added: {intro_count}")
    
    doc.save(filepath)
    print(f"  Saved: {filepath}")


def process_bab(bab_num, filename):
    filepath = rf'd:\laragon\www\erp-prototype\bab_terpisah\{filename}'
    print(f"\n{'='*60}")
    print(f"Processing BAB {bab_num} - {filename}")
    print(f"{'='*60}")
    
    doc = Document(filepath)
    
    # Italicize foreign words
    italic_count = 0
    for para in doc.paragraphs:
        if italicize_foreign_words_in_paragraph(para):
            italic_count += 1
    print(f"  Foreign words italicized in {italic_count} paragraphs")
    
    # Add intro paragraphs
    intro_count = add_intro_paragraphs(doc, f'BAB {bab_num}')
    print(f"  Introductory paragraphs added: {intro_count}")
    
    doc.save(filepath)
    print(f"  Saved: {filepath}")


if __name__ == '__main__':
    # Create backups first
    import shutil
    import os
    
    backup_dir = r'd:\laragon\www\erp-prototype\bab_terpisah\backup'
    os.makedirs(backup_dir, exist_ok=True)
    
    files = [
        'BAB 1 - Pendahuluan.docx',
        'BAB 2 - Landasan Teori.docx',
        'BAB 3 - Metodologi Penelitian.docx',
        'BAB 4 - Hasil dan Pembahasan.docx',
        'BAB 5 - Kesimpulan dan Saran.docx',
    ]
    
    for f in files:
        src = rf'd:\laragon\www\erp-prototype\bab_terpisah\{f}'
        dst = rf'{backup_dir}\{f}'
        if not os.path.exists(dst):
            shutil.copy2(src, dst)
            print(f"Backup: {f}")
    
    # Process each BAB
    process_bab1()
    process_bab2()
    process_bab(3, 'BAB 3 - Metodologi Penelitian.docx')
    process_bab(4, 'BAB 4 - Hasil dan Pembahasan.docx')
    process_bab(5, 'BAB 5 - Kesimpulan dan Saran.docx')
    
    print(f"\n{'='*60}")
    print("ALL DONE! Backups saved in bab_terpisah/backup/")
    print(f"{'='*60}")
