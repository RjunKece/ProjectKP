from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()
for section in doc.sections:
    section.top_margin = Cm(3)
    section.bottom_margin = Cm(3)
    section.left_margin = Cm(4)
    section.right_margin = Cm(3)

style_normal = doc.styles['Normal']
style_normal.font.name = 'Times New Roman'
style_normal.font.size = Pt(12)
style_normal.font.color.rgb = RGBColor(0, 0, 0)
style_normal.paragraph_format.line_spacing = 1.25
style_normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
style_normal.element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

def hc(text, size=14):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text); r.bold = True; r.font.name = 'Times New Roman'; r.font.size = Pt(size); r.font.color.rgb = RGBColor(0,0,0)
    p.paragraph_format.line_spacing = 1.25
    return p

def hl(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(text); r.bold = True; r.font.name = 'Times New Roman'; r.font.size = Pt(12); r.font.color.rgb = RGBColor(0,0,0)
    p.paragraph_format.line_spacing = 1.25; p.paragraph_format.space_before = Pt(12)
    return p

def ap(text, indent=True):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY; p.paragraph_format.line_spacing = 1.25
    if indent: p.paragraph_format.first_line_indent = Cm(1.27)
    parts = text.split('**')
    for i, part in enumerate(parts):
        r = p.add_run(part); r.font.name = 'Times New Roman'; r.font.size = Pt(12); r.font.color.rgb = RGBColor(0,0,0)
        if i % 2 == 1: r.bold = True
    return p

def li(text, num=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY; p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.left_indent = Cm(1.27); p.paragraph_format.first_line_indent = Cm(0)
    prefix = f"{num}. " if num else "\u2022 "
    parts = (prefix + text).split('**')
    for i, part in enumerate(parts):
        r = p.add_run(part); r.font.name = 'Times New Roman'; r.font.size = Pt(12); r.font.color.rgb = RGBColor(0,0,0)
        if i % 2 == 1: r.bold = True
    return p

def tbl(headers, rows):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = 'Table Grid'; table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = table.rows[0].cells[i]; c.text = ''; p = c.paragraphs[0]
        r = p.add_run(h); r.bold = True; r.font.name = 'Times New Roman'; r.font.size = Pt(11); r.font.color.rgb = RGBColor(0,0,0)
        p.paragraph_format.line_spacing = 1.25
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            c = table.rows[ri+1].cells[ci]; c.text = ''; p = c.paragraphs[0]
            r = p.add_run(val); r.font.name = 'Times New Roman'; r.font.size = Pt(11); r.font.color.rgb = RGBColor(0,0,0)
            p.paragraph_format.line_spacing = 1.25
    doc.add_paragraph()
    return table

# ===== BAB II =====
hc('BAB II')
hc('LANDASAN TEORI')

# ===================== 2.1 Sistem Informasi =====================
hl('2.1 Sistem Informasi')
ap('Sistem informasi merupakan kombinasi dari teknologi informasi dan aktivitas manusia yang menggunakan teknologi tersebut untuk mendukung operasi dan manajemen organisasi. Menurut Laudon dan Laudon (2020), sistem informasi dapat didefinisikan sebagai sekumpulan komponen yang saling berhubungan untuk mengumpulkan, memproses, menyimpan, dan mendistribusikan informasi guna mendukung pengambilan keputusan, koordinasi, dan pengendalian dalam suatu organisasi.')
ap('Sistem informasi terdiri dari beberapa komponen utama yang saling berinteraksi, yaitu: perangkat keras (hardware), perangkat lunak (software), basis data (database), jaringan komunikasi (network), prosedur (procedures), dan sumber daya manusia (people). Keenam komponen ini bekerja sama untuk mengubah data mentah menjadi informasi yang berguna bagi penggunanya (O\u2019Brien dan Marakas, 2011).')
ap('Dalam konteks organisasi modern, sistem informasi memiliki peran strategis yang tidak hanya mendukung operasional harian tetapi juga membantu manajemen dalam mengambil keputusan berbasis data. Sommerville (2016) menekankan bahwa pengembangan sistem informasi yang baik harus memperhatikan kebutuhan pengguna, keandalan sistem, serta kemampuan sistem untuk beradaptasi terhadap perubahan kebutuhan bisnis.')

# ===================== 2.2 Enterprise Resource Planning =====================
hl('2.2 Enterprise Resource Planning (ERP)')
ap('Enterprise Resource Planning (ERP) merupakan sistem perangkat lunak terintegrasi yang dirancang untuk mengelola dan mengotomatisasi berbagai proses bisnis inti dalam suatu organisasi. Menurut Monk dan Wagner (2013), ERP adalah program perangkat lunak inti yang digunakan oleh perusahaan untuk mengkoordinasikan informasi di setiap area bisnis. Sistem ERP membantu mengelola proses bisnis perusahaan secara keseluruhan dengan menggunakan basis data tunggal dan alat pelaporan manajemen bersama.')
ap('Romney dan Steinbart (2018) mendefinisikan ERP sebagai sistem yang mengintegrasikan seluruh aspek aktivitas organisasi, termasuk akuntansi, keuangan, pemasaran, sumber daya manusia, manufaktur, dan manajemen persediaan ke dalam satu sistem terpadu. Integrasi ini memungkinkan informasi yang dimasukkan pada satu bagian organisasi dapat langsung tersedia bagi bagian lain yang membutuhkan.')
ap('Davenport (1998) dalam artikelnya yang dipublikasikan di Harvard Business Review menyatakan bahwa sistem ERP menjanjikan integrasi seluruh aliran informasi dalam perusahaan, termasuk informasi keuangan, informasi sumber daya manusia, informasi rantai pasokan, dan informasi pelanggan. Integrasi ini dicapai melalui penggunaan basis data tunggal yang menyimpan seluruh data operasional perusahaan.')

hl('2.2.1 Manfaat Implementasi ERP')
ap('Implementasi sistem ERP memberikan berbagai keuntungan bagi organisasi. Leon (2014) mengidentifikasi beberapa manfaat utama dari implementasi ERP, antara lain:')
li('**Integrasi data keuangan** \u2014 Seluruh data keuangan tersimpan dalam satu sistem sehingga memudahkan konsolidasi dan pelaporan.', 1)
li('**Standarisasi proses operasional** \u2014 ERP mendorong standarisasi proses bisnis di seluruh departemen sehingga meningkatkan efisiensi.', 2)
li('**Standarisasi data SDM** \u2014 Data sumber daya manusia dikelola secara terpusat dan konsisten.', 3)
li('**Peningkatan pengambilan keputusan** \u2014 Ketersediaan data real-time mendukung pengambilan keputusan yang lebih cepat dan akurat.', 4)

ap('Umble, Haft, dan Umble (2003) dalam penelitiannya yang dipublikasikan di European Journal of Operational Research menambahkan bahwa keberhasilan implementasi ERP bergantung pada beberapa faktor kritis, di antaranya: pemahaman yang jelas tentang tujuan strategis, komitmen manajemen puncak, manajemen proyek yang baik, pelatihan pengguna yang memadai, serta pengukuran kinerja yang tepat.')

hl('2.2.2 ERP Berbasis Web')
ap('Perkembangan teknologi web telah mendorong evolusi sistem ERP dari aplikasi desktop tradisional menjadi aplikasi berbasis web. ERP berbasis web memungkinkan pengguna mengakses sistem melalui browser tanpa perlu menginstal aplikasi khusus di perangkat mereka. Laudon dan Laudon (2020) menyatakan bahwa tren ini didorong oleh kebutuhan akan fleksibilitas akses, pengurangan biaya infrastruktur, serta kemudahan dalam pemeliharaan dan pembaruan sistem.')
ap('Dalam konteks penelitian ini, sistem ERP Prototype PT. Golden Intan Berlian dibangun sebagai aplikasi berbasis web menggunakan framework Laravel 12 yang memungkinkan akses melalui browser. Pendekatan ini sesuai dengan kebutuhan perusahaan yang memiliki delapan divisi operasional yang memerlukan akses terpusat terhadap data dan informasi perusahaan.')

# ===================== 2.3 Metode RAD =====================
hl('2.3 Metode Rapid Application Development (RAD)')
ap('Rapid Application Development (RAD) merupakan metode pengembangan perangkat lunak yang menekankan pada siklus pengembangan yang singkat dengan kualitas yang tetap terjaga. Metode RAD pertama kali diperkenalkan oleh James Martin pada tahun 1991 sebagai alternatif dari metode pengembangan tradisional (waterfall) yang memiliki siklus pengembangan yang panjang dan kaku (Martin, 1991).')
ap('Menurut Kendall dan Kendall (2019), RAD merupakan pendekatan berorientasi objek untuk pengembangan sistem yang mencakup suatu metode pengembangan serta perangkat-perangkat lunak. Tujuan utama RAD adalah memperpendek waktu yang biasanya diperlukan dalam siklus pengembangan sistem tradisional antara perancangan dan penerapan suatu sistem informasi.')
ap('Pressman dan Maxim (2019) menegaskan bahwa RAD sangat cocok untuk proyek-proyek yang membutuhkan penyelesaian cepat namun tetap memperhatikan kualitas perangkat lunak. RAD memanfaatkan prototipe dan penggunaan kembali komponen perangkat lunak (reusable components) untuk mempercepat proses pengembangan.')

hl('2.3.1 Tahapan Metode RAD')
ap('Metode RAD terdiri dari empat tahapan utama yang dilaksanakan secara berurutan namun dengan durasi yang relatif singkat. Berikut adalah penjelasan masing-masing tahapan (Martin, 1991; Kendall dan Kendall, 2019):')

ap('**1. Requirements Planning (Perencanaan Kebutuhan)**')
ap('Tahap ini bertujuan untuk mengidentifikasi tujuan sistem, ruang lingkup proyek, dan kebutuhan informasi. Pengguna dan pengembang bekerja sama untuk menentukan kebutuhan fungsional dan non-fungsional sistem. Hasil dari tahap ini berupa daftar kebutuhan yang menjadi acuan pengembangan.', False)

ap('**2. User Design (Desain Pengguna)**')
ap('Tahap ini merupakan tahap perancangan yang dilakukan secara kolaboratif antara pengembang dan pengguna. Prototipe antarmuka dibuat dan dievaluasi secara iteratif hingga desain sesuai dengan kebutuhan. Keunggulan tahap ini adalah penggunaan prototipe visual yang dapat langsung dievaluasi oleh pengguna.', False)

ap('**3. Construction (Konstruksi)**')
ap('Tahap ini merupakan tahap pembangunan sistem secara intensif berdasarkan desain yang telah disetujui. Pengembang memanfaatkan fitur-fitur framework, komponen siap pakai, dan alat pengembangan modern untuk mempercepat proses pembangunan sistem.', False)

ap('**4. Cutover (Peralihan)**')
ap('Tahap akhir yang mencakup pengujian sistem secara menyeluruh, pelatihan pengguna, konversi data, dan peralihan dari sistem lama ke sistem baru. Pengujian dilakukan untuk memastikan seluruh fitur berjalan sesuai kebutuhan.', False)

hl('2.3.2 Kelebihan Metode RAD')
ap('Metode RAD memiliki beberapa kelebihan dibandingkan dengan metode pengembangan tradisional (Pressman dan Maxim, 2019):')
li('Waktu pengembangan yang lebih singkat karena menggunakan pendekatan iteratif dan komponen siap pakai.', 1)
li('Keterlibatan pengguna yang tinggi memastikan hasil akhir sesuai dengan kebutuhan.', 2)
li('Fleksibilitas terhadap perubahan kebutuhan selama proses pengembangan.', 3)
li('Penggunaan prototipe memungkinkan evaluasi dini terhadap desain sistem.', 4)
li('Cocok untuk proyek dengan skala kecil hingga menengah dengan batasan waktu yang ketat.', 5)

# ===================== 2.4 Arsitektur MVC =====================
hl('2.4 Arsitektur Model-View-Controller (MVC)')
ap('Model-View-Controller (MVC) merupakan pola arsitektur perangkat lunak yang memisahkan aplikasi menjadi tiga komponen utama yang saling berinteraksi. Menurut Sommerville (2016), pola MVC bertujuan untuk memisahkan logika bisnis dari tampilan antarmuka pengguna sehingga memudahkan pengembangan, pengujian, dan pemeliharaan aplikasi.')
ap('Tiga komponen utama dalam arsitektur MVC adalah sebagai berikut (Pressman dan Maxim, 2019):')
li('**Model** \u2014 Komponen yang bertanggung jawab terhadap logika bisnis dan pengelolaan data. Model berinteraksi langsung dengan basis data untuk melakukan operasi CRUD (Create, Read, Update, Delete).', 1)
li('**View** \u2014 Komponen yang bertanggung jawab terhadap tampilan antarmuka pengguna. View menerima data dari Model melalui Controller dan menampilkannya kepada pengguna.', 2)
li('**Controller** \u2014 Komponen yang berfungsi sebagai perantara antara Model dan View. Controller menerima input dari pengguna, memprosesnya melalui Model, dan menentukan View yang akan ditampilkan.', 3)
ap('Penerapan arsitektur MVC pada pengembangan sistem ERP Prototype ini menggunakan framework Laravel yang secara native mendukung pola MVC. Pemisahan komponen ini memastikan kode program lebih terstruktur, mudah dipelihara, dan memungkinkan pengembangan secara paralel oleh tim pengembang.')

# ===================== 2.5 Framework Laravel =====================
hl('2.5 Framework Laravel')
ap('Laravel merupakan framework aplikasi web berbasis PHP yang bersifat open-source dan mengikuti pola arsitektur Model-View-Controller (MVC). Laravel diciptakan oleh Taylor Otwell pada tahun 2011 dengan tujuan menyediakan alternatif yang lebih elegan dan ekspresif untuk pengembangan aplikasi web dibandingkan framework PHP lainnya (Otwell, 2024).')
ap('Laravel menyediakan berbagai fitur bawaan yang mempercepat proses pengembangan, antara lain:')
li('**Eloquent ORM** \u2014 Object-Relational Mapping yang memudahkan interaksi dengan basis data melalui pendekatan berorientasi objek. Setiap tabel pada basis data direpresentasikan oleh model Eloquent yang mendefinisikan relasi, atribut, dan operasi data.', 1)
li('**Blade Template Engine** \u2014 Mesin template yang menyediakan fitur template inheritance dan data binding untuk membangun antarmuka pengguna secara modular.', 2)
li('**Artisan CLI** \u2014 Command-line interface yang menyediakan berbagai perintah untuk mempercepat tugas-tugas pengembangan seperti pembuatan model, controller, migration, dan seeder.', 3)
li('**Migration dan Seeder** \u2014 Fitur untuk mengelola struktur basis data secara terversikan (version control) dan mengisi data awal untuk pengujian.', 4)
li('**Middleware** \u2014 Mekanisme untuk memfilter permintaan HTTP yang masuk ke aplikasi, digunakan untuk autentikasi, otorisasi, dan proteksi CSRF.', 5)
li('**Routing** \u2014 Sistem routing yang fleksibel untuk mendefinisikan URL dan menghubungkannya dengan controller yang sesuai.', 6)
ap('Dalam penelitian ini, sistem ERP Prototype dibangun menggunakan Laravel versi 12.0 dengan PHP 8.2 sebagai bahasa pemrograman server-side.')

# ===================== 2.6 MySQL =====================
hl('2.6 Basis Data MySQL')
ap('MySQL merupakan sistem manajemen basis data relasional (RDBMS) yang bersifat open-source dan merupakan salah satu RDBMS paling populer di dunia. MySQL menggunakan Structured Query Language (SQL) untuk mengelola dan memanipulasi data. Menurut Connolly dan Begg (2015), basis data relasional menyimpan data dalam bentuk tabel-tabel yang saling berhubungan melalui relasi kunci utama (primary key) dan kunci asing (foreign key).')
ap('Dalam sistem ERP Prototype ini, MySQL digunakan untuk menyimpan seluruh data operasional termasuk data pengguna, role, divisi, aktivitas, dan laporan. Fitur migration pada Laravel memungkinkan pengelolaan struktur basis data secara terversikan sehingga perubahan skema dapat dilacak dan dikelola dengan baik.')

# ===================== 2.7 TailwindCSS =====================
hl('2.7 TailwindCSS')
ap('TailwindCSS merupakan framework CSS berbasis pendekatan utility-first yang memungkinkan pengembang membangun antarmuka pengguna secara cepat dengan menulis class-class utilitas langsung pada elemen HTML. Berbeda dengan framework CSS tradisional yang menyediakan komponen siap pakai, TailwindCSS memberikan kontrol penuh terhadap desain melalui class-class atomik yang dapat dikombinasikan secara fleksibel (TailwindCSS Documentation, 2024).')
ap('Keunggulan TailwindCSS antara lain:')
li('Desain responsif bawaan dengan sistem breakpoint yang mudah digunakan.', 1)
li('Konsistensi desain melalui sistem design token (spacing, warna, typography).', 2)
li('Ukuran file CSS akhir yang kecil karena fitur tree-shaking yang menghapus class yang tidak digunakan.', 3)
li('Dukungan dark mode melalui konfigurasi class-based atau media query.', 4)
ap('Dalam penelitian ini, TailwindCSS versi 4.0 digunakan untuk membangun antarmuka pengguna sistem ERP yang responsif, modern, dan mendukung fitur dark mode.')

# ===================== 2.8 Alpine.js =====================
hl('2.8 Alpine.js')
ap('Alpine.js merupakan framework JavaScript minimalis yang dirancang untuk menambahkan interaktivitas pada halaman web tanpa memerlukan framework JavaScript yang besar seperti React atau Vue.js. Alpine.js menyediakan pendekatan deklaratif untuk menangani state management, event handling, dan manipulasi DOM langsung pada markup HTML (Alpine.js Documentation, 2024).')
ap('Dalam sistem ERP Prototype ini, Alpine.js digunakan untuk mengimplementasikan fitur-fitur interaktif seperti toggle dark mode, dropdown menu, modal dialog, dan notifikasi. Penggunaan Alpine.js sejalan dengan prinsip RAD karena memungkinkan penambahan interaktivitas secara cepat tanpa kompleksitas yang tinggi.')

# ===================== 2.9 Chart.js =====================
hl('2.9 Chart.js')
ap('Chart.js merupakan library JavaScript open-source untuk membuat visualisasi data berbasis grafik pada halaman web. Library ini mendukung berbagai jenis grafik seperti line chart, bar chart, pie chart, doughnut chart, radar chart, dan scatter plot. Chart.js menggunakan elemen HTML5 Canvas untuk merender grafik yang responsif dan interaktif (Chart.js Documentation, 2024).')
ap('Dalam sistem ERP Prototype ini, Chart.js digunakan pada dashboard admin untuk menampilkan grafik line chart yang memvisualisasikan tren aktivitas karyawan dan super admin selama 12 bulan terakhir. Visualisasi data ini membantu manajemen dalam memahami pola aktivitas dan mengambil keputusan berbasis data.')

# ===================== 2.10 Vite =====================
hl('2.10 Vite')
ap('Vite merupakan build tool generasi baru untuk pengembangan frontend yang diciptakan oleh Evan You (pencipta Vue.js). Vite menawarkan kecepatan pengembangan yang jauh lebih tinggi dibandingkan build tool tradisional seperti Webpack melalui fitur native ES modules dan Hot Module Replacement (HMR) yang sangat cepat (Vite Documentation, 2024).')
ap('Laravel 12 secara resmi menggunakan Vite sebagai build tool default menggantikan Laravel Mix (berbasis Webpack). Dalam penelitian ini, Vite versi 7.0 digunakan untuk melakukan asset bundling yang mencakup kompilasi CSS (TailwindCSS), pengelolaan JavaScript, dan optimasi aset untuk lingkungan produksi.')

# ===================== 2.11 Laravel Breeze =====================
hl('2.11 Laravel Breeze')
ap('Laravel Breeze merupakan paket scaffolding autentikasi resmi dari Laravel yang menyediakan implementasi fitur-fitur autentikasi dasar secara minimal dan sederhana. Fitur yang disediakan mencakup login, registrasi, reset password, verifikasi email, dan konfirmasi password (Laravel Documentation, 2024).')
ap('Dalam sistem ERP Prototype ini, Laravel Breeze versi 2.3 digunakan sebagai dasar sistem autentikasi. Proses autentikasi memanfaatkan facade Auth::attempt() untuk memvalidasi kredensial pengguna, session regeneration untuk keamanan, serta redirect berdasarkan role pengguna (Super Admin atau Karyawan).')

# ===================== 2.12 Black-Box Testing =====================
hl('2.12 Pengujian Black-Box Testing')
ap('Black-box testing merupakan metode pengujian perangkat lunak yang berfokus pada fungsionalitas sistem dari perspektif pengguna tanpa memperhatikan struktur internal kode program. Menurut Pressman dan Maxim (2019), black-box testing mengevaluasi apakah input yang diberikan menghasilkan output yang sesuai dengan spesifikasi kebutuhan yang telah ditetapkan.')
ap('Sommerville (2016) menjelaskan bahwa black-box testing memiliki beberapa teknik pengujian, antara lain:')
li('**Equivalence Partitioning** \u2014 Membagi domain input menjadi kelas-kelas ekuivalensi dan menguji satu nilai dari setiap kelas.', 1)
li('**Boundary Value Analysis** \u2014 Menguji nilai-nilai pada batas domain input.', 2)
li('**Decision Table Testing** \u2014 Menggunakan tabel keputusan untuk menguji kombinasi kondisi input.', 3)
ap('Dalam penelitian ini, metode black-box testing digunakan untuk menguji seluruh fungsionalitas sistem ERP Prototype meliputi proses login, manajemen pengguna, monitoring aktivitas, pembuatan laporan, dan fitur-fitur pendukung lainnya.')

# ===================== 2.13 Penelitian Terdahulu =====================
hl('2.13 Penelitian Terdahulu')
ap('Beberapa penelitian terdahulu yang relevan dengan penelitian ini dijadikan sebagai referensi dan bahan perbandingan:')

tbl(['No', 'Peneliti (Tahun)', 'Judul', 'Hasil Penelitian'], [
    ['1', 'Grabski, Leech, dan Schmidt (2011)',
     'A Review of ERP Research: A Future Agenda for Accounting Information Systems',
     'Implementasi ERP yang sukses meningkatkan efisiensi operasional melalui otomatisasi proses bisnis dan eliminasi redundansi data. Dipublikasikan di Journal of Information Systems, Vol. 25(1).'],
    ['2', 'Umble, Haft, dan Umble (2003)',
     'Enterprise Resource Planning: Implementation Procedures and Critical Success Factors',
     'Mengidentifikasi faktor kritis keberhasilan implementasi ERP meliputi komitmen manajemen, pelatihan pengguna, dan manajemen proyek. Dipublikasikan di European Journal of Operational Research, Vol. 146(2).'],
    ['3', 'Al-Mashari, Al-Mudimigh, dan Zairi (2003)',
     'Enterprise Resource Planning: A Taxonomy of Critical Factors',
     'Menyusun taksonomi faktor kritis implementasi ERP yang mencakup aspek strategis, taktis, dan operasional. Dipublikasikan di European Journal of Operational Research, Vol. 146(2).'],
    ['4', 'Ahmad, Mehmood, dan Yanez (2022)',
     'RAPD: Rapid and Participatory Application Development of Usable Systems During COVID19 Crisis',
     'Menunjukkan efektivitas metode RAD yang dikombinasikan dengan Participatory Design dalam mengembangkan sistem yang usable secara cepat. Dipublikasikan di IEEE Access, Vol. 10.'],
])

ap('Berdasarkan penelitian terdahulu di atas, dapat disimpulkan bahwa implementasi sistem ERP memberikan dampak positif terhadap efisiensi operasional perusahaan, dan metode RAD terbukti efektif dalam menghasilkan sistem yang fungsional dalam waktu singkat. Penelitian ini melengkapi penelitian terdahulu dengan mengimplementasikan sistem ERP Prototype menggunakan teknologi web modern (Laravel 12, TailwindCSS 4.0, Alpine.js) pada perusahaan skala menengah yang bergerak di bidang digital marketing.')

doc.save(r'd:\laragon\www\erp-prototype\laporan_bab2.docx')
print('File saved: laporan_bab2.docx')
