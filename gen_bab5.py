from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os

doc = Document()
for section in doc.sections:
    section.top_margin = Cm(3)
    section.bottom_margin = Cm(3)
    section.left_margin = Cm(4)
    section.right_margin = Cm(3)

style_normal = doc.styles['Normal']
style_normal.font.name = 'Times New Roman'
style_normal.font.size = Pt(12)
style_normal.font.color.rgb = RGBColor(0, 0, 0)
style_normal.paragraph_format.line_spacing = 1.5
style_normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
style_normal.element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

def hc(text, size=14):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text); r.bold = True; r.font.name = 'Times New Roman'; r.font.size = Pt(size); r.font.color.rgb = RGBColor(0,0,0)
    p.paragraph_format.line_spacing = 1.5
    return p

def hl(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(text); r.bold = True; r.font.name = 'Times New Roman'; r.font.size = Pt(12); r.font.color.rgb = RGBColor(0,0,0)
    p.paragraph_format.line_spacing = 1.5; p.paragraph_format.space_before = Pt(18)
    return p

def ap(text, indent=True, hang=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY; p.paragraph_format.line_spacing = 1.5
    if indent: 
        p.paragraph_format.first_line_indent = Cm(1.27)
    if hang:
        p.paragraph_format.left_indent = Cm(1.27)
        p.paragraph_format.first_line_indent = Cm(-1.27)
    r = p.add_run(text); r.font.name = 'Times New Roman'; r.font.size = Pt(12); r.font.color.rgb = RGBColor(0,0,0)
    return p

def li(text, num=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY; p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.left_indent = Cm(1.27); p.paragraph_format.first_line_indent = Cm(0)
    prefix = f"{num}. " if num else "\u2022 "
    parts = (prefix + text).split('**')
    for i, part in enumerate(parts):
        r = p.add_run(part); r.font.name = 'Times New Roman'; r.font.size = Pt(12); r.font.color.rgb = RGBColor(0,0,0)
        if i % 2 == 1: r.bold = True
    return p

# ===== BAB V =====
hc('BAB V')
hc('KESIMPULAN DAN SARAN')

hl('5.1 Kesimpulan')
ap('Berdasarkan hasil perancangan, implementasi, dan pengujian Sistem Enterprise Resource Planning (ERP) Prototype di PT. Golden Intan Berlian yang telah diuraikan pada bab-bab sebelumnya, dapat ditarik beberapa kesimpulan sebagai berikut:')
li('Sistem ERP Prototype telah berhasil dibangun berbasis web menggunakan framework Laravel 12 dengan arsitektur Model-View-Controller (MVC) dan basis data MySQL. Antarmuka pengguna (user interface) dirancang secara responsif menggunakan TailwindCSS 4.0 dan Alpine.js, memberikan pengalaman pengguna yang modern dan intuitif.', 1)
li('Penerapan metode Rapid Application Development (RAD) terbukti efektif dalam pengembangan sistem ini. Metode RAD memungkinkan proses analisis kebutuhan, perancangan antarmuka, dan konstruksi kode berjalan secara cepat dan iteratif, sehingga prototipe dapat segera dievaluasi dan disesuaikan dengan kebutuhan proses bisnis perusahaan meskipun dengan keterbatasan waktu Kerja Praktik.', 2)
li('Sistem telah mampu mengintegrasikan beberapa proses pengelolaan data pengguna, pemantauan aktivitas secara otomatis, serta pembuatan laporan (Reports) ke dalam satu platform terpusat. Hal ini memecahkan masalah inefisiensi waktu, risiko kesalahan data, dan ketiadaan audit trail pada sistem konvensional yang berjalan sebelumnya.', 3)
li('Hasil pengujian menggunakan metode Black Box Testing menunjukkan bahwa seluruh fungsionalitas utama sistem, termasuk sistem autentikasi, manajemen pengguna, pencatatan log aktivitas, pembuatan laporan, serta validasi input, telah berjalan dengan baik sesuai dengan skenario dan spesifikasi kebutuhan yang diharapkan.', 4)
li('Meskipun sistem telah beroperasi dengan baik, sistem ini masih berstatus prototype di mana beberapa fungsi operasional khusus, terutama pada dashboard dan fitur Karyawan, belum diimplementasikan secara menyeluruh namun fondasi arsitektur telah dipersiapkan untuk pengembangan lebih lanjut.', 5)

hl('5.2 Saran')
ap('Untuk pengembangan dan penyempurnaan sistem di masa yang akan datang, beberapa saran yang dapat dipertimbangkan adalah sebagai berikut:')
li('**Pengembangan Fitur Karyawan:** Mengimplementasikan fitur Karyawan secara menyeluruh, termasuk pencatatan aktivitas harian yang lebih spesifik, modul penugasan (task management), serta sistem notifikasi komunikasi antara Super Admin dan Karyawan.', 1)
li('**Penerapan Role Middleware:** Menambahkan proteksi tingkat lanjut menggunakan middleware berbasis role pada setiap rute (route) administratif, guna memastikan validasi hak akses secara ketat selain pembatasan melalui redirect pada dashboard.', 2)
li('**Integrasi Layanan Cloud:** Mengalihkan penyimpanan lampiran laporan (file upload) dari penyimpanan lokal (local storage) ke layanan cloud storage seperti Amazon S3 atau Google Cloud Storage untuk mendukung skalabilitas dan keamanan data jangka panjang.', 3)
li('**Fitur Ekspor Laporan:** Menyempurnakan modul Reports Center dengan menambahkan fitur ekspor data dalam berbagai format standar seperti PDF (Portable Document Format) dan Microsoft Excel untuk mempermudah pelaporan kepada manajemen tingkat atas.', 4)
li('**Pengujian Kinerja (Performance Testing):** Melakukan pengujian kinerja secara komprehensif, seperti stress testing atau load testing, untuk memastikan keandalan (reliability) sistem dalam menangani jumlah akses dan data aktivitas pengguna yang besar secara bersamaan.', 5)

doc.add_page_break()

# ===== DAFTAR PUSTAKA =====
hc('DAFTAR PUSTAKA')
ap('A.S., Rosa dan M. Shalahuddin. (2018). Rekayasa Perangkat Lunak Terstruktur dan Berorientasi Objek. Bandung: Informatika Bandung.', False, True)
ap('Dennis, A., Wixom, B. H., & Roth, R. M. (2012). Systems Analysis and Design (5th ed.). John Wiley & Sons.', False, True)
ap('Kendall, K. E., & Kendall, J. E. (2014). Systems Analysis and Design (9th ed.). Pearson.', False, True)
ap('Laravel. (2025). Laravel 12.x Documentation. Diakses dari https://laravel.com/docs/12.x.', False, True)
ap('Nidhra, S., & Dondeti, J. (2012). Black Box and White Box Testing Techniques - A Literature Review. International Journal of Embedded Systems and Applications (IJESA), 2(2).', False, True)
ap('Pressman, R. S., & Maxim, B. R. (2015). Software Engineering: A Practitioner\'s Approach (8th ed.). McGraw-Hill Education.', False, True)
ap('Sommerville, I. (2016). Software Engineering (10th ed.). Pearson.', False, True)
ap('Tailwind Labs. (2025). Tailwind CSS Documentation. Diakses dari https://tailwindcss.com/docs.', False, True)
ap('Wati, E. F., & Kusumo, A. A. (2016). Penerapan Metode Unified Modeling Language (UML) Berbasis Desktop Pada Sistem Informasi Point of Sale (POS). Jurnal Terapan Teknologi Informasi, 1(1).', False, True)

output = os.path.join(os.path.dirname(__file__), 'laporan_bab5_dan_daftar_pustaka.docx')
doc.save(output)
print(f'File saved: {output}')
