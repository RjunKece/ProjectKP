from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

doc = Document()
for section in doc.sections:
    section.top_margin = Cm(3)
    section.bottom_margin = Cm(3)
    section.left_margin = Cm(4)
    section.right_margin = Cm(3)

style_normal = doc.styles['Normal']
style_normal.font.name = 'Times New Roman'
style_normal.font.size = Pt(12)
style_normal.font.color.rgb = RGBColor(0, 0, 0)
style_normal.paragraph_format.line_spacing = 1.5
style_normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
style_normal.element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

def hc(text, size=14):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text); r.bold = True; r.font.name = 'Times New Roman'; r.font.size = Pt(size); r.font.color.rgb = RGBColor(0,0,0)
    p.paragraph_format.line_spacing = 1.5
    return p

def hl(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(text); r.bold = True; r.font.name = 'Times New Roman'; r.font.size = Pt(12); r.font.color.rgb = RGBColor(0,0,0)
    p.paragraph_format.line_spacing = 1.5; p.paragraph_format.space_before = Pt(18)
    return p

def ap(text, indent=True):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY; p.paragraph_format.line_spacing = 1.5
    if indent: p.paragraph_format.first_line_indent = Cm(1.27)
    parts = text.split('**')
    for i, part in enumerate(parts):
        r = p.add_run(part); r.font.name = 'Times New Roman'; r.font.size = Pt(12); r.font.color.rgb = RGBColor(0,0,0)
        if i % 2 == 1: r.bold = True
    return p

def li(text, num=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY; p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.left_indent = Cm(1.27); p.paragraph_format.first_line_indent = Cm(0)
    prefix = f"{num}. " if num else "\u2022 "
    parts = (prefix + text).split('**')
    for i, part in enumerate(parts):
        r = p.add_run(part); r.font.name = 'Times New Roman'; r.font.size = Pt(12); r.font.color.rgb = RGBColor(0,0,0)
        if i % 2 == 1: r.bold = True
    return p

# ===== LANDASAN TEORI =====
hl('Landasan Teori')

# --- Sistem Informasi ---
hl('Sistem Informasi')
ap('Sistem informasi merupakan kombinasi dari teknologi informasi dan aktivitas manusia yang menggunakan teknologi tersebut untuk mendukung operasi dan manajemen organisasi. Dennis, Wixom, dan Roth (2012) mendefinisikan sistem informasi sebagai sekumpulan komponen yang saling terintegrasi untuk mengumpulkan, memproses, menyimpan, dan menyebarkan informasi guna mendukung pengambilan keputusan, koordinasi, dan pengendalian dalam suatu organisasi.')
ap('Sistem informasi terdiri dari beberapa komponen utama yang saling berinteraksi, yaitu: perangkat keras (hardware), perangkat lunak (software), basis data (database), jaringan komunikasi (network), prosedur (procedures), dan sumber daya manusia (people). Keenam komponen ini bekerja sama untuk mengubah data mentah menjadi informasi yang berguna bagi penggunanya (A.S. Rosa dan Shalahuddin, 2018).')
ap('Berdasarkan kedua pandangan tersebut, dapat disimpulkan bahwa sistem informasi merupakan suatu kesatuan komponen teknologi dan sumber daya manusia yang dirancang untuk mengelola data menjadi informasi yang bernilai bagi organisasi, sehingga mampu mendukung efisiensi operasional dan pengambilan keputusan yang lebih baik.')

# --- ERP ---
hl('Enterprise Resource Planning (ERP)')
ap('Enterprise Resource Planning (ERP) merupakan sistem perangkat lunak terintegrasi yang dirancang untuk mengelola dan mengotomatisasi berbagai proses bisnis inti dalam suatu organisasi. P\u00e9rez Est\u00e9banez (2024) dalam artikelnya yang dipublikasikan pada jurnal Administrative Sciences (MDPI) menjelaskan bahwa sistem ERP memungkinkan perusahaan untuk mengintegrasikan seluruh fungsi bisnis ke dalam satu platform terpadu, sehingga meningkatkan efisiensi operasional dan pengambilan keputusan yang lebih akurat.')
ap('Alam, Sharma, dan Chowdhury (2024) dalam penelitiannya yang dipublikasikan pada Journal of Computer and Communications menegaskan bahwa sistem ERP modern, khususnya yang berbasis cloud, memungkinkan integrasi seluruh aspek aktivitas organisasi termasuk akuntansi, keuangan, pemasaran, sumber daya manusia, dan manajemen persediaan ke dalam satu sistem terpadu. Integrasi ini memungkinkan informasi yang dimasukkan pada satu bagian organisasi dapat langsung tersedia bagi bagian lain yang membutuhkan.')
ap('Berdasarkan kedua penelitian tersebut, dapat disimpulkan bahwa sistem ERP merupakan solusi teknologi terintegrasi yang menghubungkan berbagai fungsi bisnis dalam satu platform, sehingga mampu meningkatkan efisiensi, mengurangi redundansi data, dan mempercepat proses pengambilan keputusan manajerial.')

# --- ERP Sederhana sebagai Prototype ---
hl('ERP Sederhana sebagai Prototype')
ap('ERP sederhana merupakan bentuk pengembangan awal dari sistem ERP yang memiliki cakupan fungsi yang lebih terbatas dibandingkan dengan sistem ERP komersial. Sulaimon, Surin, dan Hamzah (2024) dalam artikelnya yang dipublikasikan pada Journal of International Business, Economics and Entrepreneurship menjelaskan bahwa pengembangan ERP dalam bentuk prototipe, khususnya yang berbasis open-source, memungkinkan perusahaan skala kecil dan menengah untuk memahami kebutuhan sistem secara lebih jelas sebelum dilakukan implementasi secara penuh.')
ap('Jaffa dan Salim (2020) dalam tinjauan literatur yang dipublikasikan pada International Journal of Integrated Engineering menambahkan bahwa pendekatan prototipe ERP memungkinkan perusahaan untuk melakukan evaluasi terhadap sistem yang dikembangkan sehingga risiko kegagalan implementasi dapat diminimalkan. Selain itu, prototipe ERP juga dapat digunakan sebagai dasar pengembangan sistem ERP yang lebih kompleks di masa mendatang sesuai dengan kebutuhan organisasi.')
ap('Berdasarkan kedua kajian tersebut, dapat disimpulkan bahwa pengembangan ERP dalam bentuk prototipe merupakan pendekatan strategis yang memungkinkan perusahaan memahami kebutuhan sistem dan meminimalkan risiko kegagalan sebelum melakukan implementasi sistem secara penuh.')

# --- Sistem ERP Berbasis Web ---
hl('Sistem ERP Berbasis Web')
ap('Perkembangan teknologi web memungkinkan sistem ERP dikembangkan dalam bentuk aplikasi berbasis web yang dapat diakses melalui browser. Salih et al. (2021) dalam artikel yang dipublikasikan pada jurnal Sensors (MDPI) menjelaskan bahwa sistem ERP berbasis cloud dan web memiliki berbagai keunggulan seperti kemudahan akses, fleksibilitas penggunaan, serta kemudahan dalam proses pengembangan dan pemeliharaan sistem.')
ap('Alam, Sharma, dan Chowdhury (2024) menambahkan bahwa penggunaan teknologi web dalam pengembangan sistem ERP memungkinkan pengguna untuk mengakses sistem dari berbagai perangkat yang terhubung dengan jaringan internet. Hal ini membuat sistem ERP menjadi lebih fleksibel serta mempermudah proses integrasi dengan sistem lain yang telah digunakan oleh perusahaan.')
ap('Dengan adanya sistem ERP berbasis web, perusahaan dapat mengelola berbagai aktivitas operasional secara lebih efektif karena sistem dapat diakses secara mudah oleh seluruh pengguna yang memiliki hak akses. Kedua penelitian di atas menegaskan bahwa ERP berbasis web merupakan solusi yang relevan dan terjangkau bagi perusahaan skala menengah.')

# --- Proses Bisnis Perusahaan ---
hl('Proses Bisnis Perusahaan')
ap('Proses bisnis merupakan rangkaian aktivitas yang saling berkaitan untuk menghasilkan nilai atau layanan bagi organisasi atau perusahaan. Dumas, La Rosa, Mendling, dan Reijers (2018) dalam bukunya Fundamentals of Business Process Management mendefinisikan proses bisnis sebagai kumpulan aktivitas terstruktur yang menggambarkan bagaimana suatu pekerjaan dilakukan mulai dari tahap awal hingga menghasilkan output yang diharapkan.')
ap('Dalam perusahaan yang bergerak di bidang digital marketing, proses bisnis meliputi berbagai aktivitas seperti pengelolaan data pengguna, pencatatan aktivitas kerja karyawan, serta monitoring kinerja tim dalam menjalankan kegiatan pemasaran digital.')
ap('Tanpa sistem yang terintegrasi, proses bisnis dalam perusahaan dapat berjalan secara tidak efisien karena data tersebar pada berbagai media pencatatan yang berbeda. Oleh karena itu, penerapan sistem ERP dapat membantu perusahaan dalam mengintegrasikan berbagai proses bisnis sehingga pengelolaan aktivitas operasional menjadi lebih efektif dan efisien.')

# --- Pengelolaan Pengguna dan Hak Akses ---
hl('Pengelolaan Pengguna dan Hak Akses')
ap('Pengelolaan pengguna dan hak akses merupakan bagian penting dalam sistem informasi yang digunakan oleh banyak pengguna dengan peran yang berbeda. Dennis, Wixom, dan Roth (2012) menjelaskan bahwa setiap pengguna dalam sistem memiliki peran tertentu yang menentukan hak akses terhadap fitur dan data yang tersedia dalam sistem.')
ap('Pengaturan hak akses bertujuan untuk menjaga keamanan data serta memastikan bahwa setiap pengguna hanya dapat mengakses informasi yang sesuai dengan tanggung jawabnya. Dengan adanya pengelolaan hak akses yang baik, sistem dapat berjalan secara lebih tertib dan terkontrol.')
ap('Dalam sistem ERP sederhana yang dikembangkan pada kegiatan Kerja Praktik ini, pengguna sistem dibagi ke dalam beberapa peran utama seperti administrator dan karyawan. Administrator memiliki hak akses yang lebih luas terhadap sistem, sedangkan karyawan hanya dapat mengakses fitur tertentu yang berkaitan dengan pencatatan aktivitas kerja harian.')

# --- Monitoring Aktivitas Kerja ---
hl('Monitoring Aktivitas Kerja')
ap('Monitoring aktivitas kerja merupakan proses pemantauan terhadap aktivitas yang dilakukan oleh karyawan dalam menjalankan tugasnya. Amadi-Echendu (2023) dalam artikelnya yang dipublikasikan pada SA Journal of Information Management menjelaskan bahwa monitoring ini bertujuan untuk memastikan bahwa setiap aktivitas kerja berjalan sesuai dengan perencanaan serta dapat digunakan sebagai bahan evaluasi kinerja.')
ap('Dalam perusahaan yang memiliki banyak aktivitas operasional, pencatatan aktivitas kerja yang dilakukan secara manual sering kali menimbulkan berbagai kendala seperti kesulitan dalam pengelolaan data, keterlambatan pelaporan, serta kesulitan dalam melakukan evaluasi kinerja.')
ap('Dengan adanya sistem ERP sederhana berbasis web, pencatatan aktivitas kerja dapat dilakukan secara terpusat dan terdokumentasi dengan baik sehingga mempermudah manajemen perusahaan dalam melakukan monitoring serta evaluasi terhadap aktivitas operasional perusahaan.')

# --- UML ---
hl('Unified Modeling Language (UML)')
ap('Unified Modeling Language (UML) merupakan bahasa pemodelan standar yang digunakan untuk memvisualisasikan, merancang, serta mendokumentasikan sistem perangkat lunak. A.S. Rosa dan Shalahuddin (2018) menjelaskan bahwa UML membantu pengembang sistem dalam memahami struktur serta alur kerja sistem sebelum sistem tersebut diimplementasikan.')
ap('Wati dan Kusumo (2016) dalam penelitiannya yang dipublikasikan pada Jurnal Terapan Teknologi Informasi menambahkan bahwa penggunaan UML dalam pengembangan sistem informasi memungkinkan pengembang dan pengguna sistem untuk memahami bagaimana sistem akan bekerja secara keseluruhan. Dengan menggunakan UML, proses perancangan sistem dapat dilakukan secara lebih terstruktur sehingga mempermudah proses implementasi sistem yang dikembangkan.')
ap('Dalam pengembangan sistem ERP sederhana berbasis web pada kegiatan Kerja Praktik ini, UML digunakan sebagai alat bantu untuk memodelkan sistem sebelum dilakukan proses implementasi sistem. Kedua referensi di atas menegaskan bahwa UML merupakan alat bantu esensial dalam tahap perancangan perangkat lunak.')

# --- Metode RAD ---
hl('Metode Rapid Application Development (RAD)')
ap('Rapid Application Development (RAD) merupakan metode pengembangan perangkat lunak yang menekankan pada siklus pengembangan yang singkat dengan kualitas yang tetap terjaga. Metode RAD pertama kali diperkenalkan oleh James Martin pada tahun 1991 sebagai alternatif dari metode pengembangan tradisional (waterfall) yang memiliki siklus pengembangan yang panjang dan kaku (Martin, 1991).')
ap('Dennis, Wixom, dan Roth (2012) menjelaskan bahwa RAD merupakan pendekatan pengembangan sistem yang menekankan pada siklus pengembangan yang pendek, penggunaan teknik iteratif, dan keterlibatan pengguna secara aktif dalam setiap tahap pengembangan. Tujuan utama RAD adalah memperpendek waktu yang biasanya diperlukan dalam siklus pengembangan sistem tradisional antara perancangan dan penerapan suatu sistem informasi.')
ap('Berdasarkan kedua referensi tersebut, RAD sangat cocok untuk proyek-proyek yang membutuhkan penyelesaian cepat namun tetap memperhatikan kualitas perangkat lunak. RAD memanfaatkan prototipe dan penggunaan kembali komponen perangkat lunak (reusable components) untuk mempercepat proses pengembangan.')

# --- Tahapan Metode RAD ---
hl('Tahapan Metode RAD')
ap('Metode RAD terdiri dari empat tahapan utama yang dilaksanakan secara berurutan namun dengan durasi yang relatif singkat. Berikut adalah penjelasan masing-masing tahapan (Martin, 1991; Dennis, Wixom, dan Roth, 2012):')

ap('**Requirements Planning (Perencanaan Kebutuhan)**')
ap('Tahap ini bertujuan untuk mengidentifikasi tujuan sistem, ruang lingkup proyek, dan kebutuhan informasi. Pengguna dan pengembang bekerja sama untuk menentukan kebutuhan fungsional dan non-fungsional sistem. Hasil dari tahap ini berupa daftar kebutuhan yang menjadi acuan pengembangan.', False)

ap('**User Design (Desain Pengguna)**')
ap('Tahap ini merupakan tahap perancangan yang dilakukan secara kolaboratif antara pengembang dan pengguna. Prototipe antarmuka dibuat dan dievaluasi secara iteratif hingga desain sesuai dengan kebutuhan. Keunggulan tahap ini adalah penggunaan prototipe visual yang dapat langsung dievaluasi oleh pengguna.', False)

ap('**Construction (Konstruksi)**')
ap('Tahap ini merupakan tahap pembangunan sistem secara intensif berdasarkan desain yang telah disetujui. Pengembang memanfaatkan fitur-fitur framework, komponen siap pakai, dan alat pengembangan modern untuk mempercepat proses pembangunan sistem.', False)

ap('**Cutover (Peralihan)**')
ap('Tahap akhir yang mencakup pengujian sistem secara menyeluruh, pelatihan pengguna, konversi data, dan peralihan dari sistem lama ke sistem baru. Pengujian dilakukan untuk memastikan seluruh fitur berjalan sesuai kebutuhan.', False)

# --- Kelebihan Metode RAD ---
hl('Kelebihan Metode RAD')
ap('Metode RAD memiliki beberapa kelebihan dibandingkan dengan metode pengembangan tradisional (Dennis, Wixom, dan Roth, 2012):')
li('Waktu pengembangan yang lebih singkat karena menggunakan pendekatan iteratif dan komponen siap pakai.', 1)
li('Keterlibatan pengguna yang tinggi memastikan hasil akhir sesuai dengan kebutuhan.', 2)
li('Fleksibilitas terhadap perubahan kebutuhan selama proses pengembangan.', 3)
li('Penggunaan prototipe memungkinkan evaluasi dini terhadap desain sistem.', 4)
li('Cocok untuk proyek dengan skala kecil hingga menengah dengan batasan waktu yang ketat.', 5)

# --- Arsitektur MVC ---
hl('Arsitektur Model-View-Controller (MVC)')
ap('Model-View-Controller (MVC) merupakan pola arsitektur perangkat lunak yang memisahkan aplikasi menjadi tiga komponen utama yang saling berinteraksi. Pop dan Altar (2014) dalam artikelnya yang dipublikasikan pada jurnal Procedia Engineering menjelaskan bahwa pola MVC bertujuan untuk memisahkan logika bisnis dari tampilan antarmuka pengguna sehingga memudahkan pengembangan, pengujian, dan pemeliharaan aplikasi.')
ap('Tiga komponen utama dalam arsitektur MVC adalah sebagai berikut (Pop dan Altar, 2014):')
li('**Model** \u2014 Komponen yang bertanggung jawab terhadap logika bisnis dan pengelolaan data. Model berinteraksi langsung dengan basis data untuk melakukan operasi CRUD (Create, Read, Update, Delete).', 1)
li('**View** \u2014 Komponen yang bertanggung jawab terhadap tampilan antarmuka pengguna. View menerima data dari Model melalui Controller dan menampilkannya kepada pengguna.', 2)
li('**Controller** \u2014 Komponen yang berfungsi sebagai perantara antara Model dan View. Controller menerima input dari pengguna, memprosesnya melalui Model, dan menentukan View yang akan ditampilkan.', 3)
ap('Penerapan arsitektur MVC pada pengembangan sistem ERP Prototype ini menggunakan framework Laravel yang secara native mendukung pola MVC. Pemisahan komponen ini memastikan kode program lebih terstruktur, mudah dipelihara, dan memungkinkan pengembangan secara paralel oleh tim pengembang.')

# --- Framework Laravel ---
hl('Framework Laravel')
ap('Laravel merupakan framework aplikasi web berbasis PHP yang bersifat open-source dan mengikuti pola arsitektur Model-View-Controller (MVC). Laravel diciptakan oleh Taylor Otwell pada tahun 2011 dengan tujuan menyediakan alternatif yang lebih elegan dan ekspresif untuk pengembangan aplikasi web dibandingkan framework PHP lainnya (Laravel Documentation, 2025).')
ap('Laravel menyediakan berbagai fitur bawaan yang mempercepat proses pengembangan, antara lain:')
li('**Eloquent ORM** \u2014 Object-Relational Mapping yang memudahkan interaksi dengan basis data melalui pendekatan berorientasi objek. Setiap tabel pada basis data direpresentasikan oleh model Eloquent yang mendefinisikan relasi, atribut, dan operasi data.', 1)
li('**Blade Template Engine** \u2014 Mesin template yang menyediakan fitur template inheritance dan data binding untuk membangun antarmuka pengguna secara modular.', 2)
li('**Artisan CLI** \u2014 Command-line interface yang menyediakan berbagai perintah untuk mempercepat tugas-tugas pengembangan seperti pembuatan model, controller, migration, dan seeder.', 3)
li('**Migration dan Seeder** \u2014 Fitur untuk mengelola struktur basis data secara terversikan (version control) dan mengisi data awal untuk pengujian.', 4)
li('**Middleware** \u2014 Mekanisme untuk memfilter permintaan HTTP yang masuk ke aplikasi, digunakan untuk autentikasi, otorisasi, dan proteksi CSRF.', 5)
li('**Routing** \u2014 Sistem routing yang fleksibel untuk mendefinisikan URL dan menghubungkannya dengan controller yang sesuai.', 6)
ap('Dalam penelitian ini, sistem ERP Prototype dibangun menggunakan Laravel versi 12.0 dengan PHP 8.2 sebagai bahasa pemrograman server-side.')

# --- Basis Data MySQL ---
hl('Basis Data MySQL')
ap('MySQL merupakan sistem manajemen basis data relasional (RDBMS) yang bersifat open-source dan merupakan salah satu RDBMS paling populer di dunia. MySQL menggunakan Structured Query Language (SQL) untuk mengelola dan memanipulasi data. A.S. Rosa dan Shalahuddin (2018) menjelaskan bahwa basis data relasional menyimpan data dalam bentuk tabel-tabel yang saling berhubungan melalui relasi kunci utama (primary key) dan kunci asing (foreign key).')
ap('Dalam sistem ERP Prototype ini, MySQL digunakan untuk menyimpan seluruh data operasional termasuk data pengguna, role, divisi, aktivitas, dan laporan. Fitur migration pada Laravel memungkinkan pengelolaan struktur basis data secara terversikan sehingga perubahan skema dapat dilacak dan dikelola dengan baik.')

# --- TailwindCSS ---
hl('Tailwind CSS')
ap('TailwindCSS merupakan framework CSS berbasis pendekatan utility-first yang memungkinkan pengembang membangun antarmuka pengguna secara cepat dengan menulis class-class utilitas langsung pada elemen HTML. Berbeda dengan framework CSS tradisional yang menyediakan komponen siap pakai, TailwindCSS memberikan kontrol penuh terhadap desain melalui class-class atomik yang dapat dikombinasikan secara fleksibel (Tailwind Labs, 2025).')
ap('Keunggulan TailwindCSS antara lain:')
li('Desain responsif bawaan dengan sistem breakpoint yang mudah digunakan.', 1)
li('Konsistensi desain melalui sistem design token (spacing, warna, typography).', 2)
li('Ukuran file CSS akhir yang kecil karena fitur tree-shaking yang menghapus class yang tidak digunakan.', 3)
li('Dukungan dark mode melalui konfigurasi class-based atau media query.', 4)
ap('Dalam penelitian ini, TailwindCSS versi 4.0 digunakan untuk membangun antarmuka pengguna sistem ERP yang responsif, modern, dan mendukung fitur dark mode.')

# --- Alpine.js ---
hl('Alpine.js')
ap('Alpine.js merupakan framework JavaScript minimalis yang dirancang untuk menambahkan interaktivitas pada halaman web tanpa memerlukan framework JavaScript yang besar seperti React atau Vue.js. Alpine.js menyediakan pendekatan deklaratif untuk menangani state management, event handling, dan manipulasi DOM langsung pada markup HTML (Alpine.js Documentation, 2024).')
ap('Dalam sistem ERP Prototype ini, Alpine.js digunakan untuk mengimplementasikan fitur-fitur interaktif seperti toggle dark mode, dropdown menu, modal dialog, dan notifikasi. Penggunaan Alpine.js sejalan dengan prinsip RAD karena memungkinkan penambahan interaktivitas secara cepat tanpa kompleksitas yang tinggi.')

# --- Chart.js ---
hl('Chart.js')
ap('Chart.js merupakan library JavaScript open-source untuk membuat visualisasi data berbasis grafik pada halaman web. Library ini mendukung berbagai jenis grafik seperti line chart, bar chart, pie chart, doughnut chart, radar chart, dan scatter plot. Chart.js menggunakan elemen HTML5 Canvas untuk merender grafik yang responsif dan interaktif (Chart.js Documentation, 2024).')
ap('Dalam sistem ERP Prototype ini, Chart.js digunakan pada dashboard admin untuk menampilkan grafik line chart yang memvisualisasikan tren aktivitas karyawan dan super admin selama 12 bulan terakhir. Visualisasi data ini membantu manajemen dalam memahami pola aktivitas dan mengambil keputusan berbasis data.')

# --- Laravel Breeze ---
hl('Laravel Breeze')
ap('Laravel Breeze merupakan paket scaffolding autentikasi resmi dari Laravel yang menyediakan implementasi fitur-fitur autentikasi dasar secara minimal dan sederhana. Fitur yang disediakan mencakup login, registrasi, reset password, verifikasi email, dan konfirmasi password (Laravel Documentation, 2025).')
ap('Dalam sistem ERP Prototype ini, Laravel Breeze versi 2.3 digunakan sebagai dasar sistem autentikasi. Proses autentikasi memanfaatkan facade Auth::attempt() untuk memvalidasi kredensial pengguna, session regeneration untuk keamanan, serta redirect berdasarkan role pengguna (Super Admin atau Karyawan).')

# --- Black-Box Testing ---
hl('Pengujian Black-Box Testing')
ap('Black-box testing merupakan metode pengujian perangkat lunak yang berfokus pada fungsionalitas sistem dari perspektif pengguna tanpa memperhatikan struktur internal kode program. Nidhra dan Dondeti (2012) dalam artikelnya yang dipublikasikan pada International Journal of Embedded Systems and Applications mendefinisikan black-box testing sebagai teknik pengujian yang mengevaluasi apakah input yang diberikan menghasilkan output yang sesuai dengan spesifikasi kebutuhan yang telah ditetapkan.')
ap('Nidhra dan Dondeti (2012) menjelaskan bahwa black-box testing memiliki beberapa teknik pengujian, antara lain:')
li('**Equivalence Partitioning** \u2014 Membagi domain input menjadi kelas-kelas ekuivalensi dan menguji satu nilai dari setiap kelas.', 1)
li('**Boundary Value Analysis** \u2014 Menguji nilai-nilai pada batas domain input.', 2)
li('**Decision Table Testing** \u2014 Menggunakan tabel keputusan untuk menguji kombinasi kondisi input.', 3)
ap('Dalam penelitian ini, metode black-box testing digunakan untuk menguji seluruh fungsionalitas sistem ERP Prototype meliputi proses login, manajemen pengguna, monitoring aktivitas, pembuatan laporan, dan fitur-fitur pendukung lainnya.')

doc.save(r'd:\laragon\www\erp-prototype\bab_2_update.docx')
print('File saved: bab_2_update.docx')
