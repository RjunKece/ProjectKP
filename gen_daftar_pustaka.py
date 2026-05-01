from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

doc = Document()
for s in doc.sections:
    s.top_margin=Cm(3); s.bottom_margin=Cm(3); s.left_margin=Cm(4); s.right_margin=Cm(3)
sn = doc.styles['Normal']
sn.font.name='Times New Roman'; sn.font.size=Pt(12); sn.font.color.rgb=RGBColor(0,0,0)
sn.paragraph_format.line_spacing=1.5; sn.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
sn.element.rPr.rFonts.set(qn('w:eastAsia'),'Times New Roman')

def hc(t,sz=14):
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run(t); r.bold=True; r.font.name='Times New Roman'; r.font.size=Pt(sz); r.font.color.rgb=RGBColor(0,0,0)
    p.paragraph_format.line_spacing=1.5

def ref(text, italic_parts=None):
    """Add a reference with hanging indent. italic_parts is list of strings to italicize."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.left_indent = Cm(1.27)
    p.paragraph_format.first_line_indent = Cm(-1.27)
    p.paragraph_format.space_after = Pt(6)
    
    if italic_parts:
        remaining = text
        for ip in italic_parts:
            idx = remaining.find(ip)
            if idx >= 0:
                before = remaining[:idx]
                if before:
                    r = p.add_run(before)
                    r.font.name = 'Times New Roman'; r.font.size = Pt(12); r.font.color.rgb = RGBColor(0,0,0)
                r = p.add_run(ip)
                r.font.name = 'Times New Roman'; r.font.size = Pt(12); r.font.color.rgb = RGBColor(0,0,0)
                r.italic = True
                remaining = remaining[idx+len(ip):]
        if remaining:
            r = p.add_run(remaining)
            r.font.name = 'Times New Roman'; r.font.size = Pt(12); r.font.color.rgb = RGBColor(0,0,0)
    else:
        r = p.add_run(text)
        r.font.name = 'Times New Roman'; r.font.size = Pt(12); r.font.color.rgb = RGBColor(0,0,0)
    return p

hc('DAFTAR PUSTAKA')

# === Semua referensi dari Bab 1, 2 (update), 3, 4, 5 - diurutkan alfabetis ===

ref('Abd Elmonem, M. A., Nasr, E. S., & Geith, M. H. (2017). Benefits and Challenges of Cloud ERP Systems \u2013 A Systematic Literature Review. Future Computing and Informatics Journal, 2(1), 1\u20139. https://doi.org/10.1016/j.fcij.2017.03.003',
    ['Future Computing and Informatics Journal'])

ref('Alam, M. J., Sharma, S. D., & Chowdhury, T. (2024). A Model of Cloud-Based Enterprise Resource Planning (ERP) for Small and Medium Enterprise. Journal of Computer and Communications, 12, 37\u201350. https://doi.org/10.4236/jcc.2024.1210004',
    ['Journal of Computer and Communications'])

ref('Alpine.js. (2024). Alpine.js Documentation. Diakses dari https://alpinejs.dev/docs',
    ['Alpine.js Documentation'])

ref('Amadi-Echendu, A. P. (2023). Using Enterprise Resource Planning Systems to Manage Property Processes. SA Journal of Information Management, 25(1), a1607. https://doi.org/10.4102/sajim.v25i1.1607',
    ['SA Journal of Information Management'])

ref('A.S., Rosa & Shalahuddin, M. (2018). Rekayasa Perangkat Lunak Terstruktur dan Berorientasi Objek. Bandung: Informatika Bandung.',
    ['Rekayasa Perangkat Lunak Terstruktur dan Berorientasi Objek'])

ref('Chart.js. (2024). Chart.js Documentation. Diakses dari https://www.chartjs.org/docs/latest/',
    ['Chart.js Documentation'])

ref('Daraghmi, Y.-A., & Daraghmi, E.-Y. (2022). RAPD: Rapid and Participatory Application Development of Usable Systems During COVID19 Crisis. IEEE Access, 10, 93601\u201393614. https://doi.org/10.1109/ACCESS.2022.3203582',
    ['IEEE Access'])

ref('Dennis, A., Wixom, B. H., & Roth, R. M. (2012). Systems Analysis and Design (5th ed.). New York: John Wiley & Sons.',
    ['Systems Analysis and Design'])

ref('Dumas, M., La Rosa, M., Mendling, J., & Reijers, H. A. (2018). Fundamentals of Business Process Management (2nd ed.). Berlin: Springer. https://doi.org/10.1007/978-3-662-56509-4',
    ['Fundamentals of Business Process Management'])

ref('Haddara, M., & Fagerstr\u00f8m, A. (2015). Cloud ERP Systems: Anatomy of Adoption Factors & Attitudes. Journal of Enterprise Resource Planning Studies, 2015, Article 521212. https://doi.org/10.5171/2015.521212',
    ['Journal of Enterprise Resource Planning Studies'])

ref('Jaffa, S., & Salim, S. A. (2020). A Review of Cloud-Based ERP Systems in SMEs. International Journal of Integrated Engineering, 12(7), 113\u2013120. https://doi.org/10.30880/ijie.2020.12.07.014',
    ['International Journal of Integrated Engineering'])

ref('Katuu, S. (2020). Enterprise Resource Planning: Past, Present, and Future. New Review of Information Networking, 25(1), 37\u201346. https://doi.org/10.1080/13614576.2020.1742770',
    ['New Review of Information Networking'])

ref('Kendall, K. E., & Kendall, J. E. (2019). Systems Analysis and Design (10th ed.). London: Pearson.',
    ['Systems Analysis and Design'])

ref('Laravel. (2025). Laravel 12.x Documentation. Diakses dari https://laravel.com/docs/12.x',
    ['Laravel 12.x Documentation'])

ref('Lutfi, A., Alshira\u2019h, A. F., Alshirah, M. H., Al-Okaily, M., Alqudah, H., Saad, M., Ibrahim, N., & Abdelmaksoud, O. (2022). Antecedents and Impacts of Enterprise Resource Planning System Adoption among Jordanian SMEs. Sustainability, 14(6), 3508. https://doi.org/10.3390/su14063508',
    ['Sustainability'])

ref('Martin, J. (1991). Rapid Application Development. New York: Macmillan Publishing Co., Inc.',
    ['Rapid Application Development'])

ref('Nidhra, S., & Dondeti, J. (2012). Black Box and White Box Testing Techniques \u2013 A Literature Review. International Journal of Embedded Systems and Applications, 2(2), 29\u201350. https://doi.org/10.5121/ijesa.2012.2204',
    ['International Journal of Embedded Systems and Applications'])

ref('Panorama Consulting Group. (2024). 2024 ERP Report. Diakses dari https://www.panorama-consulting.com/resource-center/erp-report/',
    ['2024 ERP Report'])

ref('P\u00e9rez Est\u00e9banez, R. (2024). An Approach to Sustainable Enterprise Resource Planning System Implementation in Small- and Medium-Sized Enterprises. Administrative Sciences, 14(4), 86. https://doi.org/10.3390/admsci14040086',
    ['Administrative Sciences'])

ref('Pop, D. P., & Altar, A. (2014). Designing an MVC Model for Rapid Web Application Development. Procedia Engineering, 69, 1172\u20131179. https://doi.org/10.1016/j.proeng.2014.03.106',
    ['Procedia Engineering'])

ref('Pressman, R. S., & Maxim, B. R. (2019). Software Engineering: A Practitioner\u2019s Approach (9th ed.). New York: McGraw-Hill Education.',
    ['Software Engineering: A Practitioner\u2019s Approach'])

ref('Salih, S., Hamdan, M., Abdelmaboud, A., Abdelsalam, S., Althobaiti, M. M., Cheikhrouhou, O., Hamam, H., & Alotaibi, F. (2021). Prioritising Organisational Factors Impacting Cloud ERP Adoption and the Critical Issues Related to Security, Usability, and Vendors: A Systematic Literature Review. Sensors, 21(24), 8391. https://doi.org/10.3390/s21248391',
    ['Sensors'])

ref('Sommerville, I. (2016). Software Engineering (10th ed.). London: Pearson.',
    ['Software Engineering'])

ref('Sulaimon, K., Surin, E. F., & Hamzah, M. I. (2024). Open-Source Enterprise Resource Planning Systems for Small and Medium Enterprises: A Conceptual Framework. Journal of International Business, Economics and Entrepreneurship, 9(2), 1\u201316. https://doi.org/10.24191/jibe.v9i2.3575',
    ['Journal of International Business, Economics and Entrepreneurship'])

ref('Tailwind Labs. (2025). Tailwind CSS Documentation. Diakses dari https://tailwindcss.com/docs',
    ['Tailwind CSS Documentation'])

ref('Wati, E. F., & Kusumo, A. A. (2016). Penerapan Metode Unified Modeling Language (UML) Berbasis Desktop Pada Sistem Informasi Point of Sale (POS). Jurnal Terapan Teknologi Informasi, 1(1), 1\u201310.',
    ['Jurnal Terapan Teknologi Informasi'])

doc.save(r'd:\laragon\www\erp-prototype\daftar_pustaka.docx')
print('File saved: daftar_pustaka.docx')
