from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()
for s in doc.sections:
    s.top_margin=Cm(3); s.bottom_margin=Cm(3); s.left_margin=Cm(4); s.right_margin=Cm(3)

sn = doc.styles['Normal']
sn.font.name='Times New Roman'; sn.font.size=Pt(12); sn.font.color.rgb=RGBColor(0,0,0)
sn.paragraph_format.line_spacing=1.5; sn.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
sn.element.rPr.rFonts.set(qn('w:eastAsia'),'Times New Roman')

def hc(t,sz=14):
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run(t); r.bold=True; r.font.name='Times New Roman'; r.font.size=Pt(sz); r.font.color.rgb=RGBColor(0,0,0)
    p.paragraph_format.line_spacing=1.5; return p

def hl(t):
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    r=p.add_run(t); r.bold=True; r.font.name='Times New Roman'; r.font.size=Pt(12); r.font.color.rgb=RGBColor(0,0,0)
    p.paragraph_format.line_spacing=1.5; p.paragraph_format.space_before=Pt(18); return p

def ap(t,indent=True):
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY; p.paragraph_format.line_spacing=1.5
    if indent: p.paragraph_format.first_line_indent=Cm(1.27)
    parts=t.split('**')
    for i,part in enumerate(parts):
        r=p.add_run(part); r.font.name='Times New Roman'; r.font.size=Pt(12); r.font.color.rgb=RGBColor(0,0,0)
        if i%2==1: r.bold=True
    return p

def li(t,num=None):
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY; p.paragraph_format.line_spacing=1.5
    p.paragraph_format.left_indent=Cm(1.27); p.paragraph_format.first_line_indent=Cm(0)
    prefix=f"{num}. " if num else "\u2022 "
    parts=(prefix+t).split('**')
    for i,part in enumerate(parts):
        r=p.add_run(part); r.font.name='Times New Roman'; r.font.size=Pt(12); r.font.color.rgb=RGBColor(0,0,0)
        if i%2==1: r.bold=True
    return p

def cap(t):
    """Caption centered italic"""
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.line_spacing=1.5
    r=p.add_run(t); r.font.name='Times New Roman'; r.font.size=Pt(11); r.font.color.rgb=RGBColor(0,0,0); r.italic=True
    return p

def tbl(headers, rows):
    table=doc.add_table(rows=1+len(rows),cols=len(headers))
    table.style='Table Grid'; table.alignment=WD_TABLE_ALIGNMENT.CENTER
    for i,h in enumerate(headers):
        c=table.rows[0].cells[i]; c.text=''; p=c.paragraphs[0]
        r=p.add_run(h); r.bold=True; r.font.name='Times New Roman'; r.font.size=Pt(11); r.font.color.rgb=RGBColor(0,0,0)
        p.paragraph_format.line_spacing=1.5
    for ri,row in enumerate(rows):
        for ci,val in enumerate(row):
            c=table.rows[ri+1].cells[ci]; c.text=''; p=c.paragraphs[0]
            r=p.add_run(val); r.font.name='Times New Roman'; r.font.size=Pt(11); r.font.color.rgb=RGBColor(0,0,0)
            p.paragraph_format.line_spacing=1.5
    doc.add_paragraph()
    return table

# ===== BAB III =====
hc('BAB III')
hc('METODOLOGI PENELITIAN')

# --- 3.1 Metode Pengembangan Sistem ---
hl('Metode Pengembangan Sistem')
ap('Pengembangan prototipe sistem ERP pada PT. Golden Intan Berlian menggunakan metode Rapid Application Development (RAD). Metode RAD dipilih karena sifatnya yang iteratif dan berfokus pada pembuatan prototipe fungsional dalam waktu singkat, sehingga sangat sesuai dengan kebutuhan proyek Kerja Praktik yang memiliki keterbatasan durasi pelaksanaan. Dibandingkan dengan metode pengembangan tradisional seperti Waterfall yang memiliki siklus pengembangan panjang dan linear, RAD menawarkan fleksibilitas yang lebih tinggi dalam mengakomodasi perubahan kebutuhan selama proses pengembangan berlangsung.')
ap('RAD adalah pendekatan pengembangan sistem informasi yang menekankan pada siklus pengembangan yang pendek, penggunaan teknik iteratif, dan keterlibatan pengguna secara aktif dalam setiap tahap pengembangan. Dengan pendekatan ini, pengembang dapat menghasilkan prototipe yang fungsional secara cepat untuk kemudian dievaluasi dan diperbaiki secara bertahap. Pendekatan RAD juga memungkinkan pengembang untuk memanfaatkan komponen perangkat lunak yang sudah tersedia (reusable components) seperti framework Laravel beserta ekosistem pendukungnya, sehingga proses pembangunan sistem dapat dilakukan secara lebih efisien tanpa harus membangun setiap komponen dari awal.')

# --- 3.1.1 Fase-Fase Metode RAD ---
hl('Fase-Fase Metode RAD')
ap('Metode RAD yang diterapkan terdiri dari empat fase utama:')

ap('**Fase 1: Requirements Planning (Perencanaan Kebutuhan).**')
ap('Pada fase ini dilakukan identifikasi kebutuhan sistem melalui wawancara dan observasi langsung di PT. Golden Intan Berlian. Kegiatan yang dilakukan meliputi analisis proses bisnis yang sedang berjalan, identifikasi permasalahan pada sistem manual yang digunakan sebelumnya, penentuan kebutuhan fungsional dan non-fungsional sistem, serta pendefinisian aktor dan peran pengguna dalam sistem. Hasil dari fase ini adalah dokumen kebutuhan sistem yang menjadi acuan utama dalam tahap perancangan dan implementasi selanjutnya.', False)

ap('**Fase 2: Workshop Design (Perancangan).**')
ap('Fase ini merupakan inti dari metode RAD dimana perancangan dilakukan secara iteratif dengan melibatkan umpan balik dari pengguna. Pada fase ini dilakukan perancangan basis data menggunakan Entity Relationship Diagram (ERD), perancangan alur sistem menggunakan Use Case Diagram dan Activity Diagram, serta perancangan desain antarmuka pengguna (UI/UX) dalam bentuk wireframe. Prototipe antarmuka yang dihasilkan dievaluasi dan diperbaiki secara berulang hingga sesuai dengan kebutuhan pengguna.', False)

ap('**Fase 3: Construction (Implementasi).**')
ap('Fase implementasi mencakup proses pengkodean sistem menggunakan framework Laravel 12 dengan arsitektur Model-View-Controller (MVC), pembuatan struktur basis data MySQL melalui mekanisme migration, serta pengembangan fitur-fitur sistem berdasarkan prioritas kebutuhan yang telah ditetapkan pada fase sebelumnya. Pada fase ini, pengembang memanfaatkan berbagai fitur bawaan Laravel seperti Eloquent ORM, Blade Template Engine, dan Artisan CLI untuk mempercepat proses pembangunan sistem.', False)

ap('**Fase 4: Cutover (Peralihan).**')
ap('Fase terakhir mencakup pengujian sistem secara menyeluruh menggunakan metode Black-Box Testing untuk memvalidasi seluruh fungsionalitas sistem dari perspektif pengguna akhir. Selain pengujian, pada fase ini juga dilakukan penyusunan dokumentasi teknis, pelatihan penggunaan sistem kepada pengguna, serta penyerahan prototipe sistem kepada pihak perusahaan sebagai hasil akhir dari proyek Kerja Praktik.', False)

# --- 3.2 Teknik Pengumpulan Data ---
hl('Teknik Pengumpulan Data')
ap('Teknik pengumpulan data yang digunakan dalam penelitian ini meliputi tiga metode utama sebagai berikut:')

ap('**Observasi.**')
ap('Observasi dilakukan secara langsung di lingkungan kerja PT. Golden Intan Berlian selama periode pelaksanaan Kerja Praktik. Aspek-aspek yang diamati meliputi proses pengelolaan data karyawan dari delapan divisi operasional, mekanisme pencatatan aktivitas harian yang dilakukan secara manual menggunakan spreadsheet, alur komunikasi antara manajemen dan karyawan melalui aplikasi pesan instan, serta mekanisme pemantauan kinerja karyawan yang masih bersifat konvensional. Hasil observasi ini menjadi dasar dalam mengidentifikasi kelemahan sistem berjalan dan menentukan kebutuhan fungsional sistem yang akan dikembangkan.', False)

ap('**Wawancara.**')
ap('Wawancara dilakukan dengan pihak manajemen dan karyawan PT. Golden Intan Berlian untuk mendapatkan pemahaman mendalam mengenai kebutuhan fungsional sistem, permasalahan yang dihadapi pada sistem manual yang berjalan saat ini, serta prioritas fitur yang paling dibutuhkan oleh pengguna. Melalui wawancara ini diperoleh informasi penting mengenai kebutuhan akan sistem monitoring aktivitas secara real-time, pengelolaan pengguna berbasis role, dan kemampuan pembuatan laporan yang terintegrasi.', False)

ap('**Studi Pustaka.**')
ap('Studi pustaka dilakukan untuk mempelajari konsep-konsep teoritis yang mendasari pengembangan sistem, meliputi konsep Enterprise Resource Planning (ERP), metode Rapid Application Development (RAD), framework Laravel dan arsitektur MVC, teknik perancangan basis data relasional, serta metode pengujian Black-Box Testing. Sumber literatur yang digunakan mencakup buku referensi, jurnal ilmiah, dan dokumentasi resmi dari teknologi yang digunakan dalam pengembangan sistem.', False)

# --- 3.3 Alat dan Bahan Penelitian ---
hl('Alat dan Bahan Penelitian')
ap('Berikut adalah alat dan bahan yang digunakan dalam pengembangan sistem ERP Prototype ini, yang terdiri dari perangkat keras dan perangkat lunak.')

hl('Perangkat Keras')
ap('Tabel 3.1 menampilkan spesifikasi perangkat keras minimal yang digunakan selama proses pengembangan sistem ERP Prototype. Perangkat keras ini dipilih berdasarkan kebutuhan minimum untuk menjalankan server lokal Laragon, editor kode VS Code, browser untuk pengujian, serta proses kompilasi aset frontend menggunakan Vite dan Node.js secara bersamaan.')
cap('Tabel 3.1 Perangkat Keras')
tbl(['No','Perangkat','Spesifikasi'],[
    ['1','Laptop/PC','Processor Intel/AMD, RAM minimal 8 GB'],
    ['2','Monitor','Resolusi minimal 1366x768'],
])

hl('Perangkat Lunak')
ap('Tabel 3.2 menampilkan daftar lengkap perangkat lunak beserta versi dan fungsinya yang digunakan dalam pengembangan sistem. Pemilihan perangkat lunak ini disesuaikan dengan kebutuhan pengembangan sistem ERP berbasis web menggunakan framework Laravel 12. Seluruh perangkat lunak yang digunakan bersifat open-source atau memiliki lisensi gratis untuk penggunaan pengembangan, sehingga tidak memerlukan biaya lisensi tambahan.')
cap('Tabel 3.2 Perangkat Lunak')
tbl(['No','Software','Versi','Fungsi'],[
    ['1','PHP','8.3','Bahasa pemrograman server-side'],
    ['2','Laravel','12','Framework MVC backend'],
    ['3','MySQL','8.0','Sistem manajemen basis data'],
    ['4','Node.js','20.x','Runtime pengelolaan aset frontend'],
    ['5','Vite','7.x','Build tool dan module bundler'],
    ['6','TailwindCSS','4.0','Framework CSS utility-first'],
    ['7','Alpine.js','3.x','Library JavaScript reaktif'],
    ['8','Chart.js','4.x','Library visualisasi data'],
    ['9','Laragon','6.0','Environment lokal pengembangan'],
    ['10','VS Code','Latest','Editor kode sumber'],
    ['11','Git','Latest','Kontrol versi'],
])

# --- 3.4 Analisis Sistem ---
hl('Analisis Sistem')

hl('Sistem Berjalan')
ap('Sebelum dikembangkannya sistem ERP, PT. Golden Intan Berlian mengelola seluruh operasional bisnis secara manual menggunakan spreadsheet dan aplikasi pesan instan. Kondisi ini menimbulkan beberapa kelemahan signifikan yang berdampak pada efisiensi operasional perusahaan, yaitu: (1) pencatatan aktivitas karyawan dilakukan tanpa format standar sehingga data tidak konsisten antar divisi, (2) monitoring kinerja karyawan tidak dapat dilakukan secara real-time karena data tersebar di berbagai media, (3) proses pelaporan manajerial memerlukan waktu yang lama karena harus mengkompilasi data dari berbagai sumber secara manual, dan (4) manajemen akun pengguna tidak terpusat sehingga sulit untuk mengelola hak akses dan keamanan data.')

hl('Sistem yang Diusulkan')
ap('Untuk mengatasi permasalahan tersebut, dikembangkan sistem ERP berbasis web yang menyediakan platform terpusat mencakup lima modul utama. Modul pertama adalah Modul Autentikasi yang menerapkan Role-Based Access Control (RBAC) untuk mengelola hak akses pengguna berdasarkan perannya. Modul kedua adalah Modul Dashboard yang menyajikan visualisasi Key Performance Indicator (KPI) dan grafik tren aktivitas. Modul ketiga adalah Modul Manajemen Aktivitas yang memungkinkan pencatatan dan pemantauan aktivitas harian karyawan. Modul keempat adalah Modul Laporan yang mendukung pembuatan dan pengelolaan laporan perusahaan maupun divisi. Modul kelima adalah Modul User Management yang memfasilitasi pengelolaan data pengguna secara terpusat. Kelima modul ini dirancang secara terintegrasi untuk mengatasi seluruh permasalahan yang ditemukan pada sistem berjalan.')

# --- 3.5 Perancangan Sistem ---
hl('Perancangan Sistem')
ap('Perancangan sistem dilakukan menggunakan pendekatan UML (Unified Modeling Language) yang mencakup Use Case Diagram, Activity Diagram, dan Entity Relationship Diagram (ERD). Berikut adalah detail perancangan masing-masing komponen.')

# --- Use Case Diagram ---
hl('Use Case Diagram')
ap('Use Case Diagram merupakan diagram UML yang menggambarkan interaksi antara aktor (pengguna) dengan fungsionalitas sistem. Dalam sistem ERP Prototype ini terdapat dua aktor utama, yaitu Super Admin yang memiliki hak akses penuh terhadap seluruh fitur sistem, dan Karyawan yang memiliki hak akses terbatas sesuai dengan perannya. Gambar 3.1 menampilkan Use Case Diagram lengkap yang menunjukkan seluruh fungsionalitas sistem berdasarkan peran masing-masing aktor. Diagram ini menjadi acuan utama dalam pengembangan fitur-fitur sistem pada tahap implementasi.')
cap('Gambar 3.1 Use Case Diagram Sistem ERP PT. Golden Intan Berlian')
ap('[Gambar Use Case Diagram disisipkan di sini]', False)

ap('Berikut adalah penjelasan detail use case untuk masing-masing aktor:')
ap('**Aktor 1: Super Admin**')
ap('Tabel 3.3 menjelaskan daftar use case yang dapat dilakukan oleh Super Admin. Super Admin memiliki hak akses penuh terhadap seluruh fitur sistem, termasuk pengelolaan pengguna, monitoring aktivitas seluruh karyawan, dan pengelolaan laporan.')
cap('Tabel 3.3 Use Case Super Admin')
tbl(['No','Use Case','Deskripsi'],[
    ['UC-01','Login','Masuk ke sistem menggunakan email dan password'],
    ['UC-02','Melihat Dashboard','Melihat ringkasan KPI, grafik tren, dan statistik'],
    ['UC-03','Mengelola User','Menambah user, melihat daftar, reset password'],
    ['UC-04','Monitoring Aktivitas','Memantau seluruh aktivitas karyawan'],
    ['UC-05','Mengelola Laporan','Membuat, melihat, dan menghapus laporan'],
    ['UC-06','Mengubah Tema','Beralih antara mode terang dan gelap'],
    ['UC-07','Logout','Keluar dari sistem dengan aman'],
])

ap('**Aktor 2: Karyawan**')
ap('Tabel 3.4 menjelaskan daftar use case yang dapat dilakukan oleh Karyawan. Karyawan memiliki akses terbatas pada fitur pencatatan aktivitas harian, melihat laporan, dan mengelola profil pribadi.')
cap('Tabel 3.4 Use Case Karyawan')
tbl(['No','Use Case','Deskripsi'],[
    ['UC-08','Login','Masuk ke sistem menggunakan email dan password'],
    ['UC-09','Melihat Dashboard','Melihat ringkasan kinerja personal dan grafik'],
    ['UC-10','Mencatat Aktivitas','Menambahkan aktivitas harian beserta lampiran'],
    ['UC-11','Mengedit Aktivitas','Mengubah status, deskripsi, dan lampiran'],
    ['UC-12','Melihat Laporan','Mengakses laporan perusahaan dan divisi'],
    ['UC-13','Membalas Laporan','Memberikan tanggapan terhadap laporan'],
    ['UC-14','Mengubah Tema','Beralih antara mode terang dan gelap'],
    ['UC-15','Logout','Keluar dari sistem dengan aman'],
])

# --- Activity Diagram ---
hl('Activity Diagram')
ap('Activity Diagram merupakan diagram UML yang digunakan untuk menggambarkan alur kerja (workflow) dari setiap proses utama dalam sistem secara berurutan. Diagram ini menunjukkan langkah-langkah yang dilakukan oleh aktor dan sistem, termasuk percabangan keputusan (decision) dan alur alternatif yang mungkin terjadi. Berikut adalah Activity Diagram untuk lima proses utama dalam sistem ERP Prototype yang mencakup proses login, pencatatan aktivitas, monitoring aktivitas, pembuatan laporan, dan proses logout.')

ap('Gambar 3.2 menampilkan Activity Diagram proses login yang menggambarkan alur autentikasi pengguna mulai dari input kredensial, validasi oleh sistem, hingga redirect ke dashboard sesuai role pengguna.')
cap('Gambar 3.2 Activity Diagram Proses Login')
ap('[Gambar Activity Diagram Login disisipkan di sini]', False)

ap('Gambar 3.3 menampilkan Activity Diagram pencatatan aktivitas yang dilakukan oleh karyawan. Alur dimulai dari pengisian formulir aktivitas, pengunggahan lampiran (opsional), hingga penyimpanan data ke basis data.')
cap('Gambar 3.3 Activity Diagram Pencatatan Aktivitas')
ap('[Gambar Activity Diagram Pencatatan Aktivitas disisipkan di sini]', False)

ap('Gambar 3.4 menampilkan Activity Diagram monitoring aktivitas yang dilakukan oleh Super Admin. Admin dapat melihat, memfilter, dan menganalisis seluruh aktivitas karyawan dari berbagai divisi.')
cap('Gambar 3.4 Activity Diagram Monitoring Aktivitas')
ap('[Gambar Activity Diagram Monitoring Aktivitas disisipkan di sini]', False)

ap('Gambar 3.5 menampilkan Activity Diagram pembuatan laporan yang dapat dilakukan oleh Super Admin maupun Karyawan. Proses meliputi pengisian formulir laporan, penentuan scope dan prioritas, hingga penyimpanan laporan.')
cap('Gambar 3.5 Activity Diagram Pembuatan Laporan')
ap('[Gambar Activity Diagram Pembuatan Laporan disisipkan di sini]', False)

ap('Gambar 3.6 menampilkan Activity Diagram proses logout yang mencakup penghancuran sesi, regenerasi token CSRF, dan redirect ke halaman login dengan pesan konfirmasi.')
cap('Gambar 3.6 Activity Diagram Proses Logout')
ap('[Gambar Activity Diagram Logout disisipkan di sini]', False)

# --- ERD ---
hl('Entity Relationship Diagram (ERD)')
ap('Entity Relationship Diagram (ERD) merupakan diagram yang menggambarkan struktur basis data beserta relasi antar entitas dalam sistem. ERD menjadi acuan utama dalam pembuatan struktur tabel basis data pada tahap implementasi menggunakan mekanisme migration Laravel. Gambar 3.7 menampilkan ERD sistem ERP Prototype yang terdiri dari tujuh entitas utama beserta atribut dan relasi antar entitas.')
cap('Gambar 3.7 Entity Relationship Diagram Sistem ERP')
ap('[Gambar ERD disisipkan di sini]', False)

ap('Entitas utama dalam sistem meliputi: Roles, Divisions, Users, Activities, Daily_Targets, Reports, dan Report_Responses. Relasi antar entitas meliputi: Roles-Users (1:N), Divisions-Users (1:N), Users-Activities (1:N), Users-Reports (1:N), Divisions-Reports (1:N), Reports-ReportResponses (1:N), dan Activities-DailyTargets (N:1). Setiap relasi menggunakan foreign key dengan cascade delete untuk menjaga integritas referensial data.')

# --- Struktur Tabel Database ---
hl('Struktur Tabel Database')
ap('Berikut adalah struktur detail dari setiap tabel basis data yang digunakan dalam sistem ERP Prototype. Setiap tabel dirancang sesuai dengan kebutuhan fungsional sistem dan diimplementasikan melalui mekanisme migration pada framework Laravel. Penggunaan migration memungkinkan pengelolaan struktur basis data secara terversikan (version control) sehingga setiap perubahan skema dapat dilacak dan dikelola dengan baik oleh tim pengembang.')

ap('Tabel 3.5 menampilkan struktur tabel **roles** yang menyimpan data peran pengguna dalam sistem.')
cap('Tabel 3.5 Struktur Tabel Roles')
tbl(['Kolom','Tipe','Keterangan'],[
    ['id','BIGINT (PK, AI)','Primary key'],
    ['nama_role','VARCHAR (UNIQUE)','super_admin, admin, karyawan'],
    ['created_at','TIMESTAMP','Waktu pembuatan'],
    ['updated_at','TIMESTAMP','Waktu pembaruan'],
])

ap('Tabel 3.6 menampilkan struktur tabel **divisions** yang menyimpan data divisi operasional perusahaan.')
cap('Tabel 3.6 Struktur Tabel Divisions')
tbl(['Kolom','Tipe','Keterangan'],[
    ['id','BIGINT (PK, AI)','Primary key'],
    ['nama_divisi','VARCHAR (UNIQUE)','Nama divisi'],
    ['created_at','TIMESTAMP','Waktu pembuatan'],
    ['updated_at','TIMESTAMP','Waktu pembaruan'],
])

ap('Tabel 3.7 menampilkan struktur tabel **users** yang menyimpan data pengguna sistem beserta relasinya dengan tabel roles dan divisions.')
cap('Tabel 3.7 Struktur Tabel Users')
tbl(['Kolom','Tipe','Keterangan'],[
    ['id','BIGINT (PK, AI)','Primary key'],
    ['name','VARCHAR','Nama lengkap pengguna'],
    ['email','VARCHAR (UNIQUE)','Email untuk login'],
    ['password','VARCHAR','Password ter-hash (Bcrypt)'],
    ['role_id','BIGINT (FK)','FK ke tabel roles'],
    ['division_id','BIGINT (FK, nullable)','FK ke tabel divisions'],
    ['created_at','TIMESTAMP','Waktu pembuatan'],
    ['updated_at','TIMESTAMP','Waktu pembaruan'],
])

ap('Tabel 3.8 menampilkan struktur tabel **activities** yang menyimpan data aktivitas harian karyawan beserta lampiran pendukung.')
cap('Tabel 3.8 Struktur Tabel Activities')
tbl(['Kolom','Tipe','Keterangan'],[
    ['id','BIGINT (PK, AI)','Primary key'],
    ['user_id','BIGINT (FK)','FK ke tabel users'],
    ['tanggal','DATETIME','Tanggal dan waktu aktivitas'],
    ['deskripsi','TEXT','Deskripsi aktivitas'],
    ['status','VARCHAR','submitted, in_progress, completed'],
    ['file_path','VARCHAR (nullable)','Path file lampiran'],
    ['file_name','VARCHAR (nullable)','Nama asli file'],
    ['link','VARCHAR (nullable)','URL tautan pendukung'],
    ['target_id','BIGINT (FK, nullable)','FK ke daily_targets'],
])

ap('Tabel 3.9 menampilkan struktur tabel **reports** yang menyimpan data laporan yang dibuat oleh pengguna sistem.')
cap('Tabel 3.9 Struktur Tabel Reports')
tbl(['Kolom','Tipe','Keterangan'],[
    ['id','BIGINT (PK, AI)','Primary key'],
    ['title','VARCHAR','Judul laporan'],
    ['type','VARCHAR','financial, department, growth, custom'],
    ['scope','VARCHAR','company atau division'],
    ['division_id','BIGINT (FK, nullable)','FK ke divisions'],
    ['description','TEXT (nullable)','Isi laporan'],
    ['start_date','DATE (nullable)','Tanggal mulai'],
    ['end_date','DATE (nullable)','Tanggal akhir'],
    ['status','ENUM','generated, ready, failed'],
    ['priority','ENUM','low, normal, high, urgent'],
    ['created_by','BIGINT (FK, nullable)','FK ke users'],
])

ap('Tabel 3.10 menampilkan struktur tabel **report_responses** yang menyimpan data tanggapan terhadap laporan.')
cap('Tabel 3.10 Struktur Tabel Report_Responses')
tbl(['Kolom','Tipe','Keterangan'],[
    ['id','BIGINT (PK, AI)','Primary key'],
    ['report_id','BIGINT (FK)','FK ke reports'],
    ['user_id','BIGINT (FK)','FK ke users'],
    ['message','TEXT','Isi pesan balasan'],
    ['type','ENUM','reply atau acknowledgment'],
])

ap('Tabel 3.11 menampilkan struktur tabel **daily_targets** yang menyimpan data target harian untuk setiap divisi.')
cap('Tabel 3.11 Struktur Tabel Daily_Targets')
tbl(['Kolom','Tipe','Keterangan'],[
    ['id','BIGINT (PK, AI)','Primary key'],
    ['division_id','BIGINT (FK)','FK ke divisions'],
    ['title','VARCHAR','Judul target harian'],
    ['target_count','INT (nullable)','Jumlah target'],
    ['unit','VARCHAR (nullable)','Satuan (video, deal, dll)'],
    ['is_default','BOOLEAN','Target dasar perusahaan'],
    ['is_active','BOOLEAN','Status aktif'],
])

# --- Perancangan Antarmuka ---
hl('Perancangan Antarmuka')
ap('Perancangan antarmuka (wireframe) dilakukan untuk memberikan gambaran visual mengenai tata letak dan komponen antarmuka sebelum proses implementasi. Berikut adalah rancangan antarmuka untuk tiga halaman utama sistem.')

ap('Gambar 3.8 menampilkan rancangan halaman login dengan desain split-screen yang menampilkan formulir login di satu sisi dan branding perusahaan di sisi lainnya. Rancangan ini juga dilengkapi dengan fitur toggle tema dark/light mode.')
cap('Gambar 3.8 Rancangan Halaman Login (Dark Mode)')
ap('[Gambar wireframe halaman login disisipkan di sini]', False)

ap('Gambar 3.9 menampilkan rancangan dashboard admin yang terdiri dari sidebar navigasi, area KPI cards, dan grafik tren aktivitas. Dashboard ini dirancang untuk memberikan overview lengkap mengenai performa perusahaan kepada Super Admin.')
cap('Gambar 3.9 Rancangan Dashboard Admin')
ap('[Gambar wireframe dashboard admin disisipkan di sini]', False)

ap('Gambar 3.10 menampilkan rancangan dashboard karyawan yang menampilkan ringkasan kinerja personal, daftar aktivitas terbaru, dan grafik pencapaian target harian.')
cap('Gambar 3.10 Rancangan Dashboard Karyawan')
ap('[Gambar wireframe dashboard karyawan disisipkan di sini]', False)

# --- Metode Pengujian ---
hl('Metode Pengujian')
ap('Pengujian sistem dilakukan menggunakan metode Black-Box Testing, yaitu teknik pengujian yang berfokus pada fungsionalitas sistem tanpa memperhatikan struktur internal kode. Metode ini dipilih karena sesuai dengan kebutuhan pengujian prototipe yang mengutamakan validasi fungsionalitas dari perspektif pengguna akhir.')
ap('Aspek yang diuji meliputi: fungsionalitas autentikasi (login dan logout), operasi CRUD (Create, Read, Update, Delete) pada setiap modul, validasi input pada formulir, responsivitas tampilan pada berbagai ukuran layar, serta konsistensi tema dark/light mode di seluruh halaman sistem.')

doc.save(r'd:\laragon\www\erp-prototype\bab_3_update_v2.docx')
print('File saved: bab_3_update_v2.docx')
