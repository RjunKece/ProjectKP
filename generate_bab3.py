from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()
for section in doc.sections:
    section.top_margin = Cm(3)
    section.bottom_margin = Cm(3)
    section.left_margin = Cm(4)
    section.right_margin = Cm(3)

style_normal = doc.styles['Normal']
style_normal.font.name = 'Times New Roman'
style_normal.font.size = Pt(12)
style_normal.font.color.rgb = RGBColor(0, 0, 0)
style_normal.paragraph_format.line_spacing = 1.25
style_normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
style_normal.element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

def hc(text, size=14):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text); r.bold = True; r.font.name = 'Times New Roman'; r.font.size = Pt(size); r.font.color.rgb = RGBColor(0,0,0)
    p.paragraph_format.line_spacing = 1.25
    return p

def hl(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(text); r.bold = True; r.font.name = 'Times New Roman'; r.font.size = Pt(12); r.font.color.rgb = RGBColor(0,0,0)
    p.paragraph_format.line_spacing = 1.25; p.paragraph_format.space_before = Pt(12)
    return p

def ap(text, indent=True):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY; p.paragraph_format.line_spacing = 1.25
    if indent: p.paragraph_format.first_line_indent = Cm(1.27)
    r = p.add_run(text); r.font.name = 'Times New Roman'; r.font.size = Pt(12); r.font.color.rgb = RGBColor(0,0,0)
    return p

def li(text, num=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY; p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.left_indent = Cm(1.27); p.paragraph_format.first_line_indent = Cm(0)
    prefix = f"{num}. " if num else "\u2022 "
    parts = (prefix + text).split('**')
    for i, part in enumerate(parts):
        r = p.add_run(part); r.font.name = 'Times New Roman'; r.font.size = Pt(12); r.font.color.rgb = RGBColor(0,0,0)
        if i % 2 == 1: r.bold = True
    return p

def tbl(headers, rows):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = 'Table Grid'; table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = table.rows[0].cells[i]; c.text = ''; p = c.paragraphs[0]
        r = p.add_run(h); r.bold = True; r.font.name = 'Times New Roman'; r.font.size = Pt(11); r.font.color.rgb = RGBColor(0,0,0)
        p.paragraph_format.line_spacing = 1.25
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            c = table.rows[ri+1].cells[ci]; c.text = ''; p = c.paragraphs[0]
            r = p.add_run(val); r.font.name = 'Times New Roman'; r.font.size = Pt(11); r.font.color.rgb = RGBColor(0,0,0)
            p.paragraph_format.line_spacing = 1.25
    doc.add_paragraph()
    return table

# ===== BAB III =====
hc('BAB III')
hc('METODOLOGI PENELITIAN')

# 3.1
hl('3.1 Metode Pengembangan Sistem')
ap('Metode pengembangan sistem yang digunakan dalam pembangunan Sistem ERP Prototype PT. Golden Intan Berlian adalah metode Rapid Application Development (RAD). Metode RAD pertama kali diperkenalkan oleh James Martin pada tahun 1991 sebagai pendekatan pengembangan perangkat lunak yang menekankan pada siklus pengembangan yang singkat dengan kualitas yang tetap terjaga.')
ap('Metode RAD dipilih karena beberapa pertimbangan sebagai berikut:')
li('**Waktu pengembangan yang terbatas** \u2014 Proyek ini memiliki batasan waktu yang ketat sehingga diperlukan metode yang mampu menghasilkan sistem dalam waktu singkat.', 1)
li('**Ketersediaan framework modern** \u2014 Penggunaan framework Laravel 12 dan library pendukung seperti TailwindCSS, Alpine.js, dan Chart.js memungkinkan pengembangan secara cepat melalui pemanfaatan komponen siap pakai (reusable components).', 2)
li('**Kebutuhan prototipe** \u2014 Sistem yang dikembangkan bersifat prototipe sehingga memerlukan pendekatan iteratif yang memungkinkan evaluasi dan perbaikan secara berkelanjutan.', 3)
li('**Keterlibatan pengguna** \u2014 RAD mendorong keterlibatan aktif pengguna dalam proses desain antarmuka, sehingga hasil akhir lebih sesuai dengan kebutuhan pengguna.', 4)

# 3.2
hl('3.2 Tahapan Metode RAD')
ap('Metode RAD terdiri dari empat tahapan utama yang dilaksanakan secara berurutan namun dengan durasi yang relatif singkat pada setiap tahapannya. Berikut adalah penjelasan dari masing-masing tahapan:')

hl('3.2.1 Requirements Planning (Perencanaan Kebutuhan)')
ap('Tahap Requirements Planning merupakan tahap awal dalam metode RAD yang bertujuan untuk mengidentifikasi tujuan sistem, ruang lingkup proyek, dan kebutuhan informasi yang diperlukan. Pada tahap ini dilakukan kegiatan sebagai berikut:')
li('Melakukan analisis terhadap proses bisnis yang berjalan di PT. Golden Intan Berlian, khususnya dalam hal pengelolaan data karyawan, monitoring aktivitas, dan pembuatan laporan.', 1)
li('Mengidentifikasi permasalahan yang ada pada proses manual atau sistem yang sedang berjalan.', 2)
li('Menentukan kebutuhan fungsional dan non-fungsional sistem ERP yang akan dibangun.', 3)
li('Mengidentifikasi aktor-aktor yang akan berinteraksi dengan sistem beserta hak akses masing-masing.', 4)
li('Menyusun prioritas fitur berdasarkan tingkat kepentingan dan urgensi.', 5)
ap('Hasil dari tahap ini berupa daftar kebutuhan fungsional, kebutuhan non-fungsional, dan identifikasi aktor sistem yang menjadi acuan untuk tahap selanjutnya.')

hl('3.2.2 User Design (Desain Pengguna)')
ap('Tahap User Design merupakan tahap perancangan sistem yang dilakukan secara kolaboratif dengan calon pengguna. Pada tahap ini, pengembang dan pengguna bekerja sama untuk merancang sistem yang memenuhi kebutuhan yang telah diidentifikasi. Kegiatan pada tahap ini meliputi:')
li('Merancang struktur basis data yang terdiri dari tabel-tabel utama beserta relasi antar tabel menggunakan fitur migration pada Laravel.', 1)
li('Merancang desain antarmuka pengguna (User Interface) untuk setiap halaman sistem menggunakan pendekatan prototipe langsung dengan Blade template engine dan TailwindCSS.', 2)
li('Merancang alur navigasi dan interaksi pengguna dengan sistem.', 3)
li('Melakukan evaluasi desain bersama pengguna dan memperbaiki desain berdasarkan umpan balik yang diterima.', 4)
ap('Keunggulan tahap ini dalam metode RAD adalah penggunaan prototipe visual yang dapat langsung dievaluasi oleh pengguna, sehingga meminimalkan kesalahan desain sebelum memasuki tahap pembangunan.')

hl('3.2.3 Construction (Konstruksi)')
ap('Tahap Construction merupakan tahap pembangunan sistem secara intensif berdasarkan desain yang telah disetujui pada tahap sebelumnya. Pada tahap ini, pengembang melakukan kegiatan sebagai berikut:')
li('Mengimplementasikan arsitektur Model-View-Controller (MVC) sesuai standar framework Laravel.', 1)
li('Membangun model-model Eloquent beserta relasi antar model.', 2)
li('Mengembangkan controller untuk menangani logika bisnis setiap modul.', 3)
li('Membuat view menggunakan Blade template engine dengan layout inheritance.', 4)
li('Mengimplementasikan sistem autentikasi menggunakan Laravel Breeze.', 5)
li('Mengkonfigurasi routing dengan pengelompokan berdasarkan middleware dan prefix.', 6)
li('Mengembangkan service layer untuk fungsi-fungsi pendukung seperti logging aktivitas.', 7)
li('Mengimplementasikan fitur-fitur pendukung seperti dark mode, grafik analitik, dan data seeder.', 8)
ap('Prinsip RAD diterapkan pada tahap ini dengan memanfaatkan fitur-fitur bawaan Laravel seperti Artisan CLI, Eloquent ORM, dan Blade Components untuk mempercepat proses pengembangan secara signifikan.')

hl('3.2.4 Cutover (Peralihan)')
ap('Tahap Cutover merupakan tahap akhir dalam metode RAD yang mencakup pengujian sistem secara menyeluruh, pelatihan pengguna, dan peralihan dari sistem lama ke sistem baru. Kegiatan pada tahap ini meliputi:')
li('Melakukan pengujian fungsional menggunakan metode black-box testing untuk memastikan setiap fitur berjalan sesuai dengan kebutuhan yang telah ditetapkan.', 1)
li('Melakukan pengujian validasi input untuk memastikan keamanan dan integritas data.', 2)
li('Melakukan migrasi data awal menggunakan database seeder.', 3)
li('Mengevaluasi kinerja sistem secara keseluruhan.', 4)
li('Mendokumentasikan hasil pengujian dan temuan selama proses pengembangan.', 5)

# 3.3
hl('3.3 Alur Metode RAD')
ap('Berikut adalah ilustrasi alur tahapan metode RAD yang diterapkan dalam pengembangan Sistem ERP Prototype PT. Golden Intan Berlian:')
tbl(['Tahap', 'Kegiatan Utama', 'Output'], [
    ['1. Requirements Planning', 'Analisis kebutuhan, identifikasi aktor, penentuan prioritas fitur', 'Daftar kebutuhan fungsional dan non-fungsional'],
    ['2. User Design', 'Perancangan database, desain UI/UX, evaluasi prototipe', 'Struktur database, prototipe antarmuka'],
    ['3. Construction', 'Implementasi MVC, coding modul, integrasi komponen', 'Sistem ERP fungsional'],
    ['4. Cutover', 'Black-box testing, validasi input, migrasi data', 'Laporan pengujian, sistem siap digunakan'],
])

# 3.4
hl('3.4 Alat dan Bahan Penelitian')

hl('3.4.1 Perangkat Keras (Hardware)')
ap('Perangkat keras yang digunakan dalam pengembangan sistem ini adalah sebagai berikut:')
tbl(['No', 'Perangkat Keras', 'Spesifikasi'], [
    ['1', 'Laptop/PC', 'Processor Intel/AMD, RAM minimal 8 GB'],
    ['2', 'Penyimpanan', 'SSD minimal 256 GB'],
    ['3', 'Monitor', 'Resolusi minimal 1920 x 1080'],
])

hl('3.4.2 Perangkat Lunak (Software)')
ap('Perangkat lunak yang digunakan dalam pengembangan sistem ini adalah sebagai berikut:')
tbl(['No', 'Perangkat Lunak', 'Fungsi', 'Versi'], [
    ['1', 'Laravel', 'Backend framework', '12.0'],
    ['2', 'PHP', 'Bahasa pemrograman server-side', '8.2'],
    ['3', 'MySQL', 'Database management system', '-'],
    ['4', 'Laragon', 'Local development server', '-'],
    ['5', 'Visual Studio Code', 'Code editor', 'Latest'],
    ['6', 'TailwindCSS', 'CSS framework', '4.0'],
    ['7', 'Alpine.js', 'JavaScript framework (reactive UI)', '3.x'],
    ['8', 'Vite', 'Frontend build tool', '7.0'],
    ['9', 'Chart.js', 'Library grafik/chart', 'Latest'],
    ['10', 'Laravel Breeze', 'Authentication scaffolding', '2.3'],
    ['11', 'Composer', 'PHP dependency manager', '-'],
    ['12', 'Node.js & NPM', 'JavaScript runtime & package manager', '-'],
    ['13', 'Git', 'Version control system', '-'],
    ['14', 'Google Chrome', 'Browser untuk pengujian', 'Latest'],
])

# 3.5
hl('3.5 Teknik Pengumpulan Data')
ap('Teknik pengumpulan data yang digunakan dalam penelitian ini meliputi:')
li('**Observasi** \u2014 Melakukan pengamatan langsung terhadap proses bisnis yang berjalan di PT. Golden Intan Berlian, khususnya pada delapan divisi yang ada yaitu Marketing, Sales, Keuangan, Konten Kreator, Gudang, CRM, YouTube, dan Admin Marketplace.', 1)
li('**Wawancara** \u2014 Melakukan wawancara dengan pihak manajemen dan perwakilan karyawan untuk memahami kebutuhan sistem ERP yang diperlukan.', 2)
li('**Studi Literatur** \u2014 Mempelajari referensi terkait pengembangan sistem ERP, framework Laravel, dan metode RAD dari buku, jurnal, dan dokumentasi resmi.', 3)

# 3.6
hl('3.6 Teknik Pengujian Sistem')
ap('Pengujian sistem dilakukan menggunakan metode Black-Box Testing. Metode ini dipilih karena fokus pada pengujian fungsionalitas sistem dari perspektif pengguna tanpa memperhatikan struktur internal kode program. Pengujian black-box mengevaluasi apakah input yang diberikan menghasilkan output yang sesuai dengan kebutuhan yang telah ditetapkan.')
ap('Aspek-aspek yang diuji dalam black-box testing meliputi:')
li('**Validasi input** \u2014 Menguji apakah sistem menolak input yang tidak valid dan menampilkan pesan error yang sesuai.', 1)
li('**Fungsionalitas fitur** \u2014 Menguji apakah setiap fitur (login, CRUD user, monitoring aktivitas, pembuatan laporan) berjalan sesuai kebutuhan.', 2)
li('**Navigasi dan routing** \u2014 Menguji apakah sistem mengarahkan pengguna ke halaman yang tepat berdasarkan role dan hak akses.', 3)
li('**Keamanan dasar** \u2014 Menguji apakah mekanisme autentikasi dan proteksi CSRF berfungsi dengan baik.', 4)

doc.save(r'd:\laragon\www\erp-prototype\laporan_bab3_rad.docx')
print('File saved: laporan_bab3_rad.docx')
