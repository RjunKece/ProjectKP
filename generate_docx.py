from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

# Page setup
for s in doc.sections:
    s.top_margin = Cm(3)
    s.bottom_margin = Cm(3)
    s.left_margin = Cm(4)
    s.right_margin = Cm(3)

style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)
style.paragraph_format.line_spacing = 1.5

def heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.name = 'Times New Roman'
        r.font.color.rgb = RGBColor(0,0,0)
        r.font.size = Pt(14) if level==1 else Pt(12)
    h.paragraph_format.line_spacing = 1.5
    return h

def para(text, bold=False, align=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(12)
    r.bold = bold
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.first_line_indent = Cm(1.25)
    if align: p.alignment = align
    return p

def add_table(headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = h
        for p in c.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.name = 'Times New Roman'
                r.font.size = Pt(11)
    for row in rows:
        r = t.add_row()
        for i, val in enumerate(row):
            r.cells[i].text = str(val)
            for p in r.cells[i].paragraphs:
                for run in p.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(11)
    doc.add_paragraph()

def img_placeholder(caption):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f'[Lampirkan Gambar Di Sini]')
    r.font.name = 'Times New Roman'
    r.font.size = Pt(11)
    r.italic = True
    r.font.color.rgb = RGBColor(128,128,128)
    c = doc.add_paragraph()
    c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cr = c.add_run(caption)
    cr.font.name = 'Times New Roman'
    cr.font.size = Pt(11)
    cr.bold = True
    doc.add_paragraph()

# ============================================================
# BAB III
# ============================================================
heading('BAB III\nMETODOLOGI PENELITIAN', 1).alignment = WD_ALIGN_PARAGRAPH.CENTER

heading('3.1 Metode Pengembangan Sistem', 2)
para('Pengembangan prototipe sistem ERP pada PT. Golden Intan Berlian menggunakan metode Rapid Application Development (RAD). Metode RAD dipilih karena sifatnya yang iteratif dan berfokus pada pembuatan prototipe fungsional dalam waktu singkat, sehingga sangat sesuai dengan kebutuhan proyek Kerja Praktik yang memiliki keterbatasan durasi pelaksanaan.')
para('Menurut Kendall & Kendall (2014), RAD adalah pendekatan pengembangan sistem informasi yang menekankan pada siklus pengembangan yang pendek, penggunaan teknik iteratif, dan keterlibatan pengguna secara aktif dalam setiap tahap pengembangan.')

heading('3.1.1 Fase-Fase Metode RAD', 3)
para('Metode RAD yang diterapkan terdiri dari empat fase utama:')
para('Fase 1: Requirements Planning (Perencanaan Kebutuhan). Pada fase ini dilakukan identifikasi kebutuhan sistem melalui wawancara dan observasi langsung di PT. Golden Intan Berlian, meliputi analisis proses bisnis, identifikasi permasalahan sistem manual, penentuan kebutuhan fungsional dan non-fungsional, serta pendefinisian aktor dan peran dalam sistem.')
para('Fase 2: Workshop Design (Perancangan). Fase ini merupakan inti dari metode RAD dimana perancangan dilakukan secara iteratif, meliputi perancangan basis data (ERD), alur sistem (Use Case dan Activity Diagram), serta desain antarmuka pengguna (UI/UX).')
para('Fase 3: Construction (Implementasi). Fase implementasi mencakup pengkodean menggunakan framework Laravel 12, pembuatan basis data MySQL melalui migration, dan pengembangan fitur berdasarkan prioritas kebutuhan.')
para('Fase 4: Cutover (Peralihan). Fase terakhir mencakup pengujian akhir dengan Black-Box Testing, dokumentasi, dan penyerahan prototipe.')

heading('3.2 Teknik Pengumpulan Data', 2)
para('3.2.1 Observasi. Observasi dilakukan secara langsung di lingkungan kerja PT. Golden Intan Berlian untuk mengamati proses pengelolaan data karyawan, pencatatan aktivitas harian, alur komunikasi manajemen-karyawan, dan mekanisme pemantauan kinerja.', bold=True)
para('3.2.2 Wawancara. Wawancara dilakukan untuk mendapatkan pemahaman mendalam mengenai kebutuhan fungsional, permasalahan sistem manual, dan prioritas fitur.', bold=True)
para('3.2.3 Studi Pustaka. Studi pustaka dilakukan untuk mempelajari konsep ERP, metode RAD, framework Laravel, dan teknik perancangan basis data relasional.', bold=True)

heading('3.3 Alat dan Bahan Penelitian', 2)
heading('3.3.1 Perangkat Keras', 3)
add_table(['No','Perangkat','Spesifikasi'],[
    ['1','Laptop/PC','Processor Intel/AMD, RAM minimal 8 GB'],
    ['2','Monitor','Resolusi minimal 1366x768'],
])

heading('3.3.2 Perangkat Lunak', 3)
add_table(['No','Software','Versi','Fungsi'],[
    ['1','PHP','8.3','Bahasa pemrograman server-side'],
    ['2','Laravel','12','Framework MVC backend'],
    ['3','MySQL','8.0','Sistem manajemen basis data'],
    ['4','Node.js','20.x','Runtime pengelolaan aset frontend'],
    ['5','Vite','7.x','Build tool dan module bundler'],
    ['6','TailwindCSS','4.0','Framework CSS utility-first'],
    ['7','Alpine.js','3.x','Library JavaScript reaktif'],
    ['8','Chart.js','4.x','Library visualisasi data'],
    ['9','Laragon','6.0','Environment lokal pengembangan'],
    ['10','VS Code','Latest','Editor kode sumber'],
    ['11','Git','Latest','Kontrol versi'],
])

heading('3.4 Analisis Sistem', 2)
heading('3.4.1 Sistem Berjalan', 3)
para('Sebelum dikembangkannya sistem ERP, PT. Golden Intan Berlian mengelola operasional secara manual menggunakan spreadsheet dan aplikasi pesan instan. Kelemahan: (1) pencatatan aktivitas tanpa format standar, (2) monitoring kinerja tidak real-time, (3) pelaporan memerlukan waktu lama, (4) manajemen akun tidak terpusat.')

heading('3.4.2 Sistem yang Diusulkan', 3)
para('Sistem ERP yang dikembangkan menyediakan platform terpusat mencakup: Modul Autentikasi (RBAC), Modul Dashboard (visualisasi KPI), Modul Manajemen Aktivitas, Modul Laporan, dan Modul User Management.')

heading('3.5 Perancangan Sistem', 2)
heading('3.5.1 Use Case Diagram', 3)
img_placeholder('Gambar 3.1 — Use Case Diagram Sistem ERP PT. Golden Intan Berlian')

para('Aktor 1: Super Admin', bold=True)
add_table(['No','Use Case','Deskripsi'],[
    ['UC-01','Login','Masuk ke sistem menggunakan email dan password'],
    ['UC-02','Melihat Dashboard','Melihat ringkasan KPI, grafik tren, dan statistik'],
    ['UC-03','Mengelola User','Menambah user, melihat daftar, reset password'],
    ['UC-04','Monitoring Aktivitas','Memantau seluruh aktivitas karyawan'],
    ['UC-05','Mengelola Laporan','Membuat, melihat, dan menghapus laporan'],
    ['UC-06','Mengubah Tema','Beralih antara mode terang dan gelap'],
    ['UC-07','Logout','Keluar dari sistem dengan aman'],
])

para('Aktor 2: Karyawan', bold=True)
add_table(['No','Use Case','Deskripsi'],[
    ['UC-08','Login','Masuk ke sistem menggunakan email dan password'],
    ['UC-09','Melihat Dashboard','Melihat ringkasan kinerja personal dan grafik'],
    ['UC-10','Mencatat Aktivitas','Menambahkan aktivitas harian beserta lampiran'],
    ['UC-11','Mengedit Aktivitas','Mengubah status, deskripsi, dan lampiran'],
    ['UC-12','Melihat Laporan','Mengakses laporan perusahaan dan divisi'],
    ['UC-13','Membalas Laporan','Memberikan tanggapan terhadap laporan'],
    ['UC-14','Mengubah Tema','Beralih antara mode terang dan gelap'],
    ['UC-15','Logout','Keluar dari sistem dengan aman'],
])

heading('3.5.2 Activity Diagram', 3)
for title, num in [('Proses Login','3.2'),('Pencatatan Aktivitas','3.3'),('Monitoring Aktivitas','3.4'),('Pembuatan Laporan','3.5'),('Proses Logout','3.6')]:
    img_placeholder(f'Gambar {num} — Activity Diagram {title}')

heading('3.5.3 Entity Relationship Diagram (ERD)', 3)
img_placeholder('Gambar 3.7 — Entity Relationship Diagram Sistem ERP')
para('Entitas utama: Roles, Divisions, Users, Activities, Daily_Targets, Reports, dan Report_Responses. Relasi meliputi: Roles-Users (1:N), Divisions-Users (1:N), Users-Activities (1:N), Users-Reports (1:N), Divisions-Reports (1:N), Reports-ReportResponses (1:N), Activities-DailyTargets (N:1).')

heading('3.5.4 Struktur Tabel Database', 3)

para('Tabel roles', bold=True)
add_table(['Kolom','Tipe','Keterangan'],[
    ['id','BIGINT (PK, AI)','Primary key'],
    ['nama_role','VARCHAR (UNIQUE)','super_admin, admin, karyawan'],
    ['created_at','TIMESTAMP','Waktu pembuatan'],
    ['updated_at','TIMESTAMP','Waktu pembaruan'],
])

para('Tabel divisions', bold=True)
add_table(['Kolom','Tipe','Keterangan'],[
    ['id','BIGINT (PK, AI)','Primary key'],
    ['nama_divisi','VARCHAR (UNIQUE)','Nama divisi'],
    ['created_at','TIMESTAMP','Waktu pembuatan'],
    ['updated_at','TIMESTAMP','Waktu pembaruan'],
])

para('Tabel users', bold=True)
add_table(['Kolom','Tipe','Keterangan'],[
    ['id','BIGINT (PK, AI)','Primary key'],
    ['name','VARCHAR','Nama lengkap pengguna'],
    ['email','VARCHAR (UNIQUE)','Email untuk login'],
    ['password','VARCHAR','Password ter-hash (Bcrypt)'],
    ['role_id','BIGINT (FK)','FK ke tabel roles'],
    ['division_id','BIGINT (FK, nullable)','FK ke tabel divisions'],
    ['created_at','TIMESTAMP','Waktu pembuatan'],
    ['updated_at','TIMESTAMP','Waktu pembaruan'],
])

para('Tabel activities', bold=True)
add_table(['Kolom','Tipe','Keterangan'],[
    ['id','BIGINT (PK, AI)','Primary key'],
    ['user_id','BIGINT (FK)','FK ke tabel users'],
    ['tanggal','DATETIME','Tanggal dan waktu aktivitas'],
    ['deskripsi','TEXT','Deskripsi aktivitas'],
    ['status','VARCHAR','submitted, in_progress, completed'],
    ['file_path','VARCHAR (nullable)','Path file lampiran'],
    ['file_name','VARCHAR (nullable)','Nama asli file'],
    ['link','VARCHAR (nullable)','URL tautan pendukung'],
    ['target_id','BIGINT (FK, nullable)','FK ke daily_targets'],
])

para('Tabel reports', bold=True)
add_table(['Kolom','Tipe','Keterangan'],[
    ['id','BIGINT (PK, AI)','Primary key'],
    ['title','VARCHAR','Judul laporan'],
    ['type','VARCHAR','financial, department, growth, custom'],
    ['scope','VARCHAR','company atau division'],
    ['division_id','BIGINT (FK, nullable)','FK ke divisions'],
    ['description','TEXT (nullable)','Isi laporan'],
    ['start_date','DATE (nullable)','Tanggal mulai'],
    ['end_date','DATE (nullable)','Tanggal akhir'],
    ['status','ENUM','generated, ready, failed'],
    ['priority','ENUM','low, normal, high, urgent'],
    ['created_by','BIGINT (FK, nullable)','FK ke users'],
])

para('Tabel report_responses', bold=True)
add_table(['Kolom','Tipe','Keterangan'],[
    ['id','BIGINT (PK, AI)','Primary key'],
    ['report_id','BIGINT (FK)','FK ke reports'],
    ['user_id','BIGINT (FK)','FK ke users'],
    ['message','TEXT','Isi pesan balasan'],
    ['type','ENUM','reply atau acknowledgment'],
])

para('Tabel daily_targets', bold=True)
add_table(['Kolom','Tipe','Keterangan'],[
    ['id','BIGINT (PK, AI)','Primary key'],
    ['division_id','BIGINT (FK)','FK ke divisions'],
    ['title','VARCHAR','Judul target harian'],
    ['target_count','INT (nullable)','Jumlah target'],
    ['unit','VARCHAR (nullable)','Satuan (video, deal, dll)'],
    ['is_default','BOOLEAN','Target dasar perusahaan'],
    ['is_active','BOOLEAN','Status aktif'],
])

heading('3.6 Perancangan Antarmuka', 2)
for cap in ['Gambar 3.8 — Rancangan Halaman Login (Dark Mode)','Gambar 3.9 — Rancangan Halaman Login (Light Mode)','Gambar 3.10 — Rancangan Dashboard Admin','Gambar 3.11 — Rancangan Dashboard Karyawan','Gambar 3.12 — Rancangan Sidebar (Expanded & Collapsed)']:
    img_placeholder(cap)

heading('3.7 Metode Pengujian', 2)
para('Pengujian sistem dilakukan menggunakan metode Black-Box Testing, yaitu teknik pengujian yang berfokus pada fungsionalitas sistem tanpa memperhatikan struktur internal kode. Aspek yang diuji meliputi: fungsionalitas autentikasi, operasi CRUD, validasi input, responsivitas tampilan, dan konsistensi tema dark/light mode.')

# ============================================================
# BAB IV
# ============================================================
doc.add_page_break()
heading('BAB IV\nHASIL DAN PEMBAHASAN', 1).alignment = WD_ALIGN_PARAGRAPH.CENTER

heading('4.1 Hasil Implementasi Sistem', 2)
para('Berdasarkan perancangan pada Bab III, prototipe sistem ERP PT. Golden Intan Berlian berhasil diimplementasikan menggunakan framework Laravel 12 dengan arsitektur MVC. Sistem berjalan pada PHP 8.3 dan MySQL 8.0.')

heading('4.1.1 Implementasi Basis Data', 3)
para('Basis data diimplementasikan dengan mekanisme migration Laravel. Total 7 tabel utama berhasil dibuat dengan relasi foreign key dan cascade delete untuk menjaga integritas referensial.')

add_table(['No','Nama Tabel','Jumlah Kolom','Keterangan'],[
    ['1','roles','3','Data hak akses'],
    ['2','divisions','3','8 divisi perusahaan'],
    ['3','users','9','Data pengguna + relasi'],
    ['4','activities','11','Log aktivitas harian'],
    ['5','daily_targets','10','Target harian per divisi'],
    ['6','reports','12','Laporan perusahaan/divisi'],
    ['7','report_responses','6','Tanggapan laporan'],
])

heading('4.1.2 Implementasi Arsitektur MVC', 3)
para('A. Model (Eloquent ORM) — 7 model: User, Role, Division, Activity, Report, ReportResponse, DailyTarget.', bold=True)
para('B. Controller — 9 controller menangani logika bisnis:', bold=True)
add_table(['Controller','Fungsi Utama'],[
    ['LoginController','Autentikasi login dengan validasi'],
    ['AdminDashboardController','KPI, grafik tren, analytics'],
    ['UserController','CRUD user, reset password'],
    ['ActivityController','Monitoring aktivitas karyawan'],
    ['ReportController','Kelola laporan perusahaan & divisi'],
    ['KaryawanDashboardController','Dashboard personal + grafik'],
    ['KaryawanActivityController','Catat aktivitas + upload file'],
    ['KaryawanReportController','Lihat & balas laporan'],
    ['ProfileController','Kelola profil & ubah password'],
])
para('C. View — Menggunakan Blade Template dengan 3 layout utama: admin, karyawan, dan auth.', bold=True)

heading('4.2 Implementasi Antarmuka', 2)

heading('4.2.1 Halaman Login', 3)
for cap in ['Gambar 4.1 — Tampilan Login (Dark Mode)','Gambar 4.2 — Tampilan Login (Light Mode)','Gambar 4.3 — Tampilan Pesan Error Login']:
    img_placeholder(cap)
para('Fitur login: split-screen design, theme toggle, animated error message, password visibility toggle, session message setelah logout, dan CSRF protection.')

heading('4.2.2 Dashboard Admin', 3)
for cap in ['Gambar 4.4 — Dashboard Admin (KPI)','Gambar 4.5 — Dashboard Admin (Grafik)']:
    img_placeholder(cap)

heading('4.2.3 User Management', 3)
for cap in ['Gambar 4.6 — Halaman User Management','Gambar 4.7 — Modal Tambah User']:
    img_placeholder(cap)

heading('4.2.4 Monitoring Aktivitas', 3)
for cap in ['Gambar 4.8 — Monitoring Aktivitas (KPI & Grafik)','Gambar 4.9 — Tabel Monitoring Aktivitas']:
    img_placeholder(cap)

heading('4.2.5 Halaman Reports', 3)
img_placeholder('Gambar 4.10 — Halaman Reports Admin')

heading('4.2.6 Dashboard & Aktivitas Karyawan', 3)
for cap in ['Gambar 4.11 — Dashboard Karyawan','Gambar 4.12 — Halaman Aktivitas Karyawan']:
    img_placeholder(cap)

heading('4.2.7 Sidebar Navigation', 3)
for cap in ['Gambar 4.13 — Sidebar Light Mode (Expanded vs Collapsed)','Gambar 4.14 — Sidebar Dark Mode (Expanded vs Collapsed)']:
    img_placeholder(cap)

heading('4.3 Implementasi Fitur Keamanan', 2)
para('4.3.1 Role-Based Access Control (RBAC). Middleware CheckRole memvalidasi hak akses pada setiap request. Jika role tidak sesuai, sistem redirect ke dashboard yang tepat.', bold=True)
para('4.3.2 Session Management. Saat logout: session dihancurkan, CSRF token di-regenerasi, header no-cache ditambahkan untuk mencegah akses via tombol back browser.', bold=True)
para('4.3.3 Password Security. Password di-hash menggunakan Bcrypt. Tidak pernah disimpan plain text.', bold=True)

heading('4.4 Implementasi Tema Dark/Light Mode', 2)
add_table(['Aspek','Light Mode','Dark Mode'],[
    ['Background','#f1f5f9 (abu terang)','#0c1222 (biru tua)'],
    ['Sidebar','Putih + border abu','Navy gradient + gold'],
    ['Card','Putih + shadow','Slate gelap + glassmorphism'],
    ['Teks','#0f172a (hitam)','#f1f5f9 (putih)'],
    ['Aksen','Gold #d4af37','Gold #d4af37'],
])

heading('4.5 Hasil Pengujian (Black-Box Testing)', 2)

para('Tabel 4.1 — Pengujian Modul Autentikasi', bold=True)
add_table(['No','Skenario','Hasil yang Diharapkan','Status'],[
    ['1','Login kredensial valid','Redirect ke dashboard','Valid'],
    ['2','Login email salah','Pesan error muncul','Valid'],
    ['3','Login password salah','Pesan error muncul','Valid'],
    ['4','Login field kosong','Validasi required','Valid'],
    ['5','Logout','Session dihapus, redirect login','Valid'],
    ['6','Back button setelah logout','Redirect ke login + pesan','Valid'],
    ['7','Toggle tema login','Tampilan berubah dark/light','Valid'],
])

para('Tabel 4.2 — Pengujian Dashboard Admin', bold=True)
add_table(['No','Skenario','Hasil yang Diharapkan','Status'],[
    ['1','Menampilkan KPI','4 kartu KPI muncul','Valid'],
    ['2','Menampilkan grafik','Chart.js render grafik','Valid'],
    ['3','Toggle tema','Elemen berubah warna','Valid'],
    ['4','Collapse sidebar','Sidebar mengecil, tanpa scroll horizontal','Valid'],
])

para('Tabel 4.3 — Pengujian User Management', bold=True)
add_table(['No','Skenario','Hasil yang Diharapkan','Status'],[
    ['1','Daftar user','Tabel user ditampilkan','Valid'],
    ['2','Tambah user','User tersimpan + notifikasi','Valid'],
    ['3','Email duplikat','Pesan error validasi','Valid'],
    ['4','Reset password','Password direset + notifikasi','Valid'],
    ['5','Filter role','Tabel terfilter sesuai role','Valid'],
    ['6','Pencarian','Hasil pencarian sesuai','Valid'],
])

para('Tabel 4.4 — Pengujian Monitoring Aktivitas', bold=True)
add_table(['No','Skenario','Hasil yang Diharapkan','Status'],[
    ['1','Daftar aktivitas','Tabel paginasi 15/halaman','Valid'],
    ['2','Filter divisi','Terfilter sesuai divisi','Valid'],
    ['3','Filter status','Terfilter sesuai status','Valid'],
    ['4','Filter tanggal','Terfilter sesuai tanggal','Valid'],
    ['5','Grafik mingguan','Bar chart 7 hari muncul','Valid'],
])

para('Tabel 4.5 — Pengujian Modul Laporan', bold=True)
add_table(['No','Skenario','Hasil yang Diharapkan','Status'],[
    ['1','Buat laporan perusahaan','Tersimpan + log aktivitas','Valid'],
    ['2','Buat laporan divisi','Tersimpan + relasi divisi','Valid'],
    ['3','Hapus laporan','Dihapus + notifikasi','Valid'],
    ['4','Detail laporan','Modal popup muncul','Valid'],
])

para('Tabel 4.6 — Pengujian Aktivitas Karyawan', bold=True)
add_table(['No','Skenario','Hasil yang Diharapkan','Status'],[
    ['1','Catat aktivitas','Tersimpan + notifikasi','Valid'],
    ['2','Upload file','File tersimpan di storage','Valid'],
    ['3','Edit status','Status berubah di DB','Valid'],
    ['4','Filter aktivitas','Terfilter sesuai kriteria','Valid'],
])

heading('4.6 Pembahasan', 2)
heading('4.6.1 Kelebihan Sistem', 3)
para('1. Arsitektur Terstruktur: Pola MVC memastikan kode terorganisir dan mudah dipelihara.')
para('2. Keamanan Berlapis: RBAC, CSRF, password hashing, session invalidation, dan no-cache headers.')
para('3. User Experience Premium: Glassmorphism, gradient accent, dan micro-animations.')
para('4. Dual Theme System: Dark/light mode pada seluruh halaman termasuk login.')
para('5. Sidebar Responsif: Collapsible tanpa horizontal scroll.')

heading('4.6.2 Keterbatasan Sistem', 3)
para('1. Sistem masih berupa prototipe, belum production-ready.')
para('2. Menggunakan data dummy untuk demonstrasi.')
para('3. Berjalan pada lingkungan lokal, belum dikonfigurasi untuk deployment.')
para('4. Belum ada fitur backup otomatis dan export PDF.')

heading('4.6.3 Kesesuaian dengan Metode RAD', 3)
add_table(['Fase RAD','Implementasi','Status'],[
    ['Requirements Planning','15 use case, 2 aktor, 7 tabel','Selesai'],
    ['Workshop Design','ERD, Activity Diagram, UI/UX','Selesai'],
    ['Construction','Laravel 12, 9 controller, 7 model','Selesai'],
    ['Cutover','Black-Box Testing 30+ skenario valid','Selesai'],
])

# Save
output = r'd:\laragon\www\erp-prototype\Bab_3_dan_4_Laporan_KP.docx'
doc.save(output)
print(f'DOCX saved to: {output}')
