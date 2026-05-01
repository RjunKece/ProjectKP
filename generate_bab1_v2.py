from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()
for section in doc.sections:
    section.top_margin = Cm(3)
    section.bottom_margin = Cm(3)
    section.left_margin = Cm(4)
    section.right_margin = Cm(3)

style_normal = doc.styles['Normal']
style_normal.font.name = 'Times New Roman'
style_normal.font.size = Pt(12)
style_normal.font.color.rgb = RGBColor(0, 0, 0)
style_normal.paragraph_format.line_spacing = 1.5
style_normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
style_normal.element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

def hc(text, size=14):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text); r.bold = True; r.font.name = 'Times New Roman'; r.font.size = Pt(size); r.font.color.rgb = RGBColor(0,0,0)
    p.paragraph_format.line_spacing = 1.5
    return p

def hl(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(text); r.bold = True; r.font.name = 'Times New Roman'; r.font.size = Pt(12); r.font.color.rgb = RGBColor(0,0,0)
    p.paragraph_format.line_spacing = 1.5; p.paragraph_format.space_before = Pt(18)
    return p

def ap(text, indent=True):
    """Add paragraph with optional mixed bold formatting using **bold** markers."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY; p.paragraph_format.line_spacing = 1.5
    if indent: p.paragraph_format.first_line_indent = Cm(1.27)
    parts = text.split('**')
    for i, part in enumerate(parts):
        r = p.add_run(part); r.font.name = 'Times New Roman'; r.font.size = Pt(12); r.font.color.rgb = RGBColor(0,0,0)
        if i % 2 == 1: r.bold = True
    return p

def ap_italic(text, indent=True):
    """Add paragraph with mixed italic formatting using *italic* markers."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY; p.paragraph_format.line_spacing = 1.5
    if indent: p.paragraph_format.first_line_indent = Cm(1.27)
    parts = text.split('*')
    for i, part in enumerate(parts):
        r = p.add_run(part); r.font.name = 'Times New Roman'; r.font.size = Pt(12); r.font.color.rgb = RGBColor(0,0,0)
        if i % 2 == 1: r.italic = True
    return p

def li(text, num=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY; p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.left_indent = Cm(1.27); p.paragraph_format.first_line_indent = Cm(0)
    prefix = f"{num}. " if num else "\u2022 "
    parts = (prefix + text).split('**')
    for i, part in enumerate(parts):
        r = p.add_run(part); r.font.name = 'Times New Roman'; r.font.size = Pt(12); r.font.color.rgb = RGBColor(0,0,0)
        if i % 2 == 1: r.bold = True
    return p

def tbl(headers, rows):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = 'Table Grid'; table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = table.rows[0].cells[i]; c.text = ''; p = c.paragraphs[0]
        r = p.add_run(h); r.bold = True; r.font.name = 'Times New Roman'; r.font.size = Pt(11); r.font.color.rgb = RGBColor(0,0,0)
        p.paragraph_format.line_spacing = 1.5
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            c = table.rows[ri+1].cells[ci]; c.text = ''; p = c.paragraphs[0]
            r = p.add_run(val); r.font.name = 'Times New Roman'; r.font.size = Pt(11); r.font.color.rgb = RGBColor(0,0,0)
            p.paragraph_format.line_spacing = 1.5
    doc.add_paragraph()
    return table

# ===== BAB I =====
hc('BAB I')
hc('PENDAHULUAN')

# =============================================================================
# 1.1 LATAR BELAKANG
# =============================================================================
hl('1.1 Latar Belakang')

ap('Perkembangan teknologi informasi dan komunikasi di era digital telah '
   'mendorong transformasi fundamental dalam cara perusahaan mengelola '
   'operasional bisnisnya. Katuu (2020) dalam artikelnya yang dipublikasikan '
   'pada jurnal New Review of Information Networking menerangkan bahwa '
   'Enterprise Resource Planning (ERP) telah berevolusi dari sistem yang '
   'hanya menangani perencanaan kebutuhan material (MRP) pada era 1960-an '
   'menjadi platform terintegrasi yang mencakup seluruh aspek operasional '
   'perusahaan di era modern. Evolusi ini menunjukkan betapa pentingnya '
   'integrasi sistem informasi bagi keberlangsungan dan daya saing perusahaan '
   'di tengah persaingan bisnis yang semakin dinamis.')

ap('Sistem ERP merupakan perangkat lunak terintegrasi yang menghubungkan '
   'berbagai fungsi bisnis dalam satu platform terpadu. Lutfi et al. (2022) '
   'dalam penelitiannya yang dipublikasikan pada jurnal Sustainability (MDPI) '
   'menyatakan bahwa adopsi sistem ERP pada Usaha Kecil dan Menengah (UKM) '
   'memberikan dampak signifikan terhadap peningkatan efisiensi operasional, '
   'pengurangan biaya administratif, dan percepatan proses pengambilan '
   'keputusan manajerial. Penelitian tersebut mengidentifikasi bahwa faktor '
   'teknologi, organisasi, dan lingkungan secara bersama-sama mempengaruhi '
   'keberhasilan adopsi ERP pada perusahaan skala kecil dan menengah.')

ap('Abd Elmonem, Nasr, dan Geith (2017) dalam tinjauan sistematis yang '
   'dipublikasikan pada Future Computing and Informatics Journal memaparkan '
   'bahwa perkembangan teknologi cloud computing telah mengubah paradigma '
   'implementasi ERP dari sistem on-premise yang mahal menjadi solusi berbasis '
   'web yang lebih terjangkau dan fleksibel. Temuan ini memperkuat argumentasi '
   'bahwa pengembangan sistem ERP berbasis web merupakan pendekatan yang '
   'relevan dan strategis, khususnya bagi perusahaan skala menengah yang '
   'memiliki keterbatasan sumber daya untuk mengadopsi sistem ERP komersial '
   'berskala besar seperti SAP atau Oracle.')

ap('Daraghmi dan Daraghmi (2022) dalam artikel yang dipublikasikan pada '
   'IEEE Access mendemonstrasikan efektivitas penerapan metode Rapid '
   'Application Development (RAD) dalam menghasilkan sistem yang fungsional '
   'dan usable dalam waktu pengembangan yang singkat. Penelitian tersebut '
   'membuktikan bahwa pendekatan RAD yang dikombinasikan dengan partisipasi '
   'aktif pengguna mampu mempercepat siklus pengembangan secara signifikan '
   'tanpa mengorbankan kualitas perangkat lunak yang dihasilkan. Temuan ini '
   'sejalan dengan kebutuhan proyek kerja praktik yang memiliki batasan '
   'waktu pengembangan relatif singkat.')

ap('Panorama Consulting Group (2024) dalam laporan tahunannya melaporkan '
   'bahwa lebih dari separuh organisasi yang mengimplementasikan sistem ERP '
   'berhasil menyelesaikan proyek sesuai timeline yang direncanakan, dengan '
   'median durasi proyek selama 15,5 bulan. Data ini menunjukkan bahwa '
   'implementasi ERP, meskipun kompleks, dapat dikelola dengan baik melalui '
   'pemilihan metode pengembangan yang tepat dan perencanaan yang matang. '
   'Dalam konteks pengembangan prototipe, penggunaan metode RAD memungkinkan '
   'pencapaian hasil yang fungsional dalam durasi yang jauh lebih singkat '
   'dibandingkan implementasi ERP skala penuh.')

ap('PT. Golden Intan Berlian merupakan perusahaan yang bergerak di bidang '
   'perdagangan dan distribusi yang memiliki delapan divisi operasional, '
   'yaitu Marketing, Sales, Keuangan, Konten Kreator, Gudang, CRM (Customer '
   'Relationship Management), YouTube, dan Admin Marketplace. Dengan jumlah '
   'divisi yang cukup banyak dan variasi fungsi bisnis yang beragam, '
   'pengelolaan data karyawan, monitoring aktivitas harian, serta pembuatan '
   'laporan manajerial menjadi tantangan tersendiri bagi manajemen perusahaan.')

ap('Saat ini, PT. Golden Intan Berlian masih mengelola sebagian besar proses '
   'bisnisnya secara manual atau menggunakan aplikasi yang berdiri sendiri '
   '(standalone) tanpa integrasi satu sama lain. Pengelolaan data karyawan '
   'dilakukan melalui spreadsheet, pencatatan aktivitas harian tidak '
   'terdokumentasi secara sistematis, dan pembuatan laporan manajerial '
   'memerlukan waktu yang lama karena harus mengkompilasi data dari berbagai '
   'sumber secara manual. Kondisi ini menyebabkan beberapa permasalahan '
   'operasional, antara lain: sulitnya melakukan monitoring aktivitas karyawan '
   'secara real-time, potensi terjadinya duplikasi dan inkonsistensi data '
   'antar divisi, serta lambatnya proses pengambilan keputusan manajerial '
   'karena keterbatasan akses terhadap informasi yang terintegrasi.')

ap('Permasalahan yang dihadapi PT. Golden Intan Berlian tersebut konsisten '
   'dengan temuan Haddara dan Fagerstr\u00f8m (2015) yang dipublikasikan pada '
   'Journal of Enterprise Resource Planning Studies, yang mengidentifikasi '
   'bahwa faktor-faktor utama yang mendorong adopsi ERP pada perusahaan '
   'mencakup kebutuhan akan integrasi data, efisiensi proses bisnis, dan '
   'kemampuan pelaporan yang lebih baik. Oleh karena itu, pengembangan '
   'sistem ERP berbasis web menjadi solusi yang tepat untuk mengatasi '
   'permasalahan yang ada.')

ap('Berdasarkan kondisi tersebut, diperlukan sebuah sistem ERP berbasis web '
   'yang dapat mengintegrasikan pengelolaan data karyawan, monitoring '
   'aktivitas, dan pembuatan laporan pada PT. Golden Intan Berlian. '
   'Pengembangan sistem ini menggunakan metode Rapid Application Development '
   '(RAD) yang diperkenalkan oleh James Martin pada tahun 1991. Menurut '
   'Kendall dan Kendall (2019), RAD merupakan pendekatan berorientasi objek '
   'untuk pengembangan sistem yang mencakup suatu metode pengembangan dan '
   'perangkat lunak. RAD bertujuan untuk memperpendek waktu yang diperlukan '
   'dalam siklus pengembangan sistem tradisional antara perancangan dan '
   'penerapan suatu sistem informasi.')

ap('Pemilihan metode RAD didasarkan pada beberapa pertimbangan utama. '
   'Pertama, proyek ini memiliki batasan waktu pengembangan yang relatif '
   'singkat sehingga memerlukan metode yang mampu menghasilkan prototipe '
   'fungsional secara cepat. Kedua, ketersediaan framework modern seperti '
   'Laravel 12 beserta ekosistem pendukungnya (TailwindCSS, Alpine.js, Vite, '
   'Chart.js) memungkinkan penerapan prinsip RAD melalui pemanfaatan komponen '
   'siap pakai (reusable components). Ketiga, sifat proyek yang berupa '
   'prototipe memerlukan pendekatan iteratif sebagaimana ditekankan dalam '
   'metode RAD. Pressman dan Maxim (2019) menegaskan bahwa RAD sangat cocok '
   'untuk proyek-proyek yang membutuhkan penyelesaian cepat dengan tetap '
   'memperhatikan kualitas perangkat lunak yang dihasilkan.')

ap('Sistem ERP yang dibangun merupakan sebuah prototipe (prototype) yang '
   'dirancang untuk menunjukkan bagaimana teknologi web modern dapat digunakan '
   'untuk mengembangkan sistem informasi terintegrasi pada perusahaan skala '
   'menengah. Sistem ini dibangun menggunakan framework Laravel 12 dengan '
   'arsitektur Model-View-Controller (MVC), basis data MySQL, dan antarmuka '
   'pengguna berbasis TailwindCSS 4.0 yang responsif dan modern. Diharapkan '
   'sistem ini dapat menjadi proof of concept yang memberikan gambaran nyata '
   'mengenai potensi digitalisasi proses bisnis di PT. Golden Intan Berlian.')

# =============================================================================
# 1.2 IDENTIFIKASI MASALAH
# =============================================================================
hl('1.2 Identifikasi Masalah')

ap('Berdasarkan latar belakang yang telah diuraikan di atas, dapat diidentifikasi '
   'beberapa permasalahan sebagai berikut:')
li('Pengelolaan data karyawan pada PT. Golden Intan Berlian masih dilakukan secara '
   'manual menggunakan spreadsheet, sehingga rentan terhadap kesalahan input dan '
   'inkonsistensi data antar divisi.', 1)
li('Belum adanya sistem monitoring aktivitas karyawan secara terpusat yang memungkinkan '
   'manajemen untuk memantau produktivitas dan kinerja karyawan dari delapan divisi '
   'secara real-time.', 2)
li('Proses pembuatan laporan manajerial masih memerlukan kompilasi data manual dari '
   'berbagai sumber yang terpisah, sehingga mengakibatkan lambatnya proses pengambilan '
   'keputusan strategis.', 3)
li('Belum tersedianya sistem informasi terintegrasi berbasis web yang dapat diakses '
   'oleh seluruh pengguna (Super Admin dan Karyawan) sesuai dengan hak akses '
   'masing-masing.', 4)

# =============================================================================
# 1.3 RUMUSAN MASALAH
# =============================================================================
hl('1.3 Rumusan Masalah')

ap('Berdasarkan identifikasi masalah yang telah diuraikan, maka rumusan masalah dalam '
   'penelitian ini adalah sebagai berikut:')
li('Bagaimana merancang dan membangun sistem ERP berbasis web untuk PT. Golden Intan '
   'Berlian yang dapat mengintegrasikan pengelolaan data karyawan, monitoring aktivitas, '
   'dan pembuatan laporan manajerial menggunakan metode Rapid Application Development '
   '(RAD)?', 1)
li('Bagaimana mengimplementasikan arsitektur Model-View-Controller (MVC) pada framework '
   'Laravel 12 untuk membangun sistem ERP yang terstruktur, aman, dan mudah '
   'dikembangkan?', 2)
li('Bagaimana hasil pengujian fungsionalitas sistem ERP Prototype yang dibangun '
   'menggunakan metode black-box testing?', 3)

# =============================================================================
# 1.4 BATASAN MASALAH
# =============================================================================
hl('1.4 Batasan Masalah')

ap('Agar penelitian ini lebih terarah dan terfokus, maka ditetapkan batasan masalah '
   'sebagai berikut:')
li('Sistem yang dibangun merupakan prototipe (prototype) ERP yang berfokus pada tiga '
   'modul utama, yaitu: manajemen pengguna (User Management), monitoring aktivitas '
   '(Activity Monitoring), dan pusat laporan (Reports Center).', 1)
li('Sistem dikembangkan menggunakan framework Laravel 12 dengan bahasa pemrograman '
   'PHP 8.2, basis data MySQL, dan antarmuka pengguna menggunakan TailwindCSS 4.0, '
   'Alpine.js, dan Chart.js.', 2)
li('Sistem memiliki dua peran pengguna, yaitu Super Admin yang memiliki akses penuh '
   'terhadap seluruh fitur, dan Karyawan yang memiliki akses terbatas pada dashboard '
   'pribadi dan profil.', 3)
li('Metode pengembangan sistem yang digunakan adalah Rapid Application Development (RAD) '
   'dengan empat tahapan: Requirements Planning, User Design, Construction, dan '
   'Cutover.', 4)
li('Pengujian sistem dilakukan menggunakan metode black-box testing yang berfokus pada '
   'pengujian fungsionalitas sistem dari perspektif pengguna akhir.', 5)
li('Sistem berjalan di lingkungan server lokal (localhost) menggunakan Laragon sebagai '
   'local development server dan belum di-deploy ke server produksi.', 6)
li('Studi kasus penelitian dilakukan pada PT. Golden Intan Berlian dengan delapan '
   'divisi operasional: Marketing, Sales, Keuangan, Konten Kreator, Gudang, CRM, '
   'YouTube, dan Admin Marketplace.', 7)

# =============================================================================
# 1.5 TUJUAN PENELITIAN
# =============================================================================
hl('1.5 Tujuan Penelitian')

ap('Tujuan dari penelitian ini adalah sebagai berikut:')
li('Merancang dan membangun sistem ERP Prototype berbasis web untuk PT. Golden Intan '
   'Berlian yang mengintegrasikan pengelolaan data karyawan, monitoring aktivitas, '
   'dan pembuatan laporan manajerial menggunakan metode Rapid Application Development '
   '(RAD).', 1)
li('Mengimplementasikan arsitektur Model-View-Controller (MVC) pada framework Laravel 12 '
   'untuk menghasilkan sistem ERP yang terstruktur, aman, dan mudah dikembangkan '
   'lebih lanjut.', 2)
li('Melakukan pengujian fungsionalitas sistem menggunakan metode black-box testing '
   'untuk memastikan setiap fitur berjalan sesuai dengan kebutuhan yang telah '
   'ditetapkan.', 3)

# =============================================================================
# 1.6 MANFAAT PENELITIAN
# =============================================================================
hl('1.6 Manfaat Penelitian')

ap('Penelitian ini diharapkan dapat memberikan manfaat baik secara teoritis maupun '
   'praktis sebagai berikut:')

hl('1.6.1 Manfaat Teoritis')
li('Memberikan kontribusi dalam pengembangan ilmu pengetahuan di bidang sistem '
   'informasi, khususnya terkait implementasi metode Rapid Application Development '
   '(RAD) dalam pembangunan sistem ERP berbasis web.', 1)
li('Menjadi referensi bagi peneliti selanjutnya yang ingin mengembangkan sistem ERP '
   'menggunakan framework Laravel dan metode RAD.', 2)
li('Membuktikan efektivitas penerapan metode RAD dalam menghasilkan prototipe '
   'sistem informasi yang fungsional dalam waktu pengembangan yang relatif singkat.', 3)

hl('1.6.2 Manfaat Praktis')
li('**Bagi PT. Golden Intan Berlian** \u2014 Sistem ERP Prototype ini dapat menjadi '
   'landasan (proof of concept) untuk pengembangan sistem ERP yang lebih komprehensif '
   'di masa mendatang, serta memberikan gambaran mengenai potensi digitalisasi proses '
   'bisnis perusahaan.', 1)
li('**Bagi Manajemen Perusahaan** \u2014 Menyediakan dashboard analitik yang memudahkan '
   'manajemen dalam memantau aktivitas karyawan, mengelola data pengguna dari delapan '
   'divisi, dan menghasilkan laporan manajerial secara cepat dan akurat.', 2)
li('**Bagi Pengembang Perangkat Lunak** \u2014 Memberikan contoh implementasi nyata '
   'pengembangan sistem ERP menggunakan teknologi web modern (Laravel 12, TailwindCSS 4.0, '
   'Alpine.js, Chart.js, dan Vite 7) dengan penerapan best practices seperti arsitektur '
   'MVC, Eloquent ORM, dan CSRF protection.', 3)

# =============================================================================
# 1.7 SISTEMATIKA PENULISAN
# =============================================================================
hl('1.7 Sistematika Penulisan')

ap('Sistematika penulisan laporan penelitian ini disusun dalam lima bab dengan '
   'rincian sebagai berikut:')

ap('**BAB I PENDAHULUAN**')
ap('Bab ini berisi latar belakang masalah yang menguraikan urgensi pengembangan '
   'sistem ERP berbasis web, identifikasi masalah, rumusan masalah, batasan masalah, '
   'tujuan penelitian, manfaat penelitian, dan sistematika penulisan laporan.', False)

ap('**BAB II LANDASAN TEORI**')
ap('Bab ini memuat teori-teori yang menjadi dasar dalam penelitian, meliputi '
   'konsep Enterprise Resource Planning (ERP), metode Rapid Application Development '
   '(RAD), arsitektur Model-View-Controller (MVC), framework Laravel, dan '
   'teknologi-teknologi pendukung lainnya yang digunakan dalam pengembangan sistem.', False)

ap('**BAB III METODOLOGI PENELITIAN**')
ap('Bab ini menjelaskan metode pengembangan sistem yang digunakan yaitu Rapid '
   'Application Development (RAD) beserta tahapan-tahapannya, alat dan bahan '
   'penelitian, teknik pengumpulan data, dan teknik pengujian sistem.', False)

ap('**BAB IV HASIL DAN PEMBAHASAN**')
ap('Bab ini menguraikan hasil implementasi sistem ERP Prototype yang disusun '
   'berdasarkan empat fase metode RAD, meliputi: Requirements Planning, User Design, '
   'Construction, dan Cutover. Pembahasan mencakup implementasi basis data, '
   'antarmuka pengguna, arsitektur backend, pengujian fungsional, serta kelebihan '
   'dan keterbatasan sistem.', False)

ap('**BAB V PENUTUP**')
ap('Bab ini berisi kesimpulan dari hasil penelitian yang telah dilakukan serta '
   'saran-saran untuk pengembangan sistem ERP lebih lanjut di masa mendatang.', False)

# =============================================================================
# DAFTAR PUSTAKA
# =============================================================================
doc.add_page_break()
hc('DAFTAR PUSTAKA')

refs = [
    # Jurnal Open Access (terverifikasi, 10 tahun terakhir)
    'Abd Elmonem, M. A., Nasr, E. S., & Geith, M. H. (2017). Benefits and Challenges of Cloud ERP Systems \u2013 A Systematic Literature Review. Future Computing and Informatics Journal, 2(1), 1\u20139. https://doi.org/10.1016/j.fcij.2017.03.003',
    'Daraghmi, Y.-A., & Daraghmi, E.-Y. (2022). RAPD: Rapid and Participatory Application Development of Usable Systems During COVID19 Crisis. IEEE Access, 10, 93601\u201393614. https://doi.org/10.1109/ACCESS.2022.3203582',
    'Haddara, M., & Fagerstr\u00f8m, A. (2015). Cloud ERP Systems: Anatomy of Adoption Factors & Attitudes. Journal of Enterprise Resource Planning Studies, 2015, Article 521212. https://doi.org/10.5171/2015.521212',
    'Katuu, S. (2020). Enterprise Resource Planning: Past, Present, and Future. New Review of Information Networking, 25(1), 37\u201346. https://doi.org/10.1080/13614576.2020.1742770',
    'Lutfi, A., Alshira\u2019h, A. F., Alshirah, M. H., Al-Okaily, M., Alqudah, H., Saad, M., Ibrahim, N., & Abdelmaksoud, O. (2022). Antecedents and Impacts of Enterprise Resource Planning System Adoption among Jordanian SMEs. Sustainability, 14(6), 3508. https://doi.org/10.3390/su14063508',
    # Buku referensi utama
    'Kendall, K. E., & Kendall, J. E. (2019). Systems Analysis and Design (10th ed.). Pearson.',
    'Laudon, K. C., & Laudon, J. P. (2020). Management Information Systems: Managing the Digital Firm (16th ed.). Pearson.',
    'Martin, J. (1991). Rapid Application Development. Macmillan Publishing Co., Inc.',
    'Panorama Consulting Group. (2024). 2024 ERP Report. Diakses dari https://www.panorama-consulting.com/resource-center/erp-report/',
    'Pressman, R. S., & Maxim, B. R. (2019). Software Engineering: A Practitioner\u2019s Approach (9th ed.). McGraw-Hill Education.',
    'Sommerville, I. (2016). Software Engineering (10th ed.). Pearson.',
]

all_refs = sorted(refs, key=lambda x: x.split(',')[0].split('(')[0].strip().lower())

for ref_text in all_refs:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.left_indent = Cm(1.27)
    p.paragraph_format.first_line_indent = Cm(-1.27)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(ref_text)
    r.font.name = 'Times New Roman'; r.font.size = Pt(12); r.font.color.rgb = RGBColor(0,0,0)

doc.save(r'd:\laragon\www\erp-prototype\laporan_bab1_v2.docx')
print('File saved: laporan_bab1_v2.docx')
