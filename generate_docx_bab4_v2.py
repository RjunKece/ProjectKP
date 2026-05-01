from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()
for section in doc.sections:
    section.top_margin = Cm(3)
    section.bottom_margin = Cm(3)
    section.left_margin = Cm(4)
    section.right_margin = Cm(3)

style_normal = doc.styles['Normal']
style_normal.font.name = 'Times New Roman'
style_normal.font.size = Pt(12)
style_normal.font.color.rgb = RGBColor(0, 0, 0)
style_normal.paragraph_format.line_spacing = 1.5
style_normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
style_normal.element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

def hc(text, size=14):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text); r.bold = True; r.font.name = 'Times New Roman'; r.font.size = Pt(size); r.font.color.rgb = RGBColor(0,0,0)
    p.paragraph_format.line_spacing = 1.5
    return p

def hl(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(text); r.bold = True; r.font.name = 'Times New Roman'; r.font.size = Pt(12); r.font.color.rgb = RGBColor(0,0,0)
    p.paragraph_format.line_spacing = 1.5; p.paragraph_format.space_before = Pt(12)
    return p

def ap(text, indent=True):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY; p.paragraph_format.line_spacing = 1.5
    if indent: p.paragraph_format.first_line_indent = Cm(1.27)
    r = p.add_run(text); r.font.name = 'Times New Roman'; r.font.size = Pt(12); r.font.color.rgb = RGBColor(0,0,0)
    return p

def img_placeholder(num, caption):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(f'[Gambar 4.{num} {caption}]')
    r.font.name = 'Times New Roman'; r.font.size = Pt(12); r.font.color.rgb = RGBColor(128,128,128)
    r.italic = True
    
    pc = doc.add_paragraph()
    pc.alignment = WD_ALIGN_PARAGRAPH.CENTER; pc.paragraph_format.line_spacing = 1.5
    rc = pc.add_run(f'Gambar 4.{num} '); rc.bold = True; rc.font.name = 'Times New Roman'; rc.font.size = Pt(12); rc.font.color.rgb = RGBColor(0,0,0)
    rc2 = pc.add_run(caption); rc2.font.name = 'Times New Roman'; rc2.font.size = Pt(12); rc2.font.color.rgb = RGBColor(0,0,0)

def li(text, num=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY; p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.left_indent = Cm(1.27); p.paragraph_format.first_line_indent = Cm(0)
    prefix = f"{num}. " if num else "\u2022 "
    parts = (prefix + text).split('**')
    for i, part in enumerate(parts):
        r = p.add_run(part); r.font.name = 'Times New Roman'; r.font.size = Pt(12); r.font.color.rgb = RGBColor(0,0,0)
        if i % 2 == 1: r.bold = True
    return p

def tbl(headers, rows):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = 'Table Grid'; table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = table.rows[0].cells[i]; c.text = ''; p = c.paragraphs[0]
        r = p.add_run(h); r.bold = True; r.font.name = 'Times New Roman'; r.font.size = Pt(11); r.font.color.rgb = RGBColor(0,0,0)
        p.paragraph_format.line_spacing = 1.5
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            c = table.rows[ri+1].cells[ci]; c.text = ''; p = c.paragraphs[0]
            r = p.add_run(val); r.font.name = 'Times New Roman'; r.font.size = Pt(11); r.font.color.rgb = RGBColor(0,0,0)
            p.paragraph_format.line_spacing = 1.5
    doc.add_paragraph()
    return table

# ===== BAB IV =====
hc('BAB IV')
hc('HASIL DAN PEMBAHASAN')

# ===== 4.1 =====
hl('4.1 Gambaran Umum Sistem')
ap('Pada Bab IV ini diuraikan hasil perancangan dan implementasi Sistem Enterprise Resource Planning (ERP) sederhana berbasis web yang dikembangkan selama pelaksanaan Kerja Praktik di PT. Golden Intan Berlian. Sistem ini dirancang sebagai prototype pendukung proses bisnis perusahaan dengan tujuan mengintegrasikan fungsi-fungsi operasional ke dalam satu platform terpadu yang dapat diakses melalui browser.')
ap('Pengembangan sistem dilakukan dengan menggunakan metode Rapid Application Development (RAD), yang memungkinkan proses pembangunan aplikasi secara cepat dan iteratif. Metode RAD dipilih karena sesuai dengan keterbatasan waktu pelaksanaan Kerja Praktik, serta memungkinkan penyesuaian kebutuhan secara fleksibel berdasarkan umpan balik yang diperoleh selama proses pengembangan.')
ap('Sistem ERP Prototype ini dibangun menggunakan framework Laravel 12 dengan bahasa pemrograman PHP 8.2 dan berjalan di atas arsitektur Model-View-Controller (MVC). Basis data yang digunakan adalah MySQL, sementara antarmuka pengguna dikembangkan menggunakan Blade Template Engine dengan dukungan TailwindCSS 4.0 untuk styling responsif, Alpine.js untuk interaktivitas sisi klien, dan Chart.js untuk visualisasi data dalam bentuk grafik. Vite 7 digunakan sebagai build tool untuk pengelolaan aset frontend. Sistem dijalankan pada server lokal menggunakan Laragon sebagai lingkungan pengembangan.')
ap('Sistem ini dirancang untuk dua aktor utama, yaitu Super Admin dan Karyawan. Super Admin memiliki hak akses penuh terhadap seluruh fitur sistem, termasuk manajemen pengguna, monitoring aktivitas, dan pembuatan laporan. Sementara itu, aktor Karyawan dirancang untuk memiliki akses terbatas pada dashboard pribadi dan pencatatan aktivitas harian. Meskipun implementasi fitur untuk aktor Karyawan pada tahap prototype ini belum sepenuhnya lengkap, perancangan sistem telah mempertimbangkan peran tersebut untuk pengembangan di masa mendatang.')

# ===== 4.2 =====
hl('4.2 Analisis Sistem')

hl('4.2.1 Analisis Sistem Berjalan')
ap('Berdasarkan hasil observasi selama pelaksanaan Kerja Praktik, proses bisnis yang berjalan di PT. Golden Intan Berlian masih menggunakan metode konvensional dalam beberapa aspek operasionalnya. Pencatatan aktivitas karyawan, pengelolaan data pengguna, dan pembuatan laporan dilakukan secara manual menggunakan spreadsheet atau dokumen tercetak. Hal ini menimbulkan beberapa permasalahan, antara lain:')
li('**Inefisiensi waktu** \u2014 Proses pencarian data dan penyusunan laporan memerlukan waktu yang cukup lama karena harus dilakukan secara manual.', 1)
li('**Risiko kesalahan data** \u2014 Input data secara manual rentan terhadap human error seperti duplikasi data, kesalahan pengetikan, atau inkonsistensi format.', 2)
li('**Keterbatasan akses informasi** \u2014 Data yang tersebar di berbagai dokumen menyulitkan proses monitoring dan pengambilan keputusan secara real-time.', 3)
li('**Tidak adanya audit trail** \u2014 Tidak tersedia mekanisme pencatatan otomatis atas setiap perubahan dan aktivitas yang terjadi dalam proses bisnis.', 4)

hl('4.2.2 Analisis Sistem Usulan')
ap('Berdasarkan permasalahan pada sistem berjalan, diusulkan pengembangan Sistem ERP Prototype berbasis web yang mampu mengintegrasikan proses pengelolaan data pengguna, pencatatan aktivitas, dan pembuatan laporan ke dalam satu platform terpusat. Sistem usulan ini memiliki karakteristik sebagai berikut:')
li('**Sentralisasi data** \u2014 Seluruh data tersimpan dalam satu basis data MySQL yang terstruktur dan dapat diakses secara terpusat.', 1)
li('**Otomatisasi pencatatan aktivitas** \u2014 Setiap aksi yang dilakukan pengguna dalam sistem dicatat secara otomatis melalui service ActivityLogger.', 2)
li('**Manajemen pengguna berbasis role** \u2014 Sistem mendukung pembagian hak akses berdasarkan role (Super Admin dan Karyawan) sehingga keamanan data terjaga.', 3)
li('**Pembuatan laporan digital** \u2014 Laporan dapat dibuat, disimpan, dan diakses kembali secara digital melalui modul Reports Center.', 4)
li('**Antarmuka responsif dan modern** \u2014 Tampilan sistem dirancang responsif menggunakan TailwindCSS sehingga dapat diakses dari berbagai perangkat.', 5)

# ===== 4.3 =====
hl('4.3 Perancangan Sistem')
ap('Tahap perancangan sistem merupakan fase kedua dalam metode RAD setelah fase perencanaan kebutuhan (Requirements Planning). Pada tahap ini dilakukan pemodelan proses bisnis dan struktur data menggunakan diagram UML (Unified Modeling Language) serta Entity Relationship Diagram (ERD). Perancangan ini menjadi acuan dalam proses implementasi sistem.')

hl('4.3.1 Use Case Diagram')
ap('Use Case Diagram digunakan untuk memodelkan interaksi antara aktor dengan sistem. Dalam Sistem ERP Prototype ini terdapat dua aktor utama, yaitu Super Admin dan Karyawan. Setiap aktor memiliki hak akses dan fungsionalitas yang berbeda sesuai dengan perannya masing-masing dalam sistem.')
ap('Super Admin sebagai aktor dengan hak akses tertinggi dapat melakukan seluruh fungsionalitas sistem yang meliputi: login ke sistem, mengakses dashboard admin yang menampilkan Key Performance Indicator (KPI), mengelola data pengguna (tambah user dan reset password), melakukan monitoring aktivitas seluruh pengguna, membuat dan mengelola laporan (Reports), serta mengelola profil akun. Sementara itu, aktor Karyawan memiliki akses terbatas yang meliputi: login ke sistem, mengakses dashboard karyawan, melihat profil pribadi, dan melakukan logout. Meskipun pada tahap prototype ini fitur Karyawan belum diimplementasikan secara lengkap, use case untuk aktor tersebut tetap dimodelkan sebagai bagian dari perancangan sistem secara keseluruhan.')
img_placeholder(1, 'Use Case Diagram Sistem ERP')

hl('4.3.2 Activity Diagram')
ap('Activity Diagram digunakan untuk memodelkan alur kerja (workflow) dari proses-proses utama dalam sistem. Diagram aktivitas ini menggambarkan urutan langkah-langkah yang dilakukan oleh aktor dalam menjalankan fungsionalitas tertentu, termasuk percabangan kondisi dan alur alternatif.')
ap('Alur utama yang dimodelkan dalam Activity Diagram meliputi: (1) Proses Login, yang dimulai dari pengisian kredensial oleh pengguna, validasi oleh sistem, pengecekan role, hingga redirect ke dashboard yang sesuai; (2) Proses Manajemen User, yang mencakup alur penambahan user baru dengan validasi input dan pencatatan aktivitas; (3) Proses Monitoring Aktivitas, yang menggambarkan alur akses dan filter data aktivitas; serta (4) Proses Pembuatan Laporan, yang meliputi pemilihan jenis laporan, pengisian detail, preview, konfirmasi, dan penyimpanan.')
img_placeholder(2, 'Activity Diagram Sistem ERP')

hl('4.3.3 Entity Relationship Diagram (ERD)')
ap('Entity Relationship Diagram (ERD) digunakan untuk memodelkan struktur basis data serta hubungan antar entitas dalam sistem. ERD Sistem ERP Prototype ini menggambarkan lima entitas utama beserta atribut dan relasinya. Entitas-entitas tersebut adalah: users (menyimpan data pengguna), roles (menyimpan data hak akses), divisions (menyimpan data divisi perusahaan), activities (menyimpan log aktivitas), dan reports (menyimpan data laporan).')
ap('Relasi yang terdapat dalam ERD meliputi: relasi one-to-many antara tabel roles dan users (satu role dimiliki oleh banyak user), relasi one-to-many antara tabel divisions dan users (satu divisi memiliki banyak user), relasi one-to-many antara tabel users dan activities (satu user memiliki banyak aktivitas), serta relasi one-to-many antara tabel users dan reports (satu user dapat membuat banyak laporan). Seluruh relasi diimplementasikan menggunakan foreign key constraint pada MySQL.')
img_placeholder(3, 'Entity Relationship Diagram (ERD) Sistem ERP')

# ===== 4.4 =====
hl('4.4 Implementasi Sistem')
ap('Tahap implementasi merupakan fase ketiga dalam metode RAD yang disebut Construction Phase. Pada tahap ini, hasil perancangan diterjemahkan ke dalam kode program menggunakan teknologi yang telah ditentukan. Berikut adalah uraian implementasi setiap halaman dan fitur utama dalam Sistem ERP Prototype.')

hl('4.4.1 Halaman Login')
ap('Halaman login merupakan titik masuk utama (entry point) bagi seluruh pengguna untuk mengakses Sistem ERP. Halaman ini menampilkan formulir autentikasi yang terdiri dari dua field input, yaitu email dan password. Desain halaman login menggunakan layout dua kolom (split screen) dengan bagian kiri menampilkan formulir login dan bagian kanan menampilkan branding perusahaan PT. Golden Intan Berlian dengan skema warna emas (gold) sebagai identitas visual.')
ap('Fungsi utama halaman login adalah melakukan proses autentikasi pengguna menggunakan facade Auth::attempt() pada Laravel. Setelah kredensial divalidasi, sistem melakukan regenerasi session untuk mencegah serangan session fixation, kemudian mengarahkan pengguna ke halaman dashboard sesuai dengan role yang dimiliki. Apabila Super Admin yang login, sistem akan melakukan redirect ke /admin/dashboard, sedangkan Karyawan akan diarahkan ke /karyawan/dashboard. Jika kredensial tidak valid, sistem menampilkan pesan error "Email atau password salah."')
img_placeholder(4, 'Halaman Login Sistem ERP')

hl('4.4.2 Dashboard Admin')
ap('Dashboard Admin merupakan halaman utama yang ditampilkan setelah Super Admin berhasil melakukan login. Halaman ini berfungsi sebagai pusat informasi (control center) yang menyajikan ringkasan data operasional sistem secara real-time. Dashboard dirancang untuk memberikan gambaran menyeluruh tentang kondisi sistem melalui visualisasi data yang informatif dan mudah dipahami.')
ap('Dashboard Admin terdiri dari tiga komponen utama: (1) System Snapshot yang menampilkan enam Key Performance Indicator (KPI) meliputi Total Users, Active Employees, Total Activities, Today Activity, Productivity, dan Growth; (2) Activity Analytics yang menampilkan grafik line chart menggunakan library Chart.js untuk memvisualisasikan tren aktivitas karyawan dan super admin selama 12 bulan terakhir; serta (3) Quick Actions yang menyediakan pintasan navigasi cepat ke fitur Manage Users, Review Activities, dan Generate Reports.')
img_placeholder(5, 'Halaman Dashboard Admin')

hl('4.4.3 Dashboard Karyawan')
ap('Dashboard Karyawan merupakan halaman yang dirancang untuk pengguna dengan role karyawan setelah berhasil login. Halaman ini bertujuan untuk menyediakan antarmuka yang memungkinkan karyawan melihat informasi terkait tugas, aktivitas harian, dan profil pribadi mereka. Pada tahap prototype saat ini, dashboard karyawan telah memiliki halaman dasar dengan tampilan informasi pengguna yang sedang login.')
ap('Secara konseptual, dashboard karyawan dirancang untuk menyediakan fitur-fitur seperti: ringkasan aktivitas pribadi, pencatatan aktivitas harian, notifikasi dari admin, dan akses ke profil. Meskipun implementasi fitur-fitur tersebut belum sepenuhnya lengkap pada tahap prototype ini, arsitektur sistem telah disiapkan untuk mengakomodasi pengembangan di masa mendatang melalui routing terpisah dengan prefix /karyawan dan controller yang independen.')
img_placeholder(6, 'Halaman Dashboard Karyawan')

hl('4.4.4 Manajemen User')
ap('Halaman Manajemen User (User Management) merupakan fitur yang memungkinkan Super Admin untuk mengelola seluruh akun pengguna yang terdaftar dalam sistem. Fitur ini merupakan salah satu komponen inti dari modul Human Resource Management dalam konsep ERP.')
ap('Halaman ini terdiri dari beberapa komponen: (1) Summary Cards yang menampilkan statistik Total Users, Active Users, Super Admin, dan Pending Approval; (2) Filter dan Pencarian berupa toolbar untuk mencari pengguna berdasarkan nama atau email; (3) Tabel Data Pengguna yang menampilkan daftar pengguna dengan informasi nama, email, role, divisi, status, dan tombol aksi; (4) Tambah User melalui modal form dengan field nama, email, role, dan divisi, dimana password default ditetapkan sebagai "password123" yang di-hash menggunakan bcrypt; serta (5) Reset Password untuk mereset password pengguna ke default melalui modal konfirmasi.')
img_placeholder(7, 'Halaman Manajemen User')

hl('4.4.5 Monitoring Aktivitas')
ap('Halaman Monitoring Aktivitas berfungsi sebagai pusat pemantauan seluruh log aktivitas yang tercatat dalam sistem. Fitur ini memungkinkan Super Admin untuk melacak dan mengawasi setiap aksi yang dilakukan oleh pengguna, sehingga mendukung fungsi audit trail dan pengawasan operasional.')
ap('Komponen utama halaman ini meliputi: (1) KPI Summary berupa empat kartu indikator yang menampilkan Total Aktivitas, Aktivitas Hari Ini, Aktivitas Admin, dan Aktivitas Karyawan; (2) Filter Bar berupa toolbar pencarian dengan filter berdasarkan role, divisi, dan tanggal; serta (3) Activity Table berupa tabel yang menampilkan informasi user, role, divisi, deskripsi aktivitas, dan waktu dengan paginasi 15 item per halaman menggunakan fitur pagination bawaan Laravel.')
img_placeholder(8, 'Halaman Monitoring Aktivitas')

hl('4.4.6 Halaman Reports')
ap('Halaman Reports Center merupakan pusat pembuatan dan pengelolaan laporan strategis dalam sistem ERP. Fitur ini memungkinkan Super Admin untuk menghasilkan berbagai jenis laporan yang mendukung proses pengambilan keputusan manajemen.')
ap('Halaman ini terdiri dari: (1) KPI Strip yang menampilkan lima indikator berupa Total Reports, Generated This Month, Scheduled, System Coverage, dan Storage Used; (2) Report Builder dengan empat modul pembuatan laporan yaitu Financial Intelligence, Department Performance, Growth & Trends, dan Custom Intelligence, dimana masing-masing modul membuka modal form untuk mengisi detail laporan; (3) Recent Reports berupa tabel laporan terbaru dengan informasi judul, tipe, periode, status, dan aksi (View/Download) dengan paginasi 10 item per halaman; serta (4) System Insight yang menampilkan statistik penggunaan laporan.')
img_placeholder(9, 'Halaman Reports Center')

hl('4.4.7 Halaman Profil Pengguna')
ap('Halaman Profil menampilkan informasi detail akun pengguna yang sedang login. Halaman ini dapat diakses oleh seluruh pengguna baik Super Admin maupun Karyawan. Fitur ini berfungsi untuk memberikan transparansi informasi akun kepada pengguna.')
ap('Komponen halaman profil meliputi: (1) Profile Overview berupa avatar inisial, nama, email, badge role, dan status aktif; (2) Account Information menampilkan detail role, divisi, dan tanggal registrasi; serta (3) Activity Summary yang menampilkan ringkasan total aktivitas dan aktivitas bulan berjalan.')
img_placeholder(10, 'Halaman Profil Pengguna')


# ===== 4.5 =====
hl('4.5 Implementasi Basis Data')
ap('Struktur basis data Sistem ERP Prototype dibangun menggunakan fitur migration pada Laravel yang memungkinkan pembuatan dan pengelolaan skema basis data secara terprogram dan terversioning. Basis data menggunakan MySQL dengan lima tabel utama yang saling berelasi. Berikut adalah penjelasan detail setiap tabel beserta relasinya.')

hl('4.5.1 Tabel Roles')
ap('Tabel roles berfungsi menyimpan data hak akses pengguna dalam sistem. Tabel ini menjadi referensi utama untuk menentukan level akses setiap pengguna. Terdapat tiga role yang dikonfigurasi melalui RoleSeeder, yaitu: super_admin, admin, dan karyawan.')
tbl(['Kolom','Tipe Data','Keterangan'],[
    ['id','bigint unsigned (PK, AI)','Primary key, auto increment'],
    ['nama_role','varchar (unique)','Nama role: super_admin, admin, karyawan'],
    ['created_at','timestamp','Waktu pembuatan record'],
    ['updated_at','timestamp','Waktu pembaruan record'],
])

hl('4.5.2 Tabel Divisions')
ap('Tabel divisions menyimpan data divisi atau departemen yang ada di perusahaan. Tabel ini digunakan untuk mengelompokkan pengguna berdasarkan unit kerja. Terdapat delapan divisi yang dikonfigurasi melalui DivisionSeeder, yaitu: Marketing, Sales, Keuangan, Konten Kreator, Gudang, CRM, YouTube, dan Admin Marketplace.')
tbl(['Kolom','Tipe Data','Keterangan'],[
    ['id','bigint unsigned (PK, AI)','Primary key, auto increment'],
    ['nama_divisi','varchar (unique)','Nama divisi perusahaan'],
    ['created_at','timestamp','Waktu pembuatan record'],
    ['updated_at','timestamp','Waktu pembaruan record'],
])

hl('4.5.3 Tabel Users')
ap('Tabel users merupakan tabel inti yang menyimpan data seluruh pengguna sistem. Tabel ini memiliki relasi foreign key ke tabel roles dan divisions untuk menentukan hak akses dan unit kerja setiap pengguna.')
tbl(['Kolom','Tipe Data','Keterangan'],[
    ['id','bigint unsigned (PK, AI)','Primary key, auto increment'],
    ['name','varchar(255)','Nama lengkap pengguna'],
    ['email','varchar(255) (unique)','Alamat email sebagai identitas login'],
    ['email_verified_at','timestamp (nullable)','Waktu verifikasi email'],
    ['password','varchar(255) (hashed)','Kata sandi terenkripsi bcrypt'],
    ['role_id','bigint unsigned (FK)','Foreign key ke tabel roles'],
    ['division_id','bigint unsigned (FK, nullable)','Foreign key ke tabel divisions'],
    ['remember_token','varchar(100)','Token untuk fitur remember me'],
    ['created_at','timestamp','Waktu pembuatan akun'],
    ['updated_at','timestamp','Waktu pembaruan data akun'],
])

hl('4.5.4 Tabel Activities')
ap('Tabel activities berfungsi mencatat seluruh log aktivitas yang dilakukan pengguna dalam sistem. Tabel ini memiliki foreign key ke tabel users dengan constraint cascade on delete, yang berarti apabila data user dihapus maka seluruh aktivitas terkait juga akan terhapus secara otomatis.')
tbl(['Kolom','Tipe Data','Keterangan'],[
    ['id','bigint unsigned (PK, AI)','Primary key, auto increment'],
    ['user_id','bigint unsigned (FK, cascade delete)','Foreign key ke tabel users'],
    ['tanggal','datetime','Tanggal dan waktu aktivitas dilakukan'],
    ['deskripsi','text','Deskripsi detail aktivitas'],
    ['status','varchar (default: submitted)','Status aktivitas (submitted, created, dll)'],
    ['created_at','timestamp','Waktu pembuatan record'],
    ['updated_at','timestamp','Waktu pembaruan record'],
])

hl('4.5.5 Tabel Reports')
ap('Tabel reports menyimpan data laporan yang dibuat oleh Super Admin melalui modul Reports Center. Tabel ini mendukung empat jenis laporan yaitu financial, department, growth, dan custom.')
tbl(['Kolom','Tipe Data','Keterangan'],[
    ['id','bigint unsigned (PK, AI)','Primary key, auto increment'],
    ['title','varchar(255)','Judul laporan'],
    ['type','varchar','Jenis laporan (financial, department, growth, custom)'],
    ['description','text (nullable)','Deskripsi atau isi laporan'],
    ['start_date','date (nullable)','Tanggal mulai periode laporan'],
    ['end_date','date (nullable)','Tanggal akhir periode laporan'],
    ['status','enum (generated, ready, failed)','Status pemrosesan laporan'],
    ['created_by','bigint unsigned (FK, nullable)','Foreign key ke tabel users (pembuat)'],
    ['created_at','timestamp','Waktu pembuatan laporan'],
    ['updated_at','timestamp','Waktu pembaruan laporan'],
])

hl('4.5.6 Relasi Antar Tabel')
ap('Seluruh relasi antar tabel dalam sistem ERP ini menggunakan pola Eloquent Relationship pada Laravel. Hubungan antar entitas didefinisikan dalam model masing-masing menggunakan method belongsTo() dan hasMany(). Berikut adalah ringkasan relasi antar tabel:')
tbl(['Relasi','Tipe','Penjelasan'],[
    ['Role \u2192 Users','One-to-Many (hasMany)','Satu role dapat dimiliki oleh banyak user'],
    ['Division \u2192 Users','One-to-Many (hasMany)','Satu divisi dapat memiliki banyak user'],
    ['User \u2192 Role','Many-to-One (belongsTo)','Setiap user memiliki satu role'],
    ['User \u2192 Division','Many-to-One (belongsTo)','Setiap user memiliki satu divisi (nullable)'],
    ['User \u2192 Activities','One-to-Many (hasMany)','Satu user dapat memiliki banyak aktivitas'],
    ['Activity \u2192 User','Many-to-One (belongsTo)','Setiap aktivitas dimiliki oleh satu user'],
    ['User \u2192 Reports','One-to-Many','Satu user dapat membuat banyak laporan'],
    ['Report \u2192 User','Many-to-One (belongsTo)','Setiap laporan dibuat oleh satu user'],
])
ap('Seluruh foreign key diimplementasikan dengan constraint yang sesuai. Foreign key user_id pada tabel activities menggunakan cascadeOnDelete() sehingga data aktivitas otomatis terhapus ketika user dihapus. Sementara foreign key created_by pada tabel reports menggunakan nullOnDelete() sehingga data laporan tetap tersimpan meskipun user pembuat dihapus.')

# ===== 4.6 =====
hl('4.6 Implementasi Backend (MVC Laravel)')

hl('4.6.1 Arsitektur Model-View-Controller')
ap('Sistem ERP Prototype ini mengimplementasikan pola arsitektur Model-View-Controller (MVC) yang merupakan standar bawaan framework Laravel. Pola MVC memisahkan logika aplikasi menjadi tiga komponen utama yang saling terhubung namun memiliki tanggung jawab yang berbeda:')
li('**Model** \u2014 Terdapat lima model utama yaitu User, Role, Division, Activity, dan Report. Setiap model mendefinisikan fillable attributes untuk mass assignment protection, casting untuk konversi tipe data otomatis, dan relasi Eloquent untuk menghubungkan antar entitas.', 1)
li('**View** \u2014 Menggunakan Blade Template Engine dengan sistem layout inheritance. Layout utama layouts/admin.blade.php menjadi template induk untuk seluruh halaman panel admin. TailwindCSS 4.0 dan Alpine.js digunakan untuk styling dan interaktivitas.', 2)
li('**Controller** \u2014 Terdapat delapan controller yang menangani logika bisnis, meliputi AdminDashboardController, UserController, ActivityController, ReportController, LoginController, ProfileController, DashboardController, dan AuthenticatedSessionController.', 3)

hl('4.6.2 Sistem Autentikasi')
ap('Proses autentikasi diimplementasikan melalui LoginController dengan alur sebagai berikut: (1) Pengguna mengakses halaman login melalui route GET /login; (2) Pengguna mengisi email dan password, kemudian mengirim form melalui route POST /login; (3) Sistem memvalidasi input dengan aturan required dan email; (4) Auth::attempt() memverifikasi kredensial terhadap data di tabel users; (5) Jika valid, session diregenerasi untuk keamanan dan pengguna diarahkan ke route /dashboard; (6) Route /dashboard memeriksa role pengguna dan melakukan redirect ke dashboard yang sesuai.')

hl('4.6.3 Sistem Routing')
ap('Routing sistem dikonfigurasi dalam file routes/web.php dengan pengelompokan berdasarkan middleware dan prefix. Terdapat empat kelompok route utama:')
tbl(['Kelompok Route','Prefix','Middleware','Keterangan'],[
    ['Auth Routes','/login','guest','Route untuk halaman login (GET dan POST)'],
    ['Admin Routes','/admin','auth','Dashboard, Users, Activities, Reports'],
    ['Karyawan Routes','/karyawan','auth','Dashboard karyawan'],
    ['Utility Routes','/','-','Profile, Logout, Root redirect'],
])

hl('4.6.4 Fitur Dark Mode')
ap('Sistem mendukung fitur dark mode yang diimplementasikan menggunakan Alpine.js dengan penyimpanan preferensi di localStorage browser. TailwindCSS dikonfigurasi dengan darkMode: \'class\' untuk mendukung styling kondisional berdasarkan tema aktif. Fitur ini meningkatkan kenyamanan pengguna dalam menggunakan sistem.')

# ===== 4.7 =====
hl('4.7 Pengujian Sistem')
ap('Pengujian sistem dilakukan menggunakan metode Black Box Testing, yaitu metode pengujian yang berfokus pada fungsionalitas sistem tanpa memperhatikan struktur internal kode program. Pengujian dilakukan untuk memastikan bahwa setiap fitur sistem berjalan sesuai dengan kebutuhan yang telah didefinisikan pada tahap perancangan.')

hl('4.7.1 Hasil Pengujian Black Box')
ap('Berikut adalah tabel hasil pengujian Black Box Testing pada Sistem ERP Prototype:', indent=False)
tbl(['No','Skenario Pengujian','Langkah Pengujian','Hasil yang Diharapkan','Status'],[
    ['1','Login Valid','Memasukkan email dan password yang benar, klik Sign In','Pengguna berhasil login dan diarahkan ke dashboard sesuai role','Berhasil'],
    ['2','Login Tidak Valid','Memasukkan email atau password yang salah, klik Sign In','Muncul pesan error "Email atau password salah"','Berhasil'],
    ['3','Redirect Super Admin','Login menggunakan akun Super Admin','Diarahkan ke halaman /admin/dashboard','Berhasil'],
    ['4','Redirect Karyawan','Login menggunakan akun Karyawan','Diarahkan ke halaman /karyawan/dashboard','Berhasil'],
    ['5','Tampilan Dashboard Admin','Mengakses halaman /admin/dashboard','Menampilkan KPI, grafik analytics, dan quick actions','Berhasil'],
    ['6','Menambahkan User Baru','Klik Add User, mengisi form lengkap, klik submit','User baru tersimpan dan muncul di daftar pengguna','Berhasil'],
    ['7','Validasi Form User','Submit form tambah user dengan field kosong','Muncul pesan validasi error pada field yang kosong','Berhasil'],
    ['8','Pencarian User','Mengetik nama/email pada kolom search','Tabel menampilkan hasil pencarian yang sesuai','Berhasil'],
    ['9','Reset Password User','Klik Reset Password pada user, konfirmasi reset','Password user berhasil direset ke default','Berhasil'],
    ['10','Monitoring Aktivitas','Mengakses halaman /admin/activities','Menampilkan log aktivitas dengan KPI dan paginasi','Berhasil'],
    ['11','Generate Report','Memilih jenis report, mengisi form, submit','Report berhasil disimpan dan tampil di daftar','Berhasil'],
    ['12','Melihat Detail Report','Klik View pada report di tabel recent reports','Menampilkan halaman detail report','Berhasil'],
    ['13','Akses Halaman Profil','Klik avatar pengguna, pilih Profile','Menampilkan informasi akun dan ringkasan aktivitas','Berhasil'],
    ['14','Toggle Dark Mode','Klik tombol toggle Dark/Light di header','Tampilan berubah sesuai tema yang dipilih','Berhasil'],
    ['15','Logout','Klik tombol Logout pada dropdown profil','Session dihapus dan diarahkan kembali ke halaman login','Berhasil'],
    ['16','Akses Tanpa Login','Mengakses URL /admin/dashboard tanpa login','Diarahkan kembali ke halaman login','Berhasil'],
    ['17','Duplikasi Email','Menambahkan user dengan email yang sudah terdaftar','Muncul pesan error validasi email unique','Berhasil'],
    ['18','Paginasi Data','Mengakses halaman activities dengan data >15','Data ditampilkan dengan pagination 15 per halaman','Berhasil'],
])

hl('4.7.2 Pengujian Validasi Input')
ap('Sistem menerapkan validasi input pada sisi server (server-side validation) menggunakan fitur Request Validation bawaan Laravel. Validasi ini memastikan bahwa data yang dikirimkan oleh pengguna memenuhi aturan yang telah ditetapkan sebelum diproses lebih lanjut. Berikut adalah aturan validasi yang diterapkan pada setiap form dalam sistem:')
tbl(['Form','Field','Aturan Validasi'],[
    ['Login','email','required, email'],
    ['Login','password','required'],
    ['Tambah User','name','required, string, max:255'],
    ['Tambah User','email','required, email, unique:users'],
    ['Tambah User','role_id','required, exists:roles,id'],
    ['Tambah User','division_id','required, exists:divisions,id'],
    ['Generate Report','title','required, string, max:255'],
    ['Generate Report','type','required, string'],
    ['Generate Report','report_date','required, date'],
    ['Generate Report','attachment','nullable, file, max:2048'],
])

# ===== 4.8 =====
hl('4.8 Pembahasan')
ap('Pada subbab ini diuraikan pembahasan mengenai hasil implementasi Sistem ERP Prototype yang telah dibangun selama pelaksanaan Kerja Praktik. Pembahasan mencakup analisis kelebihan dan kekurangan sistem berdasarkan hasil implementasi dan pengujian yang telah dilakukan.')

hl('4.8.1 Kelebihan Sistem')
ap('Berdasarkan hasil implementasi dan pengujian, Sistem ERP Prototype PT. Golden Intan Berlian memiliki beberapa kelebihan sebagai berikut:')
li('**Arsitektur Terstruktur** \u2014 Penggunaan pola MVC pada framework Laravel memastikan pemisahan yang jelas antara logika bisnis (Controller), tampilan antarmuka (View), dan pengelolaan data (Model), sehingga kode program lebih terorganisir, mudah dipelihara, dan dapat dikembangkan secara modular.', 1)
li('**Keamanan Data** \u2014 Sistem menerapkan beberapa mekanisme keamanan yang meliputi: hashing password menggunakan algoritma bcrypt, proteksi CSRF (Cross-Site Request Forgery) pada setiap form, regenerasi session setelah proses login, dan validasi input pada sisi server untuk mencegah injeksi data berbahaya.', 2)
li('**Antarmuka Responsif dan Modern** \u2014 Antarmuka pengguna dibangun menggunakan TailwindCSS 4.0 yang memastikan tampilan responsif di berbagai ukuran layar. Dukungan fitur dark mode meningkatkan kenyamanan dan fleksibilitas penggunaan sistem.', 3)
li('**Pencatatan Aktivitas Otomatis** \u2014 Setiap aksi penting yang dilakukan dalam sistem dicatat secara otomatis ke dalam tabel activities, sehingga tersedia audit trail yang komprehensif untuk keperluan monitoring dan evaluasi.', 4)
li('**Visualisasi Data Informatif** \u2014 Penggunaan Chart.js untuk menampilkan grafik tren aktivitas pada dashboard membantu pengambilan keputusan berbasis data secara visual dan intuitif.', 5)
li('**Pengembangan Cepat dengan RAD** \u2014 Penerapan metode RAD memungkinkan pengembangan sistem secara iteratif dan cepat, sesuai dengan keterbatasan waktu pelaksanaan Kerja Praktik.', 6)

hl('4.8.2 Kekurangan Sistem')
ap('Meskipun demikian, sebagai sebuah prototype, sistem ini juga memiliki beberapa kekurangan dan keterbatasan yang perlu diperhatikan secara realistis:')
li('**Status Prototype** \u2014 Sistem masih berupa prototype sehingga beberapa fitur, khususnya fitur untuk aktor Karyawan seperti pencatatan aktivitas harian dan dashboard karyawan, belum diimplementasikan secara lengkap dan memerlukan pengembangan lebih lanjut.', 1)
li('**Belum Terdapat Middleware Role** \u2014 Sistem belum memiliki middleware khusus untuk memvalidasi role pengguna pada setiap route admin. Saat ini pembatasan akses hanya berdasarkan redirect di route /dashboard, sehingga secara teoritis pengguna dengan role karyawan yang mengetahui URL admin dapat mengaksesnya.', 2)
li('**Penyimpanan File Lokal** \u2014 Fitur upload lampiran pada modul report masih menggunakan penyimpanan lokal (local storage) dan belum terintegrasi dengan layanan cloud storage seperti Amazon S3 atau Google Cloud Storage untuk skalabilitas yang lebih baik.', 3)
li('**Belum Terdapat Sistem Notifikasi** \u2014 Sistem belum memiliki fitur notifikasi real-time untuk memberitahu admin tentang aktivitas penting atau perubahan data dalam sistem.', 4)
li('**Ekspor Data Terbatas** \u2014 Fitur download report belum sepenuhnya diimplementasikan untuk menghasilkan file dalam format PDF atau Excel, yang merupakan kebutuhan umum dalam konteks pelaporan bisnis.', 5)
li('**Belum Dilakukan Performance Testing** \u2014 Pengujian yang dilakukan masih sebatas black box testing fungsional. Belum dilakukan pengujian performa (performance testing) untuk mengukur kemampuan sistem dalam menangani beban pengguna dalam jumlah besar.', 6)

hl('4.8.3 Kesesuaian dengan Metode RAD')
ap('Penerapan metode Rapid Application Development (RAD) dalam pengembangan Sistem ERP Prototype ini telah berjalan sesuai dengan tahapan yang didefinisikan. Fase Requirements Planning dilaksanakan melalui analisis kebutuhan dan observasi proses bisnis perusahaan. Fase User Design direalisasikan melalui perancangan diagram UML dan ERD. Fase Construction dilaksanakan melalui implementasi kode program menggunakan Laravel. Fase Cutover dilaksanakan melalui pengujian black box testing untuk memastikan fungsionalitas sistem.')
ap('Metode RAD terbukti efektif untuk pengembangan prototype ini karena memungkinkan iterasi yang cepat dalam proses pembangunan sistem. Namun demikian, keterbatasan waktu pelaksanaan Kerja Praktik menyebabkan beberapa fitur belum dapat diselesaikan secara sempurna dan memerlukan iterasi pengembangan tambahan di masa mendatang.')

doc.save(r'd:\laragon\www\erp-prototype\laporan_bab4_v2.docx')
print('File saved: laporan_bab4_v2.docx')
