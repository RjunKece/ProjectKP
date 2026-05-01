from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()
for s in doc.sections:
    s.top_margin=Cm(3); s.bottom_margin=Cm(3); s.left_margin=Cm(4); s.right_margin=Cm(3)
sn = doc.styles['Normal']
sn.font.name='Times New Roman'; sn.font.size=Pt(12); sn.font.color.rgb=RGBColor(0,0,0)
sn.paragraph_format.line_spacing=1.5; sn.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
sn.element.rPr.rFonts.set(qn('w:eastAsia'),'Times New Roman')

def hc(t,sz=14):
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run(t); r.bold=True; r.font.name='Times New Roman'; r.font.size=Pt(sz); r.font.color.rgb=RGBColor(0,0,0)
    p.paragraph_format.line_spacing=1.5
def hl(t):
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    r=p.add_run(t); r.bold=True; r.font.name='Times New Roman'; r.font.size=Pt(12); r.font.color.rgb=RGBColor(0,0,0)
    p.paragraph_format.line_spacing=1.5; p.paragraph_format.space_before=Pt(18)
def ap(t,indent=True):
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY; p.paragraph_format.line_spacing=1.5
    if indent: p.paragraph_format.first_line_indent=Cm(1.27)
    for i,part in enumerate(t.split('**')):
        r=p.add_run(part); r.font.name='Times New Roman'; r.font.size=Pt(12); r.font.color.rgb=RGBColor(0,0,0)
        if i%2==1: r.bold=True
def li(t,num=None):
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY; p.paragraph_format.line_spacing=1.5
    p.paragraph_format.left_indent=Cm(1.27); p.paragraph_format.first_line_indent=Cm(0)
    pf=f"{num}. " if num else "\u2022 "
    for i,part in enumerate((pf+t).split('**')):
        r=p.add_run(part); r.font.name='Times New Roman'; r.font.size=Pt(12); r.font.color.rgb=RGBColor(0,0,0)
        if i%2==1: r.bold=True
def cap(t):
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.line_spacing=1.5
    r=p.add_run(t); r.font.name='Times New Roman'; r.font.size=Pt(11); r.font.color.rgb=RGBColor(0,0,0); r.italic=True
def tbl(headers,rows):
    table=doc.add_table(rows=1+len(rows),cols=len(headers))
    table.style='Table Grid'; table.alignment=WD_TABLE_ALIGNMENT.CENTER
    for i,h in enumerate(headers):
        c=table.rows[0].cells[i]; c.text=''; p=c.paragraphs[0]
        r=p.add_run(h); r.bold=True; r.font.name='Times New Roman'; r.font.size=Pt(11); r.font.color.rgb=RGBColor(0,0,0); p.paragraph_format.line_spacing=1.5
    for ri,row in enumerate(rows):
        for ci,val in enumerate(row):
            c=table.rows[ri+1].cells[ci]; c.text=''; p=c.paragraphs[0]
            r=p.add_run(val); r.font.name='Times New Roman'; r.font.size=Pt(11); r.font.color.rgb=RGBColor(0,0,0); p.paragraph_format.line_spacing=1.5
    doc.add_paragraph()

# ===== BAB IV =====
hc('BAB IV')
hc('HASIL DAN PEMBAHASAN')

hl('Hasil Implementasi Sistem')
ap('Berdasarkan perancangan yang telah dilakukan pada Bab III, prototipe sistem ERP PT. Golden Intan Berlian berhasil diimplementasikan menggunakan framework Laravel 12 dengan arsitektur Model-View-Controller (MVC). Sistem berjalan pada PHP 8.3 dan MySQL 8.0 di lingkungan server lokal menggunakan Laragon. Pada bab ini akan diuraikan secara detail hasil implementasi sistem yang mencakup implementasi basis data, arsitektur MVC, antarmuka pengguna, fitur keamanan, serta hasil pengujian menggunakan metode Black-Box Testing.')

# --- Implementasi Basis Data ---
hl('Implementasi Basis Data')
ap('Basis data sistem ERP Prototype diimplementasikan menggunakan MySQL 8.0 melalui mekanisme migration Laravel. Mekanisme migration memungkinkan pembuatan dan modifikasi struktur tabel basis data secara terprogram dan terversikan, sehingga setiap perubahan skema dapat dilacak melalui sistem kontrol versi (Git). Total tujuh tabel utama berhasil dibuat dengan relasi foreign key dan penerapan cascade delete untuk menjaga integritas referensial data antar tabel.')
ap('Tabel 4.1 menampilkan ringkasan seluruh tabel yang terdapat dalam basis data sistem beserta jumlah kolom dan keterangan fungsi masing-masing tabel. Struktur ini dirancang untuk mendukung seluruh fungsionalitas sistem mulai dari manajemen pengguna, pencatatan aktivitas, hingga pengelolaan laporan.')
cap('Tabel 4.1 Ringkasan Basis Data')
tbl(['No','Nama Tabel','Jumlah Kolom','Keterangan'],[
    ['1','roles','3','Data hak akses'],
    ['2','divisions','3','8 divisi perusahaan'],
    ['3','users','9','Data pengguna + relasi'],
    ['4','activities','11','Log aktivitas harian'],
    ['5','daily_targets','10','Target harian per divisi'],
    ['6','reports','12','Laporan perusahaan/divisi'],
    ['7','report_responses','6','Tanggapan laporan'],
])

# --- Implementasi Arsitektur MVC ---
hl('Implementasi Arsitektur MVC')
ap('Arsitektur Model-View-Controller (MVC) diimplementasikan secara konsisten pada seluruh komponen sistem. Pemisahan tanggung jawab antara Model, View, dan Controller memastikan kode program terorganisir dengan baik, mudah dipelihara, dan memungkinkan pengembangan fitur secara modular.')

ap('**Model (Eloquent ORM)**')
ap('Lapisan Model diimplementasikan menggunakan Eloquent ORM yang menyediakan antarmuka berorientasi objek untuk berinteraksi dengan basis data. Terdapat tujuh model utama yang masing-masing merepresentasikan satu tabel pada basis data, yaitu: User, Role, Division, Activity, Report, ReportResponse, dan DailyTarget. Setiap model mendefinisikan relasi antar tabel (hasMany, belongsTo), atribut yang dapat diisi (fillable), serta query scope untuk memudahkan pencarian data.', False)

ap('**Controller**')
ap('Lapisan Controller bertanggung jawab menangani logika bisnis aplikasi. Terdapat sembilan controller yang mengelola seluruh fungsionalitas sistem. Tabel 4.2 menampilkan daftar controller beserta fungsi utamanya dalam menangani permintaan dari pengguna.', False)
cap('Tabel 4.2 Daftar Controller dan Fungsi Utama')
tbl(['Controller','Fungsi Utama'],[
    ['LoginController','Autentikasi login dengan validasi'],
    ['AdminDashboardController','KPI, grafik tren, analytics'],
    ['UserController','CRUD user, reset password'],
    ['ActivityController','Monitoring aktivitas karyawan'],
    ['ReportController','Kelola laporan perusahaan & divisi'],
    ['KaryawanDashboardController','Dashboard personal + grafik'],
    ['KaryawanActivityController','Catat aktivitas + upload file'],
    ['KaryawanReportController','Lihat & balas laporan'],
    ['ProfileController','Kelola profil & ubah password'],
])

ap('**View (Blade Template)**')
ap('Lapisan View diimplementasikan menggunakan Blade Template Engine yang merupakan mesin template bawaan Laravel. Sistem menggunakan tiga layout utama yang menjadi kerangka dasar tampilan, yaitu: layout admin untuk halaman-halaman Super Admin, layout karyawan untuk halaman-halaman karyawan, dan layout auth untuk halaman login. Setiap layout mendefinisikan struktur HTML, sidebar navigasi, header, dan area konten utama yang dapat diisi oleh halaman-halaman turunan melalui mekanisme template inheritance.', False)

# --- Implementasi Antarmuka ---
hl('Implementasi Antarmuka')
ap('Berikut adalah hasil implementasi antarmuka pengguna (user interface) dari sistem ERP Prototype yang dibangun menggunakan TailwindCSS 4.0 dengan dukungan fitur dark mode dan light mode pada seluruh halaman.')

hl('Halaman Login')
ap('Halaman login merupakan halaman pertama yang diakses oleh pengguna untuk masuk ke dalam sistem. Halaman ini diimplementasikan dengan desain split-screen yang menampilkan formulir login di sisi kanan dan branding perusahaan beserta informasi sistem di sisi kiri. Gambar 4.1 dan 4.2 menampilkan tampilan halaman login dalam mode gelap (dark mode) dan mode terang (light mode).')
cap('Gambar 4.1 Tampilan Login (Dark Mode)')
ap('[Gambar tampilan login dark mode disisipkan di sini]', False)
cap('Gambar 4.2 Tampilan Login (Light Mode)')
ap('[Gambar tampilan login light mode disisipkan di sini]', False)

ap('Gambar 4.3 menampilkan tampilan pesan error yang muncul ketika pengguna memasukkan kredensial yang salah. Pesan error ditampilkan dengan animasi smooth transition untuk memberikan umpan balik visual yang jelas kepada pengguna.')
cap('Gambar 4.3 Tampilan Pesan Error Login')
ap('[Gambar pesan error login disisipkan di sini]', False)
ap('Fitur-fitur yang diimplementasikan pada halaman login meliputi: split-screen design dengan branding perusahaan, toggle tema dark/light mode, animated error message dengan smooth transition, password visibility toggle, session message setelah logout berhasil, serta proteksi CSRF (Cross-Site Request Forgery) pada setiap formulir.')

hl('Dashboard Admin')
ap('Dashboard admin merupakan halaman utama yang ditampilkan setelah Super Admin berhasil login. Halaman ini menyajikan ringkasan informasi penting dalam bentuk KPI cards dan grafik tren aktivitas. Gambar 4.4 menampilkan bagian atas dashboard yang berisi empat kartu KPI (Total Karyawan, Total Aktivitas, Aktivitas Hari Ini, dan Laporan Aktif), sedangkan Gambar 4.5 menampilkan grafik tren aktivitas 12 bulan terakhir yang dirender menggunakan library Chart.js.')
cap('Gambar 4.4 Dashboard Admin (KPI Cards)')
ap('[Gambar dashboard admin KPI disisipkan di sini]', False)
cap('Gambar 4.5 Dashboard Admin (Grafik Tren)')
ap('[Gambar dashboard admin grafik disisipkan di sini]', False)

hl('User Management')
ap('Halaman User Management menyediakan fitur pengelolaan data pengguna sistem secara terpusat. Gambar 4.6 menampilkan halaman utama User Management yang berisi tabel daftar pengguna dengan fitur pencarian, filter berdasarkan role, dan paginasi. Gambar 4.7 menampilkan modal dialog untuk menambahkan pengguna baru dengan formulir yang mencakup nama, email, role, dan divisi.')
cap('Gambar 4.6 Halaman User Management')
ap('[Gambar user management disisipkan di sini]', False)
cap('Gambar 4.7 Modal Tambah User')
ap('[Gambar modal tambah user disisipkan di sini]', False)

hl('Monitoring Aktivitas')
ap('Halaman Monitoring Aktivitas memungkinkan Super Admin untuk memantau seluruh aktivitas yang dicatat oleh karyawan dari berbagai divisi. Gambar 4.8 menampilkan bagian atas halaman yang berisi KPI cards ringkasan aktivitas dan grafik bar chart aktivitas mingguan, sedangkan Gambar 4.9 menampilkan tabel detail aktivitas karyawan dengan fitur filter berdasarkan divisi, status, dan rentang tanggal.')
cap('Gambar 4.8 Monitoring Aktivitas (KPI & Grafik)')
ap('[Gambar monitoring KPI dan grafik disisipkan di sini]', False)
cap('Gambar 4.9 Tabel Monitoring Aktivitas')
ap('[Gambar tabel monitoring disisipkan di sini]', False)

hl('Halaman Reports')
ap('Halaman Reports menyediakan fitur pembuatan dan pengelolaan laporan baik untuk lingkup perusahaan maupun divisi. Gambar 4.10 menampilkan halaman Reports dari sisi admin yang memiliki akses untuk membuat, melihat, dan menghapus laporan. Gambar 4.11 menampilkan halaman Reports dari sisi karyawan yang dapat melihat dan membalas laporan. Gambar 4.12 menampilkan modal dialog pembuatan laporan baru dengan formulir yang mencakup judul, tipe, scope, prioritas, dan deskripsi laporan.')
cap('Gambar 4.10 Halaman Reports Admin')
ap('[Gambar reports admin disisipkan di sini]', False)
cap('Gambar 4.11 Halaman Reports Karyawan')
ap('[Gambar reports karyawan disisipkan di sini]', False)
cap('Gambar 4.12 Modal Buat Laporan Baru')
ap('[Gambar modal buat laporan disisipkan di sini]', False)

hl('Dashboard & Aktivitas Karyawan')
ap('Dashboard karyawan menyajikan ringkasan kinerja personal karyawan. Gambar 4.13 dan 4.14 menampilkan tampilan dashboard karyawan yang berisi KPI cards personal (total aktivitas, aktivitas bulan ini, tingkat penyelesaian), grafik tren aktivitas bulanan, dan daftar aktivitas terbaru. Tampilan ini dirancang untuk memberikan gambaran komprehensif mengenai produktivitas karyawan secara individual.')
cap('Gambar 4.13 Dashboard Karyawan (1)')
ap('[Gambar dashboard karyawan bagian 1 disisipkan di sini]', False)
cap('Gambar 4.14 Dashboard Karyawan (2)')
ap('[Gambar dashboard karyawan bagian 2 disisipkan di sini]', False)

ap('Gambar 4.15 dan 4.16 menampilkan halaman aktivitas karyawan yang berisi daftar seluruh aktivitas yang telah dicatat beserta status dan detail lampiran. Gambar 4.17 menampilkan modal pencatatan aktivitas baru dimana karyawan dapat mengisi deskripsi aktivitas, memilih target harian, mengunggah file lampiran, dan menyertakan tautan pendukung.')
cap('Gambar 4.15 Halaman Aktivitas Karyawan (1)')
ap('[Gambar aktivitas karyawan bagian 1 disisipkan di sini]', False)
cap('Gambar 4.16 Halaman Aktivitas Karyawan (2)')
ap('[Gambar aktivitas karyawan bagian 2 disisipkan di sini]', False)
cap('Gambar 4.17 Modal Pencatatan Aktivitas')
ap('[Gambar modal pencatatan aktivitas disisipkan di sini]', False)

# --- Implementasi Fitur Keamanan ---
hl('Implementasi Fitur Keamanan')
ap('Sistem ERP Prototype menerapkan beberapa lapisan keamanan untuk melindungi data dan memastikan hanya pengguna yang berwenang yang dapat mengakses fitur-fitur tertentu. Berikut adalah detail implementasi fitur keamanan yang diterapkan pada sistem.')

ap('**Role-Based Access Control (RBAC).**')
ap('Sistem menerapkan mekanisme kontrol akses berbasis peran menggunakan middleware CheckRole yang memvalidasi hak akses pengguna pada setiap permintaan HTTP. Jika role pengguna tidak sesuai dengan halaman yang diakses, sistem secara otomatis melakukan redirect ke dashboard yang sesuai dengan peran pengguna tersebut. Pendekatan ini memastikan bahwa Super Admin dan Karyawan hanya dapat mengakses fitur yang sesuai dengan hak aksesnya masing-masing.', False)

ap('**Session Management.**')
ap('Manajemen sesi diimplementasikan dengan mekanisme keamanan berlapis. Saat pengguna melakukan logout, sesi dihancurkan secara menyeluruh, token CSRF di-regenerasi untuk mencegah serangan replay, dan header no-cache ditambahkan pada respons HTTP untuk mencegah pengguna mengakses halaman yang telah di-cache melalui tombol back pada browser.', False)

ap('**Password Security.**')
ap('Keamanan password dijaga melalui penerapan algoritma hashing Bcrypt yang mengubah password menjadi hash satu arah sebelum disimpan ke dalam basis data. Password tidak pernah disimpan dalam bentuk teks biasa (plain text), sehingga meskipun basis data diakses oleh pihak yang tidak berwenang, password pengguna tetap terlindungi.', False)

# --- Implementasi Tema ---
hl('Implementasi Tema Dark/Light Mode')
ap('Sistem ERP Prototype mendukung dua mode tampilan, yaitu dark mode dan light mode, yang dapat dipilih oleh pengguna melalui toggle button yang tersedia di setiap halaman. Tabel 4.3 menampilkan perbandingan konfigurasi warna antara kedua mode tampilan untuk setiap elemen antarmuka utama.')
cap('Tabel 4.3 Perbandingan Tema Dark/Light Mode')
tbl(['Aspek','Light Mode','Dark Mode'],[
    ['Background','#f1f5f9 (abu terang)','#0c1222 (biru tua)'],
    ['Sidebar','Putih + border abu','Navy gradient + gold'],
    ['Card','Putih + shadow','Slate gelap + glassmorphism'],
    ['Teks','#0f172a (hitam)','#f1f5f9 (putih)'],
    ['Aksen','Gold #d4af37','Gold #d4af37'],
])

# --- Hasil Pengujian ---
hl('Hasil Pengujian (Black-Box Testing)')
ap('Pengujian sistem dilakukan menggunakan metode Black-Box Testing yang berfokus pada validasi fungsionalitas sistem dari perspektif pengguna akhir. Pengujian mencakup seluruh modul utama sistem dengan total lebih dari 30 skenario pengujian. Berikut adalah hasil pengujian untuk setiap modul.')

ap('Tabel 4.4 menampilkan hasil pengujian modul autentikasi yang mencakup skenario login dengan kredensial valid dan tidak valid, proses logout, penanganan tombol back setelah logout, dan toggle tema pada halaman login. Seluruh skenario menghasilkan status valid yang menunjukkan bahwa modul autentikasi berfungsi sesuai dengan spesifikasi kebutuhan.')
cap('Tabel 4.4 Pengujian Modul Autentikasi')
tbl(['No','Skenario','Hasil yang Diharapkan','Status'],[
    ['1','Login kredensial valid','Redirect ke dashboard','Valid'],
    ['2','Login email salah','Pesan error muncul','Valid'],
    ['3','Login password salah','Pesan error muncul','Valid'],
    ['4','Login field kosong','Validasi required','Valid'],
    ['5','Logout','Session dihapus, redirect login','Valid'],
    ['6','Back button setelah logout','Redirect ke login + pesan','Valid'],
    ['7','Toggle tema login','Tampilan berubah dark/light','Valid'],
])

ap('Tabel 4.5 menampilkan hasil pengujian dashboard admin yang memvalidasi tampilan KPI cards, rendering grafik Chart.js, fungsi toggle tema, dan perilaku sidebar yang dapat di-collapse.')
cap('Tabel 4.5 Pengujian Dashboard Admin')
tbl(['No','Skenario','Hasil yang Diharapkan','Status'],[
    ['1','Menampilkan KPI','4 kartu KPI muncul','Valid'],
    ['2','Menampilkan grafik','Chart.js render grafik','Valid'],
    ['3','Toggle tema','Elemen berubah warna','Valid'],
    ['4','Collapse sidebar','Sidebar mengecil, tanpa scroll horizontal','Valid'],
])

ap('Tabel 4.6 menampilkan hasil pengujian modul User Management yang mencakup operasi CRUD pengguna, validasi duplikasi email, reset password, serta fitur filter dan pencarian.')
cap('Tabel 4.6 Pengujian User Management')
tbl(['No','Skenario','Hasil yang Diharapkan','Status'],[
    ['1','Daftar user','Tabel user ditampilkan','Valid'],
    ['2','Tambah user','User tersimpan + notifikasi','Valid'],
    ['3','Email duplikat','Pesan error validasi','Valid'],
    ['4','Reset password','Password direset + notifikasi','Valid'],
    ['5','Filter role','Tabel terfilter sesuai role','Valid'],
    ['6','Pencarian','Hasil pencarian sesuai','Valid'],
])

ap('Tabel 4.7 menampilkan hasil pengujian modul monitoring aktivitas yang memvalidasi tampilan daftar aktivitas, fitur filter berdasarkan divisi, status, dan tanggal, serta rendering grafik mingguan.')
cap('Tabel 4.7 Pengujian Monitoring Aktivitas')
tbl(['No','Skenario','Hasil yang Diharapkan','Status'],[
    ['1','Daftar aktivitas','Tabel paginasi 15/halaman','Valid'],
    ['2','Filter divisi','Terfilter sesuai divisi','Valid'],
    ['3','Filter status','Terfilter sesuai status','Valid'],
    ['4','Filter tanggal','Terfilter sesuai tanggal','Valid'],
    ['5','Grafik mingguan','Bar chart 7 hari muncul','Valid'],
])

ap('Tabel 4.8 menampilkan hasil pengujian modul laporan yang mencakup pembuatan laporan untuk scope perusahaan dan divisi, penghapusan laporan, serta tampilan detail laporan melalui modal popup.')
cap('Tabel 4.8 Pengujian Modul Laporan')
tbl(['No','Skenario','Hasil yang Diharapkan','Status'],[
    ['1','Buat laporan perusahaan','Tersimpan + log aktivitas','Valid'],
    ['2','Buat laporan divisi','Tersimpan + relasi divisi','Valid'],
    ['3','Hapus laporan','Dihapus + notifikasi','Valid'],
    ['4','Detail laporan','Modal popup muncul','Valid'],
])

ap('Tabel 4.9 menampilkan hasil pengujian modul aktivitas karyawan yang memvalidasi pencatatan aktivitas baru, pengunggahan file lampiran, perubahan status aktivitas, dan fitur filter.')
cap('Tabel 4.9 Pengujian Aktivitas Karyawan')
tbl(['No','Skenario','Hasil yang Diharapkan','Status'],[
    ['1','Catat aktivitas','Tersimpan + notifikasi','Valid'],
    ['2','Upload file','File tersimpan di storage','Valid'],
    ['3','Edit status','Status berubah di DB','Valid'],
    ['4','Filter aktivitas','Terfilter sesuai kriteria','Valid'],
])

# --- Pembahasan ---
hl('Pembahasan')

hl('Kelebihan Sistem')
ap('Berdasarkan hasil implementasi dan pengujian yang telah dilakukan, sistem ERP Prototype yang dikembangkan memiliki beberapa kelebihan utama sebagai berikut:')
li('**Arsitektur Terstruktur** \u2014 Penerapan pola MVC pada framework Laravel memastikan kode program terorganisir dengan baik dan mudah dipelihara. Pemisahan tanggung jawab antara Model, View, dan Controller memudahkan proses debugging dan pengembangan fitur baru.', 1)
li('**Keamanan Berlapis** \u2014 Sistem menerapkan beberapa lapisan keamanan yang mencakup RBAC (Role-Based Access Control), proteksi CSRF, password hashing menggunakan Bcrypt, session invalidation pada saat logout, serta penambahan header no-cache untuk mencegah akses halaman melalui cache browser.', 2)
li('**User Experience Premium** \u2014 Antarmuka pengguna dirancang dengan estetika modern menggunakan teknik glassmorphism, gradient accent, dan micro-animations yang memberikan pengalaman pengguna yang menyenangkan dan profesional.', 3)
li('**Dual Theme System** \u2014 Sistem mendukung dua mode tampilan (dark mode dan light mode) yang dapat diaktifkan pada seluruh halaman termasuk halaman login, memberikan fleksibilitas kepada pengguna untuk memilih tampilan yang paling nyaman.', 4)
li('**Sidebar Responsif** \u2014 Sidebar navigasi dapat di-collapse untuk memberikan ruang tampilan yang lebih luas tanpa menimbulkan horizontal scroll, sehingga sistem tetap nyaman digunakan pada berbagai ukuran layar.', 5)

hl('Keterbatasan Sistem')
ap('Meskipun sistem telah berfungsi dengan baik sesuai hasil pengujian, terdapat beberapa keterbatasan yang perlu diperhatikan untuk pengembangan selanjutnya:')
li('Sistem masih berupa prototipe (prototype) yang belum sepenuhnya production-ready untuk digunakan dalam lingkungan operasional nyata.', 1)
li('Data yang digunakan pada sistem merupakan data dummy untuk keperluan demonstrasi dan pengujian, belum menggunakan data operasional perusahaan yang sesungguhnya.', 2)
li('Sistem berjalan pada lingkungan server lokal (localhost) menggunakan Laragon dan belum dikonfigurasi untuk deployment ke server produksi.', 3)
li('Beberapa fitur lanjutan seperti backup otomatis basis data dan ekspor laporan dalam format PDF belum diimplementasikan pada versi prototipe ini.', 4)

hl('Kesesuaian dengan Metode RAD')
ap('Tabel 4.10 menampilkan ringkasan kesesuaian implementasi sistem dengan empat fase metode Rapid Application Development (RAD). Tabel ini menunjukkan bahwa seluruh fase RAD telah dilaksanakan dengan baik, mulai dari perencanaan kebutuhan yang menghasilkan 15 use case dan 7 tabel basis data, perancangan iteratif menggunakan ERD dan Activity Diagram, konstruksi menggunakan Laravel 12 dengan 9 controller dan 7 model, hingga tahap cutover yang mencakup pengujian Black-Box Testing dengan lebih dari 30 skenario yang seluruhnya menghasilkan status valid.')
cap('Tabel 4.10 Kesesuaian dengan Metode RAD')
tbl(['Fase RAD','Implementasi','Status'],[
    ['Requirements Planning','15 use case, 2 aktor, 7 tabel','Selesai'],
    ['Workshop Design','ERD, Activity Diagram, UI/UX','Selesai'],
    ['Construction','Laravel 12, 9 controller, 7 model','Selesai'],
    ['Cutover','Black-Box Testing 30+ skenario valid','Selesai'],
])

doc.save(r'd:\laragon\www\erp-prototype\bab_4_update.docx')
print('File saved: bab_4_update.docx')
